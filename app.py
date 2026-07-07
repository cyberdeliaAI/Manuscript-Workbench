#!/usr/bin/env python3
"""
Manuscript Workbench — local Python edition.

Run:
  pip install -r requirements.txt
  python app.py

Then open http://127.0.0.1:8765
"""
from __future__ import annotations

import base64
import datetime as _dt
import errno
import html as html_std
import ipaddress
import io
import json
import os
import re
import sys
import tempfile
import traceback
import urllib.parse
from html.parser import HTMLParser
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    import requests
except Exception:  # pragma: no cover
    requests = None

try:
    import markdown as markdown_lib
except Exception:  # pragma: no cover
    markdown_lib = None

try:
    from lxml import html as lxml_html
except Exception:  # pragma: no cover
    lxml_html = None

try:
    import trafilatura
except Exception:  # pragma: no cover
    trafilatura = None

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except Exception:  # pragma: no cover
    Document = None
    WD_STYLE_TYPE = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None

APP_HOST = "127.0.0.1"
APP_PORT = int(os.environ.get("MANUSCRIPT_WORKBENCH_PORT") or os.environ.get("VELLUM_PREP_PORT", "8765"))
APP_NAME = "Manuscript Workbench — Python"
MAX_JSON_BYTES = 80 * 1024 * 1024
MAX_URL_IMPORT_BYTES = 12 * 1024 * 1024

PARAGRAPH_STYLE_MAP = {
    "chapter": "Heading 1",
    "part": "Heading 1",
    "subhead1": "Heading 2",
    "subhead2": "Heading 3",
    "subhead3": "Heading 4",
    "subhead4": "Heading 5",
    "subhead5": "Heading 6",
    "subtitle": "Manuscript Element Subtitle",
    "element-author": "Manuscript Element Author",
    "hidden-heading": "Manuscript Hidden Heading",
    "quote": "Manuscript Block Quote",
    "verse": "Manuscript Verse",
    "attribution": "Manuscript Attribution",
    "centered": "Manuscript Centered Text",
    "flush-left": "Manuscript Flush Left",
    "written-note": "Manuscript Written Note",
    "text-conversation": "Manuscript Text Conversation",
    "caption": "Caption",
    "inline-image": "Manuscript Inline Image",
    "ornament": "Normal",
    "scene": "Normal",
}

STYLE_TO_MARKER = {
    "Heading 1": "chapter",
    "Heading 2": "subhead1",
    "Heading 3": "subhead2",
    "Heading 4": "subhead3",
    "Heading 5": "subhead4",
    "Heading 6": "subhead5",
    "Quote": "quote",
    "Intense Quote": "quote",
    "Caption": "caption",
    "Manuscript Element Subtitle": "subtitle",
    "Manuscript Element Author": "element-author",
    "Manuscript Hidden Heading": "hidden-heading",
    "Manuscript Block Quote": "quote",
    "Manuscript Verse": "verse",
    "Manuscript Attribution": "attribution",
    "Manuscript Centered Text": "centered",
    "Manuscript Flush Left": "flush-left",
    "Manuscript Written Note": "written-note",
    "Manuscript Text Conversation": "text-conversation",
    "Manuscript Inline Image": "inline-image",
    "Vellum Element Subtitle": "subtitle",
    "Vellum Element Author": "element-author",
    "Vellum Hidden Heading": "hidden-heading",
    "Vellum Block Quote": "quote",
    "Vellum Verse": "verse",
    "Vellum Attribution": "attribution",
    "Vellum Centered Text": "centered",
    "Vellum Flush Left": "flush-left",
    "Vellum Written Note": "written-note",
    "Vellum Text Conversation": "text-conversation",
    "Vellum Inline Image": "inline-image",
}

EXPORT_PARAGRAPH_STYLES = [
    "Manuscript Element Subtitle",
    "Manuscript Element Author",
    "Manuscript Hidden Heading",
    "Manuscript Block Quote",
    "Manuscript Verse",
    "Manuscript Attribution",
    "Manuscript Centered Text",
    "Manuscript Flush Left",
    "Manuscript Written Note",
    "Manuscript Text Conversation",
    "Manuscript Inline Image",
]

EXPORT_CHARACTER_STYLES = [
    "Manuscript Small Caps",
    "Manuscript Monospace",
    "Manuscript Sans Serif",
]

CHAPTER_CANDIDATE_RE = re.compile(
    r"^\s*(?:(?:chapter|hoofdstuk|chapitre|cap[ií]tulo)\s+([0-9ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|een|twee|drie|vier|vijf|zes|zeven|acht|negen|tien)\b.*|(?:part|deel)\s+([0-9ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten|een|twee|drie|vier|vijf|zes|zeven|acht|negen|tien)\b.*|prologue|epilogue|acknowledg(?:e)?ments|about the author|copyright|dedication|preface|introduction|afterword|also by|inhoud|voorwoord|nawoord|dankwoord)\s*$",
    re.I,
)

PAGE_NUMBER_RE = re.compile(
    r"^\s*(?:[-—–]\s*)?(?:p\.?\s*|page\s+|pagina\s+)?(?:\d{1,4}|[ivxlcdm]{2,})(?:\s*(?:/|of|van)\s*\d{1,4})?(?:\s*[-—–])?\s*$",
    re.I,
)

ROMAN_VALUES = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
WORD_NUMS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6, "seven": 7,
    "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12, "thirteen": 13,
    "fourteen": 14, "fifteen": 15, "sixteen": 16, "seventeen": 17, "eighteen": 18,
    "nineteen": 19, "twenty": 20,
    "een": 1, "twee": 2, "drie": 3, "vier": 4, "vijf": 5, "zes": 6, "zeven": 7,
    "acht": 8, "negen": 9, "tien": 10,
}


def json_response(handler: BaseHTTPRequestHandler, payload: Any, status: int = 200) -> None:
    body = json.dumps(payload, ensure_ascii=False, indent=None).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def text_response(handler: BaseHTTPRequestHandler, text: str, status: int = 200, content_type: str = "text/plain; charset=utf-8") -> None:
    body = text.encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", content_type)
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def normalize_line_endings(text: str) -> Tuple[str, List[str]]:
    notes: List[str] = []
    if text.startswith("\ufeff"):
        text = text[1:]
        notes.append("BOM removed")
    if "\r\n" in text:
        text = text.replace("\r\n", "\n")
        notes.append("CRLF → LF")
    if "\r" in text:
        text = text.replace("\r", "\n")
        notes.append("CR → LF")
    return text, notes


def decode_text_bytes(data: bytes) -> Tuple[str, List[str]]:
    notes: List[str] = []
    if data.startswith(b"\xff\xfe"):
        text = data[2:].decode("utf-16-le", errors="replace")
        notes.append("Read as UTF-16 LE")
    elif data.startswith(b"\xfe\xff"):
        text = data[2:].decode("utf-16-be", errors="replace")
        notes.append("Read as UTF-16 BE")
    elif data.startswith(b"\xef\xbb\xbf"):
        text = data[3:].decode("utf-8", errors="replace")
        notes.append("Read as UTF-8 BOM")
    else:
        text = data.decode("utf-8", errors="replace")
        # Heuristic: a UTF-16 file without BOM usually contains many NUL bytes after UTF-8 decode.
        if text.count("\x00") > max(4, len(text) // 50):
            even_nuls = data[0::2].count(0)
            odd_nuls = data[1::2].count(0)
            try:
                if odd_nuls > even_nuls:
                    text = data.decode("utf-16-le", errors="replace")
                    notes.append("Read as UTF-16 LE by NUL-byte heuristic")
                else:
                    text = data.decode("utf-16-be", errors="replace")
                    notes.append("Read as UTF-16 BE by NUL-byte heuristic")
            except Exception:
                pass
        elif "\ufffd" in text:
            alt = data.decode("windows-1252", errors="replace")
            if alt.count("\ufffd") < text.count("\ufffd"):
                text = alt
                notes.append("Read as Windows-1252")
            else:
                notes.append("Read as UTF-8 with replacement characters")
        else:
            notes.append("Read as UTF-8")
    text, norm_notes = normalize_line_endings(text)
    notes.extend(norm_notes)
    return text, notes


def rtf_to_text(rtf: str) -> str:
    """Small, dependency-free RTF-to-text extractor. Good enough for manuscripts.

    It handles common paragraph, tab, escaped hex, and unicode tokens. It intentionally
    ignores most formatting and many RTF destinations.
    """
    ignorable_destinations = {
        "fonttbl", "colortbl", "stylesheet", "info", "pict", "object", "header", "footer",
        "generator", "themedata", "datastore", "xmlnstbl", "listtable", "listoverridetable",
        "revtbl", "rsidtbl", "latentstyles", "filetbl", "aftnsep", "aftnsepc", "aftncn",
    }
    stack: List[Tuple[bool, int]] = []
    out: List[str] = []
    i = 0
    n = len(rtf)
    skip_group = False
    ucskip = 1
    pending_ignorable = False

    def emit(s: str) -> None:
        if not skip_group:
            out.append(s)

    while i < n:
        c = rtf[i]
        if c == "{":
            stack.append((skip_group, ucskip))
            pending_ignorable = False
            i += 1
        elif c == "}":
            if stack:
                skip_group, ucskip = stack.pop()
            pending_ignorable = False
            i += 1
        elif c == "\\":
            i += 1
            if i >= n:
                break
            c2 = rtf[i]
            if c2 in "{}\\":
                emit(c2)
                i += 1
            elif c2 == "~":
                emit(" ")
                i += 1
            elif c2 == "-":
                emit("-")
                i += 1
            elif c2 == "_":
                emit("-")
                i += 1
            elif c2 == "*":
                pending_ignorable = True
                i += 1
            elif c2 == "'" and i + 2 < n:
                hx = rtf[i + 1:i + 3]
                try:
                    emit(bytes.fromhex(hx).decode("windows-1252", errors="replace"))
                except Exception:
                    pass
                i += 3
            elif c2.isalpha():
                start = i
                while i < n and rtf[i].isalpha():
                    i += 1
                word = rtf[start:i]
                sign = 1
                if i < n and rtf[i] == "-":
                    sign = -1
                    i += 1
                num_start = i
                while i < n and rtf[i].isdigit():
                    i += 1
                arg = None
                if i > num_start:
                    arg = int(rtf[num_start:i]) * sign
                if i < n and rtf[i] == " ":
                    i += 1

                if pending_ignorable or word in ignorable_destinations:
                    skip_group = True
                    pending_ignorable = False
                    continue
                pending_ignorable = False

                if word in ("par", "pard", "line"):
                    emit("\n")
                elif word == "tab":
                    emit("\t")
                elif word == "emdash":
                    emit("—")
                elif word == "endash":
                    emit("–")
                elif word == "bullet":
                    emit("•")
                elif word == "lquote":
                    emit("‘")
                elif word == "rquote":
                    emit("’")
                elif word == "ldblquote":
                    emit("“")
                elif word == "rdblquote":
                    emit("”")
                elif word == "uc" and arg is not None:
                    ucskip = max(0, arg)
                elif word == "u" and arg is not None:
                    codepoint = arg if arg >= 0 else arg + 65536
                    try:
                        emit(chr(codepoint))
                    except Exception:
                        pass
                    # Skip fallback chars/tokens after unicode escape.
                    skip = ucskip
                    while skip > 0 and i < n:
                        if rtf[i] == "\\":
                            # Skip an escaped fallback token as one char.
                            i += 1
                            if i < n and rtf[i] == "'" and i + 2 < n:
                                i += 3
                            elif i < n:
                                i += 1
                        else:
                            i += 1
                        skip -= 1
                # all other controls are formatting; ignore.
            else:
                # Control symbol; usually formatting. Consume it.
                i += 1
        else:
            if c not in "\r\n":
                emit(c)
            i += 1

    text = "".join(out)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip("\n")


class ManuscriptHTMLTextExtractor(HTMLParser):
    block_tags = {
        "address", "article", "blockquote", "body", "dd", "details", "dialog", "div", "dl", "dt",
        "fieldset", "figcaption", "figure", "h1", "h2", "h3", "h4", "h5", "h6",
        "hr", "li", "main", "ol", "p", "pre", "section", "table", "tbody", "td", "tfoot", "th", "thead",
        "tr", "ul",
    }
    skip_tags = {"script", "style", "noscript", "template", "head", "svg", "canvas", "nav", "header", "footer", "aside", "form", "menu"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self.skip_depth = 0

    def newline(self) -> None:
        if self.parts and not self.parts[-1].endswith("\n"):
            self.parts.append("\n")

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        tag = tag.lower()
        if tag in self.skip_tags:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag == "br":
            self.newline()
        elif tag in self.block_tags:
            self.newline()

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.skip_tags and self.skip_depth:
            self.skip_depth -= 1
            return
        if self.skip_depth:
            return
        if tag in self.block_tags:
            self.newline()

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        text = "".join(self.parts)
        text = text.replace("\xa0", " ")
        text = re.sub(r"[ \t\f\v]+", " ", text)
        text = re.sub(r" *\n *", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def html_to_text(source: str) -> str:
    if lxml_html is not None:
        try:
            source = readable_html_fragment(source)
        except Exception:
            pass
    parser = ManuscriptHTMLTextExtractor()
    parser.feed(source)
    parser.close()
    return parser.text()


def trafilatura_html_to_text(source: str, url: str = "") -> str:
    if trafilatura is None:
        raise RuntimeError("trafilatura is not installed. Run: pip install -r requirements.txt")
    extracted = trafilatura.extract(
        source or "",
        url=url or None,
        output_format="markdown",
        include_comments=False,
        include_tables=True,
        include_images=False,
        include_links=False,
        favor_precision=True,
    )
    if not extracted or not extracted.strip():
        raise ValueError("Trafilatura could not find readable main text.")
    return normalize_imported_text(extracted)


def normalize_imported_text(text: str) -> str:
    text, _notes = normalize_line_endings(text or "")
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t\f\v]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def html_import_options(source: str, filename: str = "import.html", url: str = "") -> Dict[str, Any]:
    simple = normalize_imported_text(html_to_text(source or ""))
    main = ""
    main_error = ""
    try:
        main = trafilatura_html_to_text(source or "", url=url)
    except Exception as exc:
        main_error = str(exc)
    notes = []
    if simple:
        notes.append("Simple extraction available")
    if main:
        notes.append("Trafilatura main text available")
    elif main_error:
        notes.append(f"Trafilatura unavailable: {main_error}")
    return {
        "filename": filename,
        "url": url,
        "simple": {"text": simple, "chars": len(simple)},
        "main": {"text": main, "chars": len(main), "error": main_error},
        "notes": notes,
        "engine": "trafilatura" if trafilatura is not None else "simple-only",
    }


def is_blocked_import_host(hostname: str) -> bool:
    host = (hostname or "").strip().lower().rstrip(".")
    if not host or host in {"localhost", "127.0.0.1", "::1"} or host.endswith(".localhost"):
        return True
    try:
        ip = ipaddress.ip_address(host)
        return ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved
    except ValueError:
        pass
    return False


def validate_import_url(raw_url: str) -> str:
    url = (raw_url or "").strip()
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError("Use a full http:// or https:// URL.")
    if parsed.username or parsed.password:
        raise ValueError("URLs with embedded usernames or passwords are not supported.")
    if is_blocked_import_host(parsed.hostname or ""):
        raise ValueError("Local, private, and internal addresses are blocked for URL import.")
    return urllib.parse.urlunparse(parsed)


def fetch_url_bytes(url: str) -> Tuple[bytes, str, str]:
    if requests is None:
        raise RuntimeError("requests is not installed. Run: pip install -r requirements.txt")
    safe_url = validate_import_url(url)
    headers = {
        "User-Agent": "ManuscriptWorkbench/1.0 (+local import)",
        "Accept": "text/html,application/xhtml+xml,text/plain;q=0.8,*/*;q=0.5",
    }
    with requests.get(safe_url, headers=headers, timeout=(8, 25), stream=True, allow_redirects=True) as resp:
        resp.raise_for_status()
        final_url = validate_import_url(resp.url)
        content_type = resp.headers.get("Content-Type", "")
        chunks: List[bytes] = []
        total = 0
        for chunk in resp.iter_content(chunk_size=65536):
            if not chunk:
                continue
            total += len(chunk)
            if total > MAX_URL_IMPORT_BYTES:
                raise ValueError("URL response is too large to import safely.")
            chunks.append(chunk)
    return b"".join(chunks), final_url, content_type


MARKDOWN_ALLOWED_TAGS = {
    "a", "blockquote", "br", "code", "del", "div", "em", "h1", "h2", "h3", "h4", "h5", "h6",
    "hr", "img", "li", "ol", "p", "pre", "span", "strong", "table", "tbody", "td", "th", "thead",
    "tr", "u", "ul",
}
MARKDOWN_VOID_TAGS = {"br", "hr", "img"}
MARKDOWN_ALLOWED_ATTRS = {
    "a": {"href", "title"},
    "img": {"src", "alt", "title"},
    "code": {"class"},
    "div": {"class"},
    "span": {"class"},
}
MARKDOWN_SAFE_CLASS_RE = re.compile(r"^(?:language-[A-Za-z0-9_-]+|block|block-label|block-[A-Za-z0-9_-]+|scene|sc|sans|task)$")
INLINE_MARKER_RE = re.compile(r"\[\[(sc|mono|sans|u):(.+?)\]\]", re.S)


def markdown_safe_url(value: str) -> str:
    raw = (value or "").strip()
    if not raw or re.search(r"[\s<>\"']", raw):
        return ""
    if re.match(r"^(?:https?:|mailto:|#|/|\./|\.\./)", raw, re.I):
        return raw
    return ""


class MarkdownHTMLSanitizer(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        tag = tag.lower()
        if tag not in MARKDOWN_ALLOWED_TAGS:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        clean_attrs: List[str] = []
        allowed = MARKDOWN_ALLOWED_ATTRS.get(tag, set())
        for name, value in attrs:
            name = name.lower()
            value = value or ""
            if name not in allowed:
                continue
            if name in {"href", "src"}:
                value = markdown_safe_url(value)
                if not value:
                    continue
            elif name == "class":
                classes = [c for c in value.split() if MARKDOWN_SAFE_CLASS_RE.fullmatch(c)]
                if not classes:
                    continue
                value = " ".join(classes)
            clean_attrs.append(f'{name}="{html_std.escape(value, quote=True)}"')
        attr_text = (" " + " ".join(clean_attrs)) if clean_attrs else ""
        self.parts.append(f"<{tag}{attr_text}>")

    def handle_startendtag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        self.handle_starttag(tag, attrs)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag not in MARKDOWN_ALLOWED_TAGS:
            if self.skip_depth:
                self.skip_depth -= 1
            return
        if self.skip_depth:
            return
        if tag not in MARKDOWN_VOID_TAGS:
            self.parts.append(f"</{tag}>")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(html_std.escape(data, quote=False))

    def handle_entityref(self, name: str) -> None:
        if not self.skip_depth:
            self.parts.append(f"&{name};")

    def handle_charref(self, name: str) -> None:
        if not self.skip_depth:
            self.parts.append(f"&#{name};")

    def html(self) -> str:
        return "".join(self.parts)


def sanitize_markdown_html(value: str) -> str:
    parser = MarkdownHTMLSanitizer()
    parser.feed(value or "")
    parser.close()
    return parser.html()


def render_inline_markers_for_preview(line: str) -> str:
    def repl(match: re.Match[str]) -> str:
        kind = match.group(1)
        content = html_std.escape(match.group(2) or "", quote=False)
        if kind == "u":
            return f"<u>{content}</u>"
        if kind == "mono":
            return f"<code>{content}</code>"
        return f'<span class="{kind}">{content}</span>'

    return INLINE_MARKER_RE.sub(repl, line)


def preprocess_manuscript_markdown(text: str) -> str:
    lines = (text or "").splitlines()
    out: List[str] = []
    open_block = False
    code_fence = ""
    for line in lines:
        fence = re.match(r"^\s*(```+|~~~+)", line)
        if code_fence:
            out.append(line)
            if fence and fence.group(1).startswith(code_fence[0]):
                code_fence = ""
            continue
        if fence:
            code_fence = fence.group(1)
            out.append(line)
            continue
        marker = re.match(r"^\s*:::\s*([A-Za-z0-9_-]+)?\s*$", line)
        if marker:
            if open_block:
                out.append("")
                out.append("</div>")
                open_block = False
            else:
                name = canonical_structure_marker(marker.group(1) or "block")
                safe_name = html_std.escape(name, quote=True)
                out.append(f'<div class="block block-{safe_name}" markdown="1">')
                out.append(f'<div class="block-label">{safe_name}</div>')
                out.append("")
                open_block = True
            continue
        if re.fullmatch(r"\s*(?:\*\s*){3,}\s*|\s*[-–—]{3,}\s*|\s*~{3,}\s*", line):
            out.append('<div class="scene">* * *</div>')
            continue
        list_indent = re.match(r"^( +)((?:[-+*])|(?:\d+[.)]))\s+", line)
        if list_indent and (len(list_indent.group(1)) % 4):
            depth = max(1, (len(list_indent.group(1)) + 1) // 2)
            line = (" " * (depth * 4)) + line[len(list_indent.group(1)):]
        out.append(render_inline_markers_for_preview(line))
    if open_block:
        out.append("")
        out.append("</div>")
    return "\n".join(out)


def render_markdown_html(text: str) -> str:
    if markdown_lib is None:
        raise RuntimeError("Python-Markdown is not installed. Run: pip install -r requirements.txt")
    rendered = markdown_lib.markdown(
        preprocess_manuscript_markdown(text or ""),
        extensions=["extra", "sane_lists", "nl2br"],
        output_format="html5",
    )
    return sanitize_markdown_html(rendered)


def readable_html_fragment(source: str) -> str:
    tree = lxml_html.fromstring(source)
    boilerplate_xpath = (
        "//script|//style|//noscript|//template|//svg|//canvas|//nav|//header|//footer|//aside|//form|//menu|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' sidebar ')]|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' navigation ')]|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' menu ')]|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' footer ')]|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' header ')]|"
        "//*[contains(translate(concat(' ', @id, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' sidebar ')]|"
        "//*[contains(translate(concat(' ', @id, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' navigation ')]|"
        "//*[contains(translate(concat(' ', @id, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' menu ')]|"
        "//*[contains(translate(concat(' ', @id, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' footer ')]|"
        "//*[contains(translate(concat(' ', @id, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' header ')]"
    )
    for el in tree.xpath(boilerplate_xpath):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    candidates = tree.xpath(
        "//main|//article|//*[@role='main']|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' content ')]|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' post ')]|"
        "//*[contains(translate(concat(' ', @class, ' '), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), ' entry ')]|"
        "//*[@id='content' or @id='main' or @id='article']"
    )
    if candidates:
        best = max(candidates, key=lambda el: len(" ".join(el.itertext()).strip()))
        return lxml_html.tostring(best, encoding="unicode", method="html")
    return lxml_html.tostring(tree, encoding="unicode", method="html")


def import_docx_bytes(data: bytes) -> Tuple[str, List[str]]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install -r requirements.txt")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(data)
        path = tmp.name
    try:
        doc = Document(path)
        lines: List[str] = []
        for p in doc.paragraphs:
            style_name = p.style.name if p.style is not None else ""
            marker = STYLE_TO_MARKER.get(style_name)
            text = paragraph_runs_to_marked_text(p)
            plain = text.strip()
            if not marker and plain and re.fullmatch(r"(?:[*•·]\s*){3,}|[-–—]{3,}|#", plain):
                # Practical round-trip for common ornamental/scene-break lines.
                marker = "ornament"
            if marker:
                lines.append(marker_to_markdown_text(marker, text))
            else:
                lines.append(text)
        text = "\n".join(lines)
        text, notes = normalize_line_endings(text)
        notes.insert(0, f"Imported DOCX paragraphs: {len(lines):,}")
        if getattr(doc, "tables", None):
            notes.append(f"Skipped DOCX tables: {len(doc.tables):,}")
        return text, notes
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass


def paragraph_runs_to_marked_text(p: Any) -> str:
    pieces: List[str] = []
    for run in p.runs:
        text = run.text or ""
        if not text:
            continue
        style_name = run.style.name if run.style is not None else ""
        if style_name in ("Manuscript Small Caps", "Vellum Small Caps"):
            pieces.append(f"[[sc:{text}]]")
        elif style_name in ("Manuscript Monospace", "Vellum Monospace"):
            pieces.append(f"[[mono:{text}]]")
        elif style_name in ("Manuscript Sans Serif", "Vellum Sans Serif"):
            pieces.append(f"[[sans:{text}]]")
        elif run.underline:
            pieces.append(f"[[u:{text}]]")
        elif run.bold and run.italic:
            pieces.append(f"***{text}***")
        elif run.bold:
            pieces.append(f"**{text}**")
        elif run.italic:
            pieces.append(f"*{text}*")
        else:
            pieces.append(text)
    return "".join(pieces)


def marker_to_markdown_text(marker: str, text: str) -> str:
    text = (text or "").strip()
    prefixes = {
        "chapter": "# ",
        "part": "# ",
        "subhead1": "## ",
        "subhead2": "### ",
        "subhead3": "#### ",
        "subhead4": "##### ",
        "subhead5": "###### ",
        "quote": "> ",
    }
    if marker in prefixes:
        return (prefixes[marker] + text).rstrip()
    if marker in ("scene", "ornament"):
        return "***"
    if not text:
        return f"::: {marker}\n:::"
    return f"::: {marker}\n{text}\n:::"


def import_file(filename: str, data: bytes, html_mode: str = "main") -> Dict[str, Any]:
    lower = filename.lower()
    notes: List[str] = []
    sniff = data[:4096].lstrip()
    is_docx = data.startswith(b"PK\x03\x04")
    is_rtf = sniff.startswith(b"{\\rtf")
    is_html = bool(re.match(br"(?is)^(?:<!doctype\s+html|<html\b|<!--.*?-->\s*<html\b|<head\b|<body\b)", sniff))
    if lower.endswith(".doc") and not lower.endswith(".docx"):
        raise ValueError(".doc is an old binary Word format. Please save it as .docx, .rtf, or .txt first.")
    if lower.endswith(".docx") or is_docx:
        text, notes = import_docx_bytes(data)
        if is_docx and not lower.endswith(".docx"):
            notes.insert(0, "Detected DOCX content from file signature")
    elif lower.endswith(".rtf") or is_rtf:
        decoded, notes = decode_text_bytes(data)
        text = rtf_to_text(decoded)
        notes.insert(0, "Imported RTF as plain text" if lower.endswith(".rtf") else "Detected RTF content and imported as plain text")
    elif lower.endswith((".html", ".htm")) or is_html:
        decoded, notes = decode_text_bytes(data)
        if html_mode == "simple":
            text = html_to_text(decoded)
            notes.insert(0, "Imported HTML with simple extraction" if lower.endswith((".html", ".htm")) else "Detected HTML content and imported with simple extraction")
        else:
            try:
                text = trafilatura_html_to_text(decoded)
                notes.insert(0, "Imported HTML main text with Trafilatura")
            except Exception as exc:
                text = html_to_text(decoded)
                notes.insert(0, f"Trafilatura failed; used simple HTML extraction: {exc}")
    elif lower.endswith((".md", ".markdown")):
        text, notes = decode_text_bytes(data)
        notes.insert(0, "Imported Markdown")
    else:
        text, notes = decode_text_bytes(data)
    return {"text": text, "notes": notes, "filename": filename}


def ensure_paragraph_style(doc: Any, name: str) -> None:
    try:
        _ = doc.styles[name]
        return
    except KeyError:
        pass
    if name.startswith("Heading") or name in ("Normal", "Caption"):
        return
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def ensure_character_style(doc: Any, name: str) -> None:
    try:
        _ = doc.styles[name]
        return
    except KeyError:
        pass
    style = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    if name == "Manuscript Monospace":
        style.font.name = "Courier New"
    elif name == "Manuscript Sans Serif":
        style.font.name = "Arial"
    elif name == "Manuscript Small Caps":
        style.font.small_caps = True


MARKDOWN_HEADING_TO_MARKER = {
    1: "chapter",
    2: "subhead1",
    3: "subhead2",
    4: "subhead3",
    5: "subhead4",
    6: "subhead5",
}

MARKDOWN_CONTAINER_ALIASES = {
    "note": "written-note",
    "written-note": "written-note",
    "text": "text-conversation",
    "text-conversation": "text-conversation",
    "verse": "verse",
    "quote": "quote",
    "attribution": "attribution",
    "centered": "centered",
    "flush-left": "flush-left",
    "caption": "caption",
    "image": "inline-image",
    "inline-image": "inline-image",
    "subtitle": "subtitle",
    "element-author": "element-author",
    "hidden-heading": "hidden-heading",
}

MARKDOWN_CONTAINER_RE = re.compile(r"^\s*:::\s*([a-z0-9_-]+)?\s*$", re.I)


def canonical_structure_marker(marker: str) -> str:
    return MARKDOWN_CONTAINER_ALIASES.get((marker or "").lower(), (marker or "").lower())


def is_part_heading(content: str) -> bool:
    return bool(re.match(r"^\s*(?:part|deel)\b", content or "", re.I))


def markdown_line_marker(line: str) -> Tuple[str, str]:
    stripped = line.strip()
    heading = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$", line)
    if heading:
        content = heading.group(2).strip()
        level = len(heading.group(1))
        if level == 1 and is_part_heading(content):
            return "part", content
        return MARKDOWN_HEADING_TO_MARKER.get(level, ""), content
    quote = re.match(r"^\s{0,3}>\s?(.*)$", line)
    if quote:
        return "quote", quote.group(1).strip()
    if re.fullmatch(r"(?:\*\s*){3,}|[-–—]{3,}|~{3,}", stripped):
        return "scene", ""
    return "", line


def structure_marker_for_line(line: str, block_marker: str = "") -> Tuple[str, str]:
    marker, content = markdown_line_marker(line)
    if marker:
        return marker, content
    if block_marker and line.strip():
        return block_marker, line.strip()
    return "", line


def strip_structure_line(line: str) -> str:
    marker, content = structure_marker_for_line(line)
    return content.rstrip() if marker else line.rstrip()


def create_manuscript_docx(text: str, metadata: Dict[str, str]) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install -r requirements.txt")
    doc = Document()
    for style_name in EXPORT_PARAGRAPH_STYLES:
        ensure_paragraph_style(doc, style_name)
    for style_name in EXPORT_CHARACTER_STYLES:
        ensure_character_style(doc, style_name)

    core = doc.core_properties
    if metadata.get("title"):
        core.title = metadata.get("title", "")
    if metadata.get("author"):
        core.author = metadata.get("author", "")
    if metadata.get("subtitle"):
        core.subject = metadata.get("subtitle", "")
    keywords = []
    for key in ("language", "series", "bookNumber", "publisher"):
        if metadata.get(key):
            keywords.append(f"{key}:{metadata[key]}")
    if keywords:
        core.keywords = "; ".join(keywords)
    core.comments = "Generated by Manuscript Workbench — Python"

    # Remove default empty paragraph if present.
    if doc.paragraphs and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    lines = text.split("\n")
    block_marker = ""
    for raw_line in lines:
        line = raw_line.rstrip("\r")
        container = MARKDOWN_CONTAINER_RE.match(line)
        if container:
            block_marker = canonical_structure_marker(container.group(1) or "")
            continue
        marker, content = structure_marker_for_line(line, block_marker)

        if marker == "scene":
            p = doc.add_paragraph("* * *", style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if marker == "ornament":
            content = content or "* * *"
        style_name = PARAGRAPH_STYLE_MAP.get(marker, "Normal")
        ensure_paragraph_style(doc, style_name)
        p = doc.add_paragraph(style=style_name)
        if marker in ("ornament", "centered"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif marker == "flush-left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        add_marked_runs(p, content, doc)

    if not doc.paragraphs:
        doc.add_paragraph("")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


INLINE_TOKEN_RE = re.compile(
    r"(\[\[(sc|mono|sans|u):(.+?)\]\]|\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)|`([^`\n]+?)`)",
    re.S,
)


def add_marked_runs(paragraph: Any, text: str, doc: Any) -> None:
    pos = 0
    for m in INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if m.group(2):
            kind = m.group(2)
            content = m.group(3) or ""
            run = paragraph.add_run(content)
            if kind == "u":
                run.underline = True
            else:
                style_name = {"sc": "Manuscript Small Caps", "mono": "Manuscript Monospace", "sans": "Manuscript Sans Serif"}[kind]
                ensure_character_style(doc, style_name)
                run.style = style_name
                if kind == "sc":
                    run.font.small_caps = True
        elif m.group(4) is not None:
            run = paragraph.add_run(m.group(4))
            run.bold = True
            run.italic = True
        elif m.group(5) is not None:
            run = paragraph.add_run(m.group(5))
            run.bold = True
        elif m.group(6) is not None:
            run = paragraph.add_run(m.group(6))
            run.italic = True
        elif m.group(7) is not None:
            ensure_character_style(doc, "Manuscript Monospace")
            run = paragraph.add_run(m.group(7))
            run.style = "Manuscript Monospace"
        else:
            paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def strip_vellum_markers(text: str) -> str:
    out: List[str] = []
    for line in text.split("\n"):
        if MARKDOWN_CONTAINER_RE.match(line):
            continue
        out.append(strip_structure_line(line))
    return "\n".join(out)


def roman_to_int(s: str) -> Optional[int]:
    s = s.upper()
    if not s or any(ch not in ROMAN_VALUES for ch in s):
        return None
    total = 0
    prev = 0
    for ch in reversed(s):
        val = ROMAN_VALUES[ch]
        if val < prev:
            total -= val
        else:
            total += val
            prev = val
    return total if 0 < total < 5000 else None


def parse_chapter_number(line: str) -> Optional[int]:
    clean = strip_structure_line(line).strip()
    m = re.search(r"\b(?:chapter|hoofdstuk)\s+([0-9]+|[ivxlcdm]+|[A-Za-z]+)\b", clean, re.I)
    if not m:
        return None
    token = m.group(1).lower()
    if token.isdigit():
        return int(token)
    if token in WORD_NUMS:
        return WORD_NUMS[token]
    return roman_to_int(token)


def preflight(text: str, metadata: Dict[str, str]) -> Dict[str, Any]:
    lines = text.split("\n")
    issues: List[Dict[str, Any]] = []
    info: List[Dict[str, Any]] = []

    def add(level: str, code: str, message: str, count: int = 1, line: Optional[int] = None, sample: Optional[str] = None) -> None:
        issues.append({"level": level, "code": code, "message": message, "count": count, "line": line, "sample": sample})

    marked_chapters: List[Tuple[int, str]] = []
    subheads = 0
    special = 0
    candidates: List[Tuple[int, str]] = []
    page_nums: List[Tuple[int, str]] = []
    tabs = 0
    trailing = 0
    extra_blank_runs = 0
    blank_run_len = 0
    repeated: Dict[str, List[int]] = {}
    title_toc_lines: List[Tuple[int, str]] = []
    empty_markers: List[Tuple[int, str]] = []
    block_fence_lines: List[int] = []

    for idx, line in enumerate(lines, start=1):
        container = MARKDOWN_CONTAINER_RE.match(line)
        if container:
            block_fence_lines.append(idx)
            continue
        marker, content = structure_marker_for_line(line)
        content = content.strip()
        if marker in ("chapter", "part"):
            marked_chapters.append((idx, content))
        elif marker.startswith("subhead"):
            subheads += 1
        elif marker:
            special += 1
        if marker and not content and marker not in ("scene", "ornament"):
            empty_markers.append((idx, marker))
        if not marker and CHAPTER_CANDIDATE_RE.match(line.strip()):
            candidates.append((idx, line.strip()))
        if PAGE_NUMBER_RE.match(line.strip()) and not marker and line.strip():
            page_nums.append((idx, line.strip()))
        tabs += line.count("\t")
        if re.search(r"[ \t]+$", line):
            trailing += 1
        if line.strip() == "":
            blank_run_len += 1
        else:
            if blank_run_len >= 3:
                extra_blank_runs += 1
            blank_run_len = 0
        low = line.strip().lower()
        if low in ("contents", "table of contents", "inhoud", "toc"):
            title_toc_lines.append((idx, line.strip()))
        if 4 <= len(line.strip()) <= 90 and not marker and not PAGE_NUMBER_RE.match(line.strip()):
            repeated.setdefault(re.sub(r"\s+", " ", line.strip()), []).append(idx)
    if blank_run_len >= 3:
        extra_blank_runs += 1

    docx_list_lines: List[Tuple[int, str]] = []
    table_blocks: List[Tuple[int, str]] = []
    table_separator_re = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
    for idx, line in enumerate(lines, start=1):
        marker, _content = structure_marker_for_line(line)
        if marker:
            continue
        if re.match(r"^\s*(?:[-+*]|\d+[.)])\s+\S", line):
            docx_list_lines.append((idx, line.strip()))
        if idx < len(lines) and "|" in line and table_separator_re.match(lines[idx]):
            table_blocks.append((idx, line.strip()))

    if not metadata.get("title"):
        add("warning", "missing-title", "Book title metadata is empty.")
    if not metadata.get("author"):
        add("warning", "missing-author", "Author metadata is empty.")
    if not marked_chapters:
        add("error", "no-chapters", "No chapters/parts are marked with Markdown # headings.")
    else:
        info.append({"label": "Markdown chapters/parts", "value": len(marked_chapters)})
    if subheads:
        info.append({"label": "Markdown subheads", "value": subheads})
    if special:
        info.append({"label": "Special structure blocks", "value": special})
    if candidates:
        add("warning", "chapter-candidates", "Possible chapter/front-matter lines are not marked.", len(candidates), candidates[0][0], candidates[0][1])
    if page_nums:
        add("warning", "page-numbers", "Standalone page-number-like lines are still present.", len(page_nums), page_nums[0][0], page_nums[0][1])
    if tabs:
        add("warning", "tabs", "Tabs are still present. DOCX imports are safer with styles instead of tabs.", tabs)
    if trailing:
        add("info", "trailing-whitespace", "Lines with trailing whitespace found.", trailing)
    if extra_blank_runs:
        add("warning", "extra-blank-runs", "Runs of 3+ blank lines found.", extra_blank_runs)
    if title_toc_lines:
        add("warning", "toc-detected", "Possible table of contents/title-page content detected. Most publishing/export tools create their own TOC.", len(title_toc_lines), title_toc_lines[0][0], title_toc_lines[0][1])
    if empty_markers:
        add("warning", "empty-markers", "Structure markers without text found.", len(empty_markers), empty_markers[0][0], empty_markers[0][1])
    if len(block_fence_lines) % 2:
        add("warning", "unclosed-structure-block", "A ::: structure block appears to be unclosed.", 1, block_fence_lines[-1], ":::")
    if docx_list_lines:
        add("warning", "docx-list-literal", "Markdown lists preview nicely, but DOCX export currently keeps list markers as literal text.", len(docx_list_lines), docx_list_lines[0][0], docx_list_lines[0][1])
    if table_blocks:
        add("warning", "docx-table-literal", "Markdown tables preview nicely, but DOCX export currently keeps table syntax as literal text.", len(table_blocks), table_blocks[0][0], table_blocks[0][1])

    repeated_candidates = [(k, v) for k, v in repeated.items() if len(v) >= 4 and v[-1] - v[0] >= 40]
    if repeated_candidates:
        repeated_candidates.sort(key=lambda kv: len(kv[1]), reverse=True)
        sample, positions = repeated_candidates[0]
        add("warning", "repeated-lines", "Repeated short lines may be headers/footers.", len(repeated_candidates), positions[0], sample)

    chapter_nums = [(line_no, parse_chapter_number(line)) for line_no, line in marked_chapters]
    chapter_nums = [(ln, num) for ln, num in chapter_nums if num is not None]
    if len(chapter_nums) >= 3:
        skips = []
        nums = [n for _, n in chapter_nums]
        for a, b in zip(nums, nums[1:]):
            if b != a + 1:
                skips.append((a, b))
        if skips:
            a, b = skips[0]
            add("warning", "chapter-number-sequence", f"Chapter numbering may skip or reset: {a} → {b}.", len(skips))

    invisible_count = len(re.findall(r"[\u200B\u200C\u200D\u2060\uFEFF\u00AD]", text))
    if invisible_count:
        add("warning", "invisible-chars", "Invisible/zero-width characters found.", invisible_count)
    control_count = len(re.findall(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F]", text))
    if control_count:
        add("warning", "control-chars", "Control/page-break characters found.", control_count)
    mojibake_count = len(re.findall(r"[\u00C2-\u00F4][\u0080-\u00BF\u20AC\u201A\u0192\u201E\u2026\u2020\u2021\u02C6\u2030\u0160\u2039\u0152\u017D\u2018\u2019\u201C\u201D\u2022\u2013\u2014\u02DC\u2122\u0161\u203A\u0153\u017E\u0178]", text))
    if mojibake_count:
        add("warning", "mojibake", "Possible mojibake/encoding damage found.", mojibake_count)

    levels = {"error": 0, "warning": 0, "info": 0}
    for issue in issues:
        levels[issue["level"]] = levels.get(issue["level"], 0) + 1
    return {"issues": issues, "summary": levels, "info": info, "lineCount": len(lines), "wordCount": len(re.findall(r"\S+", text))}


def get_json_body(handler: BaseHTTPRequestHandler) -> Dict[str, Any]:
    length = int(handler.headers.get("Content-Length", "0"))
    if length > MAX_JSON_BYTES:
        raise ValueError("Request is too large")
    data = handler.rfile.read(length)
    if not data:
        return {}
    return json.loads(data.decode("utf-8"))


def safe_filename(name: str, fallback: str = "manuscript") -> str:
    name = re.sub(r"[^A-Za-z0-9._-]+", "-", name or fallback).strip("-._")
    return name or fallback


AI_PROVIDER_BASE_URLS = {
    "lmstudio": "http://localhost:1234/v1",
    "omlx": "http://localhost:8000/v1",
    "ollama": "http://localhost:11434/v1",
    "custom": "http://localhost:1234/v1",
}


def normalize_base_url(base_url: str, provider: str = "lmstudio") -> str:
    default = AI_PROVIDER_BASE_URLS.get((provider or "").lower(), AI_PROVIDER_BASE_URLS["lmstudio"])
    base = (base_url or default).strip().rstrip("/")
    if base and "://" not in base:
        base = "http://" + base
    if not base.endswith("/v1"):
        # Accept either http://localhost:1234 or http://localhost:1234/v1
        parsed = urllib.parse.urlparse(base)
        if parsed.scheme and parsed.netloc and not parsed.path.rstrip("/").endswith("/v1"):
            base = base + "/v1"
    return base


def resolve_lm_model(base_url: str, model: str = "") -> str:
    if model.strip():
        return model.strip()
    if requests is None:
        return "local-model"
    try:
        resp = requests.get(f"{base_url}/models", timeout=6)
        resp.raise_for_status()
        data = resp.json()
        models = data.get("data") or []
        if models:
            return models[0].get("id") or "local-model"
    except Exception:
        pass
    return "local-model"


def ai_request_headers(api_key: str = "") -> Dict[str, str]:
    headers: Dict[str, str] = {}
    key = (api_key or "").strip()
    if key:
        headers["Authorization"] = f"Bearer {key}"
    return headers


HTML = r'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Manuscript Workbench — Python</title>
<style>
:root{
  --bg-deep:#171613; --bg-panel:#211f1b; --bg-raised:#2a2823; --bg-editor:#1c1b17;
  --line:#3a372f; --line-soft:#2c2a24; --text:#ede7da; --muted:#948c7d; --faint:#6b6558;
  --blue:#4fa3d1; --amber:#e0a458; --rust:#c1633e; --green:#6fa87a; --violet:#a986c9;
  --mono: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", monospace;
  --sans: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
}
*{box-sizing:border-box} html,body{height:100%} body{margin:0;background:var(--bg-deep);color:var(--text);font-family:var(--sans);display:flex;flex-direction:column;overflow:hidden}
header{display:grid;grid-template-columns:minmax(220px,1fr) auto;align-items:center;gap:14px;padding:12px 18px;border-bottom:1px solid var(--line);background:linear-gradient(180deg,#1c1a16,#171613)}
.brand{display:flex;flex-direction:column;gap:3px;min-width:0}.mark{font-family:var(--mono);font-weight:800;letter-spacing:.10em;text-transform:uppercase;font-size:14px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.mark span{color:var(--amber)}.sub{font-size:11px;color:var(--faint);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.actions{display:flex;align-items:center;justify-content:flex-end;gap:8px;min-width:0}.actions button,.actions .btn{flex:0 0 auto}.filename{font-family:var(--mono);font-size:12px;color:var(--muted);min-width:0;max-width:220px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
@media(max-width:980px){header{grid-template-columns:1fr}.actions{justify-content:flex-start;flex-wrap:wrap}.filename{max-width:min(100%,420px)}}
button,.btn,select,input[type=text],textarea{font-family:var(--sans)} button,.btn{font-size:12px;font-weight:700;color:var(--text);background:var(--bg-raised);border:1px solid var(--line);border-radius:6px;padding:7px 10px;cursor:pointer;transition:.12s}button:hover,.btn:hover{border-color:#356e8f;background:#31302a}button:disabled{opacity:.45;cursor:not-allowed}.primary{background:var(--blue);border-color:var(--blue);color:#0c1a22}.amber{background:var(--amber);border-color:var(--amber);color:#251704}.danger{border-color:#754333;color:#ffd7c5}.mini{font-size:11px;padding:5px 7px}.wide{width:100%}input[type=file]{display:none}
button.active{background:var(--amber);border-color:var(--amber);color:#241505}
main{flex:1;display:grid;grid-template-columns:minmax(0,1fr) 380px;grid-template-rows:minmax(0,1fr);gap:14px;padding:14px 18px 18px;min-height:0;overflow:hidden}@media(max-width:980px){main{grid-template-columns:1fr;grid-template-rows:none;overflow-y:auto}}
.card{background:var(--bg-panel);border:1px solid var(--line);border-radius:10px;display:flex;flex-direction:column;min-height:0;overflow:hidden}.editor-card{min-height:0}.statbar{display:flex;gap:14px;align-items:center;flex-wrap:wrap;padding:8px 14px;border-bottom:1px solid var(--line);font-family:var(--mono);font-size:11px;color:var(--muted)}.statbar b{color:var(--text)}.statbar .label{color:var(--faint);text-transform:uppercase;letter-spacing:.06em;margin-right:4px}
.quickbar{display:flex;align-items:center;gap:6px;flex-wrap:wrap;padding:8px 12px;border-bottom:1px solid var(--line);background:#1e1d19}.quickbar select{width:auto;min-width:150px;background:var(--bg-editor);color:var(--text);border:1px solid var(--line);border-radius:6px;padding:6px 8px;font-size:12px}.quickbar button{min-width:32px;height:32px;padding:5px 8px;display:inline-flex;align-items:center;justify-content:center}.quickbar .sep{width:1px;height:22px;background:var(--line);margin:0 3px}.quickbar .b{font-weight:900}.quickbar .i{font-style:italic}.quickbar .u{text-decoration:underline;text-underline-offset:3px}.quickbar .sc{font-variant:small-caps;letter-spacing:.04em}.quickbar .mono{font-family:var(--mono);font-size:11px}.toolbar{border-bottom:1px solid var(--line);background:#1e1d19}.toolbar-head{display:flex;align-items:center;justify-content:space-between;gap:8px;padding:8px 12px}.toolbar-title{font-family:var(--mono);font-size:11px;font-weight:800;letter-spacing:.1em;color:var(--faint);text-transform:uppercase}.toolbar-body{display:flex;flex-wrap:wrap;gap:6px;padding:0 12px 10px}.toolbar.collapsed .toolbar-body{display:none}.toolbar.collapsed{border-bottom:none}.toolbar .group-label{font-size:10px;color:var(--faint);align-self:center;margin-left:6px;text-transform:uppercase;letter-spacing:.08em}
.editor-wrap{flex:1;min-height:0;display:flex;flex-direction:column}.editor-area{flex:1;min-height:0;display:flex;background:var(--bg-editor);position:relative}.gutter{width:54px;flex:0 0 54px;overflow:hidden;background:#1a1916;border-right:1px solid var(--line-soft);font-family:var(--mono);font-size:12px;line-height:20px;padding:12px 0;color:var(--faint)}.gline{height:20px;padding-left:9px;position:relative}.gline:before{content:"";position:absolute;left:0;top:0;bottom:0;width:3px;background:transparent}.gline.issue:before,.gline.issue-amber:before{background:var(--amber)}.gline.issue-violet:before{background:var(--violet)}.gline.issue-rust:before{background:var(--rust)}.gline.issue-blue:before{background:var(--blue)}
#editor{flex:1;resize:none;border:0;outline:0;background:transparent;color:var(--text);font-family:var(--mono);font-size:12.5px;line-height:20px;padding:12px 14px;white-space:pre;overflow:auto;tab-size:4;position:relative;z-index:1}#editor.wrap,#editor:placeholder-shown{white-space:pre-wrap;overflow-wrap:break-word;overflow-x:hidden}.editor-area.wrap .gutter{display:none}#editor::placeholder{color:var(--faint);font-family:var(--sans);font-size:13px;white-space:pre-wrap}#editor::selection{background:rgba(224,164,88,.55)}.find-highlights{display:none;position:absolute;z-index:0;pointer-events:none;overflow:hidden;color:transparent;font-family:var(--mono);font-size:12.5px;line-height:20px;padding:12px 14px;white-space:pre;tab-size:4}.find-highlights.show{display:block}.find-highlights.wrap{white-space:pre-wrap;overflow-wrap:break-word}.find-hit{background:rgba(224,164,88,.32);border-radius:2px}.find-hit.current{background:rgba(79,163,209,.48);outline:1px solid rgba(79,163,209,.8)}.markdown-preview{display:none;flex:1;overflow:auto;border-left:1px solid var(--line);background:#f3f0e8;color:#20242a;padding:34px 52px;font-family:Georgia,"Times New Roman",serif;font-size:17px;line-height:1.65}.editor-card.preview-on #editor{flex:0 0 52%;border-right:1px solid var(--line-soft)}.editor-card.preview-on .markdown-preview{display:block}.markdown-preview h1,.markdown-preview h2,.markdown-preview h3,.markdown-preview h4,.markdown-preview h5,.markdown-preview h6{font-family:var(--sans);line-height:1.16;margin:1.2em 0 .55em;color:#20242a}.markdown-preview h1{font-size:30px}.markdown-preview h2{font-size:24px}.markdown-preview h3{font-size:20px}.markdown-preview h4{font-size:18px}.markdown-preview p{margin:.75em 0}.markdown-preview blockquote{border-left:4px solid #d8a04f;margin:1em 0;padding:.2em 0 .2em 1em;color:#4d463b}.markdown-preview ul,.markdown-preview ol{margin:.65em 0 .95em 1.45em;padding:0}.markdown-preview li{margin:.25em 0;padding-left:.18em}.markdown-preview li>ul,.markdown-preview li>ol{margin:.3em 0 .35em 1.25em}.markdown-preview .task{list-style:none}.markdown-preview .task input{margin-right:.45em}.markdown-preview code{font-family:var(--mono);font-size:.9em;background:#e4ded2;border-radius:4px;padding:.08em .28em}.markdown-preview pre{background:#e4ded2;border:1px solid #d4cbbc;border-radius:6px;padding:12px 14px;overflow:auto;line-height:1.45}.markdown-preview pre code{background:transparent;padding:0}.markdown-preview table{border-collapse:collapse;margin:1em 0;width:100%;font-family:var(--sans);font-size:.92em}.markdown-preview th,.markdown-preview td{border:1px solid #d2c8b6;padding:6px 8px;text-align:left}.markdown-preview th{background:#e4ded2}.markdown-preview a{color:#225f8a;text-decoration:underline}.markdown-preview img{max-width:100%;height:auto}.markdown-preview .scene{text-align:center;letter-spacing:.22em;color:#7b6d57;margin:1.4em 0}.markdown-preview .block{border-left:3px solid #c8b792;background:#ebe5d8;margin:1em 0;padding:.7em 1em}.markdown-preview .block-label{font-family:var(--mono);font-size:11px;text-transform:uppercase;letter-spacing:.08em;color:#7b6d57;margin-bottom:.35em}.markdown-preview .sc{font-variant:small-caps;letter-spacing:.04em}.markdown-preview .sans{font-family:var(--sans)}.markdown-preview .empty{color:#8b8171;font-family:var(--sans);font-size:14px}@media(max-width:760px){.editor-card.preview-on .editor-area{flex-direction:column}.editor-card.preview-on .gutter{display:none}.editor-card.preview-on #editor{flex:0 0 46%;border-right:0;border-bottom:1px solid var(--line)}.markdown-preview{padding:24px 24px}}
.editor-foot{display:grid;grid-template-columns:1fr auto 1fr;align-items:center;gap:12px;padding:10px 16px;border-top:1px solid var(--line)}.sel-info,.small{font-family:var(--mono);font-size:11px;color:var(--faint)}.sel-info{justify-self:end}.scope-toggle{justify-self:start;display:flex;gap:2px;background:var(--bg-raised);border:1px solid var(--line);border-radius:6px;padding:2px}.scope-toggle label{font-size:12px;font-weight:700;padding:6px 10px;border-radius:4px;cursor:pointer;color:var(--muted);white-space:nowrap}.scope-toggle input{display:none}.scope-toggle label:has(input:checked){background:var(--amber)}.scope-toggle label:has(input:checked) span{color:#241505}@media(max-width:760px){.editor-foot{grid-template-columns:1fr}.scope-toggle,.wrap-control,.sel-info{justify-self:start}}
.sidebar{display:flex;flex-direction:column;gap:16px;overflow-y:auto;min-height:0;height:100%;padding-right:4px;overscroll-behavior:contain}.sidebar .card{flex:0 0 auto}.panel-title{display:flex;align-items:center;justify-content:space-between;font-family:var(--mono);font-size:11px;font-weight:800;letter-spacing:.1em;text-transform:uppercase;color:var(--faint);padding:12px 16px 8px}.panel-body{padding:0 16px 12px}.row{display:flex;gap:8px;align-items:center;margin-bottom:8px}.row label{font-size:11px;color:var(--muted);min-width:82px}input[type=text],select,textarea.prompt{width:100%;background:#171613;border:1px solid var(--line);color:var(--text);border-radius:6px;padding:7px 8px;font-size:12px}.meta-grid{display:grid;grid-template-columns:1fr 1fr;gap:7px}.meta-grid label{display:flex;flex-direction:column;font-size:10px;color:var(--faint);gap:3px}.meta-grid input{font-family:var(--sans)}
.diag-row,.vellum-row{display:flex;align-items:center;gap:8px;padding:7px 16px;border-top:1px solid var(--line-soft)}.diag-row .lbl,.vellum-row .vtext{flex:1;min-width:0;font-size:12px;color:var(--muted);line-height:1.3;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}.diag-row .count{font-family:var(--mono);font-size:12px;font-weight:800;color:var(--text);min-width:24px;text-align:right}.diag-row .jump,.vellum-row .jump{font-family:var(--mono);font-size:11px;padding:4px 7px;border-radius:4px}.diag-row.zero .count{color:var(--green)}.diag-row.zero .jump{display:none}.diag-row.zero .lbl{color:var(--faint)}.dot.grey{background:var(--faint)}.dot.rust{background:var(--rust)}.dot.amber{background:var(--amber)}.dot.blue{background:var(--blue)}.dot.violet{background:var(--violet)}.vellum-help,.empty-hint{font-size:11px;line-height:1.4;color:var(--faint);padding:8px 16px 4px}.vellum-summary{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:6px;padding:10px 16px;border-top:1px solid var(--line-soft)}.vsum{background:var(--bg-raised);border:1px solid var(--line-soft);border-radius:6px;padding:6px 8px;font-family:var(--mono);font-size:11px;color:var(--muted)}.vsum b{color:var(--text)}.vellum-row .tag{font-family:var(--mono);font-size:10px;font-weight:800;color:#211507;background:var(--amber);border-radius:4px;padding:3px 5px;min-width:54px;text-align:center;white-space:nowrap}.vellum-row .tag.sub{background:var(--blue);color:#071720}.vellum-row .tag.supp{background:var(--violet);color:#170d20}.vellum-row .tag.warn{background:var(--rust);color:#210b05}.marker-legend{font-family:var(--mono);font-size:10.5px;line-height:1.45;color:var(--faint);padding:8px 16px 12px;border-top:1px solid var(--line-soft)}
.fix-group{border-top:1px solid var(--line-soft)}.fix-group-title{font-size:11px;font-weight:800;letter-spacing:.06em;text-transform:uppercase;color:var(--faint);padding:10px 16px 4px}.fix-row{display:flex;align-items:center;gap:8px;padding:5px 16px}.fix-row button{flex:1;text-align:left;font-size:12px;padding:8px 10px}.fix-row button.destructive{border-color:#6f5132}.fix-row button.destructive:after{content:" preview";float:right;color:var(--amber);font-family:var(--mono);font-size:10px;font-weight:800;letter-spacing:.04em;text-transform:uppercase}.fix-note{font-size:11px;color:var(--faint);padding:0 16px 8px;display:flex;align-items:flex-start;gap:6px;line-height:1.4}.fix-note input{margin:2px 0 0}.fix-note.caveat{color:#b58a5e}.lm-panel{border-top:1px solid var(--line-soft);padding:10px 16px 14px;display:flex;flex-direction:column;gap:10px}.ai-section{border:1px solid var(--line-soft);border-radius:8px;background:#1c1b17;overflow:hidden}.ai-section summary{cursor:pointer;list-style:none;padding:9px 10px;font-family:var(--mono);font-size:11px;font-weight:800;letter-spacing:.07em;text-transform:uppercase;color:var(--muted);display:flex;align-items:center;justify-content:space-between}.ai-section summary::-webkit-details-marker{display:none}.ai-section summary:after{content:"+";color:var(--faint);font-size:14px}.ai-section[open] summary{border-bottom:1px solid var(--line-soft);color:var(--text)}.ai-section[open] summary:after{content:"-"} .ai-section-body{padding:10px;display:flex;flex-direction:column;gap:10px}.lm-field{display:flex;flex-direction:column;gap:5px}.lm-field label{font-family:var(--mono);font-size:10.5px;letter-spacing:.06em;text-transform:uppercase;color:var(--faint)}.lm-field input,.lm-field select,.lm-field textarea{width:100%;background:var(--bg-editor);color:var(--text);border:1px solid var(--line);border-radius:6px;padding:8px 9px;font-family:var(--mono);font-size:11.5px}.lm-field textarea{min-height:160px;resize:vertical;line-height:1.45;white-space:pre-wrap}.lm-field textarea.compact{min-height:72px}.advanced-grid{display:grid;grid-template-columns:1fr 1fr;gap:8px}.lm-note{color:#b58a5e;font-size:11px;line-height:1.4}
.side-tabs{position:sticky;top:0;z-index:5;display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:4px;padding:0 0 8px;background:var(--bg-panel)}.side-tabs button{min-width:0;padding:7px 4px;font-size:11px;border-radius:6px}.side-tabs button.active{background:var(--amber);border-color:var(--amber);color:#241505}.side-panel{display:none;flex-direction:column;gap:16px}.side-panel.active{display:flex}.wrap-control{justify-self:center;display:flex;align-items:center;gap:7px;font-family:var(--mono);font-size:11px;color:var(--muted);white-space:nowrap}.wrap-control input{margin:0;accent-color:var(--amber)}
.issue-list{display:flex;flex-direction:column;gap:6px}.issue{border:1px solid var(--line-soft);border-radius:7px;padding:8px;background:#1c1b17}.issue .top{display:flex;align-items:center;gap:7px}.dot{width:8px;height:8px;border-radius:50%;background:var(--muted);flex-shrink:0}.dot.error{background:var(--rust)}.dot.warning{background:var(--amber)}.dot.info{background:var(--blue)}.issue .msg{font-size:12px;line-height:1.35}.issue .sample{font-family:var(--mono);font-size:11px;color:var(--muted);margin-top:5px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.issue button{margin-left:auto}.ok{color:var(--green)}.warn{color:var(--amber)}.err{color:#e89a7e}.hint{font-size:11px;line-height:1.45;color:var(--faint);margin:6px 0 0}.divider{height:1px;background:var(--line-soft);margin:10px -14px}.toast{position:fixed;left:50%;bottom:16px;transform:translateX(-50%);z-index:80;background:var(--bg-raised);border:1px solid #356e8f;border-radius:8px;padding:9px 14px;font-family:var(--mono);font-size:12px;opacity:0;pointer-events:none;transition:.2s;max-width:88vw}.toast.show{opacity:1;transform:translateX(-50%) translateY(-4px)}
.outline-list{border-top:1px solid var(--line-soft)}.outline-row{display:grid;grid-template-columns:auto minmax(0,1fr) auto;gap:8px;align-items:center;padding:8px 16px;border-top:1px solid var(--line-soft)}.outline-row:first-child{border-top:0}.outline-row.sub{padding-left:30px}.outline-tag{font-family:var(--mono);font-size:10px;font-weight:800;color:#211507;background:var(--amber);border-radius:4px;padding:3px 5px;min-width:34px;text-align:center}.outline-row.l1 .outline-tag{background:var(--amber);color:#211507}.outline-row.l2 .outline-tag{background:var(--blue);color:#071720}.outline-row.l3 .outline-tag{background:var(--green);color:#07190b}.outline-row.l4 .outline-tag{background:var(--violet);color:#170d20}.outline-row.l5 .outline-tag{background:var(--rust);color:#210b05}.outline-row.l6 .outline-tag{background:#777065;color:#11100d}.outline-main{min-width:0}.outline-title{font-size:12px;color:var(--muted);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.outline-meta{font-family:var(--mono);font-size:10.5px;color:var(--faint);margin-top:2px}.outline-actions{display:flex;gap:5px}.outline-actions button{font-family:var(--mono);font-size:10.5px;padding:4px 6px}.outline-toggle{width:24px;padding-left:0!important;padding-right:0!important}
.merge-list{flex:1;min-height:0;overflow:auto;padding:10px;display:flex;flex-direction:column;gap:6px}.merge-row{display:grid;grid-template-columns:minmax(0,1fr) auto;gap:8px;align-items:center;border:1px solid var(--line-soft);border-radius:7px;background:#181713;padding:8px 9px}.merge-name{font-size:12px;color:var(--text);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.merge-meta{font-family:var(--mono);font-size:10.5px;color:var(--faint);margin-top:2px}.merge-actions{display:flex;gap:5px}.merge-actions button{font-family:var(--mono);font-size:10.5px;padding:4px 6px}
.find-body{padding:12px;display:flex;flex-direction:column;gap:10px}.find-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px}.find-field{display:flex;flex-direction:column;gap:5px}.find-field label,.find-options label{font-family:var(--mono);font-size:10.5px;letter-spacing:.06em;text-transform:uppercase;color:var(--faint)}.find-field input{width:100%;background:var(--bg-editor);color:var(--text);border:1px solid var(--line);border-radius:6px;padding:8px 9px;font-family:var(--mono);font-size:12px}.find-options{display:flex;gap:12px;align-items:center;flex-wrap:wrap}.find-options label{display:flex;gap:6px;align-items:center;text-transform:none;letter-spacing:0}.find-options select{width:auto;background:var(--bg-editor);color:var(--text);border:1px solid var(--line);border-radius:6px;padding:6px 8px;font-size:12px}
.selection-menu{position:fixed;z-index:55;display:flex;align-items:center;gap:6px;background:var(--bg-panel);border:1px solid var(--line);border-radius:8px;padding:6px;box-shadow:0 12px 32px rgba(0,0,0,.35);max-width:min(92vw,520px)}.selection-menu[hidden]{display:none}.selection-menu select{width:260px;max-width:calc(92vw - 58px);background:var(--bg-editor);color:var(--text);border:1px solid var(--line);border-radius:6px;padding:7px 8px;font-size:12px}.selection-menu button{flex:0 0 auto}
*{scrollbar-color:#3a372f transparent;scrollbar-width:thin}::-webkit-scrollbar{width:10px;height:10px}::-webkit-scrollbar-track{background:transparent}::-webkit-scrollbar-corner{background:transparent}::-webkit-scrollbar-thumb{background:#3a372f;border-radius:6px}::-webkit-scrollbar-thumb:hover{background:#4a4638}#editor::-webkit-scrollbar-track,.sidebar::-webkit-scrollbar-track,.preview-box::-webkit-scrollbar-track{background:transparent}#editor::-webkit-scrollbar-thumb,.sidebar::-webkit-scrollbar-thumb,.preview-box::-webkit-scrollbar-thumb{background:#3a372f;border-radius:6px}
.recovery{display:none;align-items:center;gap:8px;background:#2b2117;border-bottom:1px solid #67451d;padding:8px 18px;font-size:12px}.recovery.show{display:flex}.recovery .spacer{flex:1}
.modal{display:none;position:fixed;inset:0;z-index:60;background:rgba(0,0,0,.55);align-items:center;justify-content:center;padding:18px}.modal.open{display:flex}.modal-card{background:var(--bg-panel);border:1px solid var(--line);border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.45);display:flex;flex-direction:column;max-width:1180px;width:96vw;height:88vh;min-height:0}.modal-head{display:flex;align-items:center;justify-content:space-between;gap:10px;padding:10px 12px;border-bottom:1px solid var(--line)}.modal-title{font-family:var(--mono);font-size:12px;font-weight:800;letter-spacing:.08em;text-transform:uppercase}.modal-actions{display:flex;gap:7px;align-items:center;flex-wrap:wrap}.integrity-panel{display:none;border-bottom:1px solid var(--line);padding:8px 12px;background:#1c1b17;font-family:var(--mono);font-size:11px;line-height:1.45;color:var(--muted)}.integrity-panel.show{display:block}.integrity-panel.warn{color:var(--amber)}.integrity-panel.ok{color:var(--green)}.integrity-panel .integrity-items{display:flex;flex-wrap:wrap;gap:6px;margin-top:5px}.integrity-panel .pill{border:1px solid var(--line);border-radius:999px;padding:3px 7px;background:var(--bg-raised);color:var(--muted)}.integrity-panel.warn .pill.issue{border-color:#6f5132;color:#ffd19a}.preview-grid{flex:1;min-height:0;display:grid;grid-template-columns:1fr 1fr;gap:10px;padding:10px}.modal.after-only .preview-grid{grid-template-columns:1fr}.modal.after-only .before-pane{display:none}.modal.before-only .preview-grid{grid-template-columns:1fr}.modal.before-only .after-pane{display:none}.pane{display:flex;flex-direction:column;min-height:0;border:1px solid var(--line);border-radius:8px;overflow:hidden;background:#181713}.pane-head{display:flex;align-items:center;justify-content:space-between;padding:7px 9px;border-bottom:1px solid var(--line);font-family:var(--mono);font-size:11px;color:var(--muted)}.preview-box{flex:1;min-height:0;overflow:auto;padding:12px 13px;white-space:pre-wrap;font-family:var(--mono);font-size:12.5px;line-height:1.6}.compare-grid{flex:1;min-height:0;display:grid;grid-template-columns:1fr 1fr;gap:10px;padding:10px}.compare-text{flex:1;min-height:0;width:100%;resize:none;border:0;outline:0;background:#181713;color:var(--text);padding:12px 13px;font-family:var(--mono);font-size:12.5px;line-height:1.6;white-space:pre-wrap}.compare-meta{font-family:var(--mono);font-size:10.5px;color:var(--faint);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.modal-foot{display:flex;align-items:center;justify-content:space-between;gap:10px;padding:10px 12px;border-top:1px solid var(--line)}
textarea.prompt{height:190px;resize:vertical;font-family:var(--mono);line-height:1.35}.lm-status{font-family:var(--mono);font-size:11px;color:var(--muted);min-height:16px}.prompt-actions{display:grid;grid-template-columns:1fr 1fr;gap:6px}.lm-grid{display:grid;grid-template-columns:1fr 1fr;gap:6px}.footer-note{font-size:10px;color:var(--faint);line-height:1.35;margin-top:8px}
</style>
</head>
<body>
<header>
  <div class="brand"><div class="mark">Manuscript <span>Workbench</span></div><div class="sub">local Python edition · Markdown-friendly structure · DOCX import/export · local AI streaming</div></div>
  <div class="actions">
    <span class="filename" id="filename">no file loaded</span>
    <label class="btn" for="fileInput">Load File</label><input type="file" id="fileInput" accept=".txt,.md,.markdown,.rtf,.docx,.doc,.html,.htm,text/plain,text/markdown,text/x-markdown,text/html,application/rtf,application/vnd.openxmlformats-officedocument.wordprocessingml.document">
    <button id="importUrlBtn">Import URL</button>
    <label class="btn" for="folderInput">Merge Folder</label><input type="file" id="folderInput" webkitdirectory directory multiple accept=".txt,.md,.markdown,.rtf,.docx,.doc,.html,.htm,text/plain,text/markdown,text/x-markdown,text/html,application/rtf,application/vnd.openxmlformats-officedocument.wordprocessingml.document">
    <button id="compareBtn">Compare</button>
    <button id="undoBtn" disabled>Undo</button>
    <button id="clearBtn">Clear</button>
    <button id="exportTxtBtn">Export TXT</button>
    <button id="exportMarkedTxtBtn">Export Markdown</button>
    <button id="exportDocxBtn" class="primary">Export DOCX</button>
  </div>
</header>
<div class="recovery" id="recovery"><strong>Previous autosave found.</strong><span>Restore your last local session?</span><span class="spacer"></span><button id="restoreBtn" class="mini">Restore</button><button id="discardRestoreBtn" class="mini">Discard</button></div>
<main>
  <div class="card editor-card">
    <div class="statbar"><span><span class="label">chars</span><b id="statChars">0</b></span><span><span class="label">words</span><b id="statWords">0</b></span><span><span class="label">lines</span><b id="statLines">0</b></span><span><span class="label">chapters</span><b id="statChapters">0</b></span><span><span class="label">subheads</span><b id="statSubheads">0</b></span></div>
    <div class="quickbar" id="quickbar" aria-label="Quick formatting">
      <select id="quickStyle" title="Text style">
        <option value="">Text style...</option>
        <option value="normal">Normal text</option>
        <option value="chapter">Title / Chapter</option>
        <option value="part">Part</option>
        <option value="subhead1">Heading 1</option>
        <option value="subhead2">Heading 2</option>
        <option value="subhead3">Heading 3</option>
        <option value="quote">Quote block</option>
        <option value="verse">Verse</option>
        <option value="written-note">Written note</option>
      </select>
      <select id="quickIcon" title="Insert icon">
        <option value="">Insert icon...</option>
        <option value="📌">📌 Important</option>
        <option value="💡">💡 Tip</option>
        <option value="⚠️">⚠️ Warning</option>
        <option value="ℹ️">ℹ️ Info</option>
        <option value="✅">✅ Done</option>
        <option value="🔧">🔧 Setup</option>
        <option value="❗">❗ Attention</option>
        <option value="🔒">🔒 Security</option>
        <option value="📝">📝 Note</option>
      </select>
      <button data-quick-inline="bold" class="b" title="Bold">B</button>
      <button data-quick-inline="italic" class="i" title="Italic">I</button>
      <button data-quick-inline="u" class="u" title="Underline">U</button>
      <button data-quick-inline="sc" class="sc" title="Small Caps">SC</button>
      <button data-quick-inline="mono" class="mono" title="Monospace">Mono</button>
      <button data-quick-inline="sans" title="Sans Serif">Sans</button>
      <span class="sep"></span>
      <button id="openFindBtn" title="Find and replace">Find</button>
      <button id="togglePreview" title="Toggle rendered Markdown preview">Preview</button>
      <button id="quickAiRewrite" title="AI Rewrite Selection">AI</button>
      <button id="quickAiClean" title="AI Clean Selection">Clean</button>
      <button id="quickClearMark" title="Clear heading/structure marker">Clear</button>
    </div>
    <div class="toolbar collapsed" id="vellumToolbar">
      <div class="toolbar-head"><div class="toolbar-title">Markdown tools</div><button id="toggleVellum" class="mini">Expand</button></div>
      <div class="toolbar-body">
        <span class="group-label">Structure</span>
        <button data-marker="chapter" title="Marks this line as a chapter heading for Markdown/DOCX export.">Chapter</button><button data-marker="part" title="Marks a larger section that contains multiple chapters, such as Part One.">Part</button><button data-marker="subhead1" title="Marks a section title inside a chapter.">Subhead 1</button><button data-marker="subhead2" title="Marks a lower-level section title inside a chapter.">Subhead 2</button><button data-marker="subhead3" title="Marks a third-level section title inside a chapter.">Subhead 3</button><button data-marker="subhead4" title="Marks a fourth-level section title inside a chapter. Use sparingly.">Subhead 4</button><button data-marker="subhead5" title="Marks a fifth-level section title inside a chapter. Use rarely.">Subhead 5</button>
        <button data-marker="subtitle" title="Adds a chapter subtitle. It appears with the chapter heading but is not normally a TOC item.">Subtitle</button><button data-marker="element-author" title="Adds an author credit for this chapter/section, useful for anthologies or guest sections.">Element Author</button><button data-marker="hidden-heading" title="Keeps the element title for structure/TOC while hiding the visible heading in the book.">Hidden Heading</button>
        <span class="group-label">Special</span>
        <button data-marker="quote" title="Formats this paragraph as an inset block quotation.">Quote</button><button data-marker="verse" title="Formats lines as poetry, lyrics, or verse where line breaks matter.">Verse</button><button data-marker="attribution" title="Formats a quote attribution, usually aligned after a block quote.">Attribution</button><button data-marker="centered" title="Centers this paragraph as an alignment block.">Centered</button><button data-marker="flush-left" title="Forces this paragraph flush left, overriding normal paragraph styling.">Flush Left</button><button data-marker="written-note" title="Formats text as a note written by a character.">Written Note</button><button data-marker="text-conversation" title="Formats text messages or chat-style dialogue between characters.">Text Conversation</button><button data-marker="caption" title="Formats this line as an image caption.">Caption</button><button data-marker="inline-image" title="Marks a placeholder line for an inline image to handle later in Word or another editor.">Inline Image</button><button data-insert="ornament" title="Inserts an ornamental break, a styled divider often used between scenes.">Ornament</button><button data-insert="scene" title="Inserts a plain scene break for a change of time, place, or narrator.">Scene Break</button><button id="clearVellumBtn" title="Removes the structure marker from the selected/current line without deleting the text.">Clear Mark</button>
        <span class="group-label">Inline</span>
        <button data-inline="italic" title="Wraps the selected text in italic markup for export.">Italic</button><button data-inline="bold" title="Wraps the selected text in bold markup for export.">Bold</button><button data-inline="u" title="Marks selected text as underlined for Word export.">Underline</button><button data-inline="sc" title="Marks selected text as small caps for DOCX export.">Small Caps</button><button data-inline="mono" title="Marks selected text as monospace, useful for code or machine text.">Monospace</button><button data-inline="sans" title="Marks selected text as sans serif for a distinct inline style.">Sans Serif</button>
      </div>
    </div>
    <div class="editor-wrap">
      <div class="editor-area" id="editorArea"><div class="gutter" id="gutter"></div><div class="find-highlights" id="findHighlights" aria-hidden="true"></div><textarea id="editor" spellcheck="false" wrap="off" placeholder="Load a TXT, MD, RTF, DOCX, or HTML file, or paste text here.&#10;&#10;Markdown structure:&#10;# Chapter One&#10;## A Subhead&#10;&gt; A quoted paragraph&#10;***&#10;&#10;Inline markup:&#10;*italic*  **bold**  [[u:underline]]&#10;[[sc:small caps]]  [[mono:code]]  [[sans:sans serif]]"></textarea><div class="markdown-preview" id="markdownPreview" aria-label="Markdown preview"></div></div>
      <div class="editor-foot">
        <div class="scope-toggle" id="scopeToggle">
          <label><input type="radio" name="scope" value="all" checked><span>Whole text</span></label>
          <label><input type="radio" name="scope" value="selection"><span>Selection only</span></label>
        </div>
        <label class="wrap-control"><input type="checkbox" id="wrapToggle"><span>Word wrap</span></label>
        <span class="sel-info" id="selInfo">no selection</span>
      </div>
    </div>
  </div>
  <div class="sidebar">
    <div class="side-tabs" role="tablist" aria-label="Workbench tools">
      <button class="active" data-side-tab="check" type="button">Check</button>
      <button data-side-tab="outline" type="button">Outline</button>
      <button data-side-tab="clean" type="button">Clean</button>
      <button data-side-tab="ai" type="button">AI</button>
      <button data-side-tab="vellum" type="button" title="Markdown">MD</button>
      <button data-side-tab="book" type="button">Book</button>
    </div>

    <div class="side-panel active" data-side-panel="check">
      <div class="card">
        <div class="panel-title">Diagnostics</div>
        <div id="diagList"></div>
      </div>
      <div class="card">
        <div class="panel-title"><span>Preflight</span><button id="preflightBtn" class="mini amber">Run</button></div>
        <div class="panel-body"><div id="preflightSummary" class="hint">Run Preflight before exporting to catch unmarked chapters, page numbers, TOC text, tabs and repeated headers.</div><div class="issue-list" id="issueList"></div></div>
      </div>
    </div>

    <div class="side-panel" data-side-panel="outline">
      <div class="card">
        <div class="panel-title">Chapter Navigator</div>
        <div class="empty-hint">Clickable outline from Markdown headings. Chapter actions use the text between this chapter and the next chapter.</div>
        <div class="outline-list" id="outlineList"></div>
      </div>
    </div>

    <div class="side-panel" data-side-panel="ai">
      <div class="card" id="lmPanel">
        <div class="panel-title">AI Provider</div>
        <div class="lm-panel">
          <details class="ai-section">
            <summary>Provider setup</summary>
            <div class="ai-section-body">
              <div class="lm-field"><label for="lmProvider">Provider</label><select id="lmProvider"><option value="lmstudio">LM Studio (easy GUI)</option><option value="omlx">oMLX (Mac MLX)</option><option value="ollama">Ollama (cross-platform)</option><option value="custom">Custom OpenAI-compatible</option></select></div>
              <div class="lm-field"><label for="lmBase">Endpoint</label><input id="lmBase" type="text" value="http://localhost:1234/v1" spellcheck="false"></div>
              <div class="lm-field"><label for="lmApiKey">API Key</label><input id="lmApiKey" type="password" placeholder="optional; required by some providers such as oMLX" spellcheck="false" autocomplete="off"></div>
              <div class="lm-field"><label for="lmModel">Model</label><select id="lmModel"><option value="">Detect first available model</option></select></div>
              <div class="lm-grid"><button id="detectModelBtn">Detect Model</button><button id="lmTestBtn">Test</button></div>
            </div>
          </details>
          <details class="ai-section" open>
            <summary>Rewrite & clean</summary>
            <div class="ai-section-body">
              <div class="lm-field"><label for="promptSelect">Prompt preset</label><select id="promptSelect"></select></div>
              <div class="lm-field"><label for="lmStrength">Rewrite strength</label><select id="lmStrength"><option value="default">Prompt default</option><option value="light">Light</option><option value="medium" selected>Medium</option><option value="strong">Strong</option></select></div>
              <div class="lm-grid"><button id="rewriteSelBtn" class="primary">Rewrite Selection</button><button id="rewriteParaBtn">Current Paragraph</button><button id="rewriteWholeBtn">Rewrite Whole</button><button id="rewriteChunkBtn" title="Uses the current selection if there is one; otherwise rewrites the whole text in chunks.">Rewrite Chunked</button><button id="cleanSelBtn">AI Clean Selection</button><button id="cleanParaBtn">AI Clean Paragraph</button><button id="cleanWholeBtn">AI Clean Whole</button><button id="cleanChunkBtn" title="Uses the current selection if there is one; otherwise cleans the whole text in chunks.">AI Clean Chunked</button></div>
            </div>
          </details>
          <details class="ai-section" open>
            <summary>Guided tools</summary>
            <div class="ai-section-body">
              <div class="lm-field"><label for="guidedContext">Goal / answers / context</label><textarea class="compact" id="guidedContext" placeholder="Optional. Example: make this warmer, less formal, for newsletter readers; or paste answers to AI questions here."></textarea></div>
              <div class="lm-grid"><button id="askImproveBtn">Improve Questions</button><button id="feedbackOnlyBtn">Feedback Only</button><button id="rewriteGuidedBtn" class="primary">Rewrite with Guidance</button><button id="translateSelBtn">Translate Selection</button></div>
              <div class="lm-field"><label for="translateTarget">Translate to</label><input id="translateTarget" type="text" value="English" spellcheck="false"></div>
            </div>
          </details>
          <details class="ai-section">
            <summary>Prompt editor</summary>
            <div class="ai-section-body">
              <div class="lm-field"><label for="promptText">System prompt</label><textarea class="prompt" id="promptText"></textarea></div>
              <div class="prompt-actions"><button id="savePromptBtn">Save Prompt</button><button id="savePromptAsBtn">Save As New</button><button id="deletePromptBtn">Delete</button><button id="resetPromptsBtn">Reset Defaults</button></div>
            </div>
          </details>
          <details class="ai-section">
            <summary>Advanced request options</summary>
            <div class="ai-section-body">
              <div class="advanced-grid">
                <div class="lm-field"><label for="advTemperature">Temperature</label><input id="advTemperature" type="text" placeholder="provider default"></div>
                <div class="lm-field"><label for="advTopP">Top P</label><input id="advTopP" type="text" placeholder="provider default"></div>
                <div class="lm-field"><label for="advTopK">Top K</label><input id="advTopK" type="text" placeholder="provider default"></div>
                <div class="lm-field"><label for="advRepeatPenalty">Repeat penalty</label><input id="advRepeatPenalty" type="text" placeholder="provider default"></div>
                <div class="lm-field"><label for="advMaxTokens">Max tokens</label><input id="advMaxTokens" type="text" placeholder="provider default"></div>
                <div class="lm-field"><label for="chunkSize">Chunk size</label><input id="chunkSize" type="text" value="6000"></div>
              </div>
              <div class="hint">Leave fields empty to use the provider or model defaults.</div>
            </div>
          </details>
          <div class="lm-grid"><button id="lmStopBtn" class="danger wide" disabled>Stop generation</button><button id="showPreviewBtn" class="wide" disabled>Show Preview</button></div>
          <div class="lm-status" id="lmStatus">Local AI idle.</div>
          <div class="lm-note" id="lmNote">This app sends only model + messages + stream to the selected OpenAI-compatible local endpoint. Runtime settings stay in your provider app.</div>
        </div>
      </div>
    </div>

    <div class="side-panel" data-side-panel="clean">
      <div class="card">
        <div class="panel-title">Operations</div>
        <div class="empty-hint">Safe Cleanup is one-pass and conservative. Structure and typography actions show a preview before applying.</div>
        <div class="fix-row"><button id="safeCleanupBtn" class="primary">Safe Cleanup: encoding, spaces, hidden chars</button></div>
        <div class="fix-row"><button id="showMarkersBtn">Strip Markers Preview</button></div>
        <div class="fix-row"><button data-fix="htmlText">Extract text from HTML</button></div>

        <div class="fix-group">
          <div class="fix-group-title">Encoding &amp; hidden characters</div>
          <div class="fix-row"><button data-fix="mojibake">Repair mojibake</button></div>
          <div class="fix-row"><button data-fix="invisible">Remove zero-width / invisible chars</button></div>
          <div class="fix-row"><button data-fix="specialSpaces">Normalize special spaces</button></div>
          <div class="fix-row"><button data-fix="controlChars">Strip control / page-break chars</button></div>
          <div class="fix-row"><button data-fix="unicodeNFC">Normalize Unicode composition</button></div>
          <div class="fix-row"><button data-fix="ligatures">Replace typographic ligatures</button></div>
        </div>
        <div class="fix-group">
          <div class="fix-group-title">Structure</div>
          <div class="fix-row"><button data-fix="join">Join wrapped lines conservatively</button></div>
          <div class="fix-note"><label><input type="checkbox" id="dehyphenate" checked> re-join words split by a hyphen at line end</label></div>
          <div class="fix-row"><button data-fix="blankRuns">Normalize blank lines</button></div>
          <div class="fix-row"><button data-fix="trailingSpaces">Trim trailing spaces only</button></div>
          <div class="fix-row"><button data-fix="indentation">Remove indentation</button></div>
          <div class="fix-row"><button data-fix="pageNumbers">Remove standalone page-number lines</button></div>
          <div class="fix-row"><button data-fix="repeatLines">Remove repeated header/footer candidates</button></div>
          <div class="fix-note caveat"><span>Previewed first: page numbers, headers, poetry, scene breaks and chapter titles can look similar.</span></div>
        </div>
        <div class="fix-group">
          <div class="fix-group-title">Spaces &amp; tabs</div>
          <div class="fix-row"><button data-fix="tabsToSpace">Tabs to spaces</button></div>
          <div class="fix-row"><button data-fix="collapseSpaces">Collapse multiple spaces</button></div>
          <div class="fix-row"><button data-fix="spaceBeforePunct">Remove space before punctuation</button></div>
          <div class="fix-row"><button data-fix="spaceAfterPunct">Add missing space after punctuation</button></div>
        </div>
        <div class="fix-group">
          <div class="fix-group-title">Typography</div>
          <div class="fix-row"><button data-fix="quotesCurly">Quotes to curly</button></div>
          <div class="fix-row"><button data-fix="quotesStraight">Quotes to straight</button></div>
          <div class="fix-row"><button data-fix="dashEm">-- to em dash</button></div>
          <div class="fix-row"><button data-fix="ellipsis">... to ellipsis</button></div>
        </div>
        <div class="fix-group">
          <div class="fix-group-title">File</div>
          <div class="fix-row"><button data-fix="trimDoc">Trim document edges</button></div>
        </div>
      </div>
    </div>

    <div class="side-panel" data-side-panel="vellum">
      <div class="card">
        <div class="panel-title">Markdown Structure</div>
        <div class="vellum-help">Marked lines are exported as real Word styles. Use the toolbar above the editor to mark structure and inline formatting.</div>
        <div class="vellum-summary" id="vellumSummary"></div>
        <div id="vellumList"></div>
        <div class="marker-legend">Markdown structure includes # chapter, ## subhead, &gt; quote, ***, and fenced blocks like ::: verse or ::: written-note.</div>
      </div>
    </div>

    <div class="side-panel" data-side-panel="book">
      <div class="card">
        <div class="panel-title">Metadata</div>
        <div class="panel-body"><div class="meta-grid">
          <label>Book Title<input id="metaTitle" type="text"></label><label>Subtitle<input id="metaSubtitle" type="text"></label><label>Author<input id="metaAuthor" type="text"></label><label>Language<input id="metaLanguage" type="text" placeholder="en / nl"></label><label>Series<input id="metaSeries" type="text"></label><label>Book #<input id="metaBookNumber" type="text"></label><label>Publisher<input id="metaPublisher" type="text"></label><label>Filename base<input id="metaFilename" type="text" placeholder="manuscript"></label>
        </div><p class="hint">Title/author go into DOCX core properties; the other fields are stored as keywords/comments where possible.</p></div>
      </div>
    </div>
  </div>
</main>
<div class="selection-menu" id="selectionMenu" hidden>
  <select id="selectionAction" aria-label="Selection action">
    <option value="">Selection actions...</option>
    <optgroup label="AI">
      <option value="ai-rewrite">AI Rewrite Selection</option>
      <option value="ai-clean">AI Clean Selection</option>
    </optgroup>
    <optgroup label="Inline">
      <option value="inline-bold">Bold</option>
      <option value="inline-italic">Italic</option>
      <option value="inline-u">Underline</option>
      <option value="inline-sc">Small Caps</option>
      <option value="inline-mono">Monospace</option>
      <option value="inline-sans">Sans Serif</option>
    </optgroup>
    <optgroup label="Markdown structure">
      <option value="mark-chapter">Chapter</option>
      <option value="mark-part">Part</option>
      <option value="mark-subhead1">Subhead 1</option>
      <option value="mark-subhead2">Subhead 2</option>
      <option value="mark-quote">Quote</option>
      <option value="mark-verse">Verse</option>
      <option value="mark-written-note">Written Note</option>
      <option value="clear-vellum">Clear Structure Mark</option>
    </optgroup>
    <optgroup label="Clean">
      <option value="safe-clean">Safe Cleanup Selection</option>
    </optgroup>
  </select>
  <button id="selectionRunBtn" class="mini primary" type="button">Run</button>
</div>
<div class="toast" id="toast"></div>
<div class="modal" id="previewModal">
  <div class="modal-card" id="modalCard">
    <div class="modal-head"><div><div class="modal-title" id="previewTitle">Preview</div><div class="small" id="previewSub">before / after</div></div><div class="modal-actions"><button data-view="split" class="mini viewBtn">Split</button><button data-view="after-only" class="mini viewBtn">After only</button><button data-view="before-only" class="mini viewBtn">Before only</button><label class="small"><input type="checkbox" id="previewLinked"> linked scroll</label><label class="small"><input type="checkbox" id="autoScroll" checked> auto-scroll</label><button id="bottomBtn" class="mini">Bottom</button><button id="closePreviewBtn" class="mini">Close</button></div></div>
    <div class="integrity-panel" id="integrityPanel"></div>
    <div class="preview-grid"><div class="pane before-pane"><div class="pane-head"><span>Before</span></div><div class="preview-box" id="beforeBox"></div></div><div class="pane after-pane"><div class="pane-head"><span>After</span><span id="streamState" class="small"></span></div><div class="preview-box" id="afterBox"></div></div></div>
    <div class="modal-foot"><span class="lm-status" id="previewStatus">Ready.</span><div class="modal-actions"><button id="comparePreviewBtn">Open in Compare</button><button id="stopBtn" class="danger" disabled>Stop generation</button><button id="applyPreviewBtn" class="primary" disabled>Apply change</button></div></div>
  </div>
</div>
<div class="modal" id="compareModal">
  <div class="modal-card">
    <div class="modal-head"><div><div class="modal-title">Compare / Parallel View</div><div class="small" id="compareSub">Original and revised text side by side</div></div><div class="modal-actions"><label class="small"><input type="checkbox" id="compareLinked" checked> linked scroll</label><button id="closeCompareBtn" class="mini">Close</button></div></div>
    <div class="compare-grid">
      <div class="pane"><div class="pane-head"><span>Original</span><span id="compareOriginalMeta" class="compare-meta">empty</span></div><textarea id="compareOriginal" class="compare-text" spellcheck="false"></textarea></div>
      <div class="pane"><div class="pane-head"><span>Revised / Translation</span><span id="compareRevisedMeta" class="compare-meta">empty</span></div><textarea id="compareRevised" class="compare-text" spellcheck="false"></textarea></div>
    </div>
    <div class="modal-foot"><span class="lm-status" id="compareStatus">Ready.</span><div class="modal-actions"><label class="btn mini" for="compareOriginalInput">Load Original</label><input type="file" id="compareOriginalInput" accept=".txt,.md,.markdown,.rtf,.docx,.doc,.html,.htm,text/plain,text/markdown,text/x-markdown,text/html,application/rtf,application/vnd.openxmlformats-officedocument.wordprocessingml.document"><label class="btn mini" for="compareRevisedInput">Load Revised</label><input type="file" id="compareRevisedInput" accept=".txt,.md,.markdown,.rtf,.docx,.doc,.html,.htm,text/plain,text/markdown,text/x-markdown,text/html,application/rtf,application/vnd.openxmlformats-officedocument.wordprocessingml.document"><button id="compareUseEditorOriginal">Editor -> Original</button><button id="compareUseEditorRevised">Editor -> Revised</button><button id="compareSwapBtn">Swap</button><button id="compareExportRevisedBtn">Export Revised</button><button id="compareCopyRevisedBtn" class="primary">Revised -> Editor</button></div></div>
  </div>
</div>
<div class="modal" id="mergeModal">
  <div class="modal-card" style="max-width:760px;height:78vh">
    <div class="modal-head"><div><div class="modal-title">Merge Folder</div><div class="small" id="mergeSub">Arrange files before merging</div></div><div class="modal-actions"><button id="mergeCancelBtn" class="mini">Cancel</button></div></div>
    <div class="merge-list" id="mergeList"></div>
    <div class="modal-foot"><span class="lm-status" id="mergeStatus">Ready.</span><div class="modal-actions"><button id="mergeBtn" class="primary">Merge into editor</button></div></div>
  </div>
</div>
<div class="modal" id="urlModal">
  <div class="modal-card" style="max-width:720px;height:auto;min-height:0">
    <div class="modal-head"><div><div class="modal-title">Import URL</div><div class="small">Fetch a web page and extract readable text</div></div><div class="modal-actions"><button id="urlCancelBtn" class="mini">Cancel</button></div></div>
    <div class="find-body">
      <div class="find-field"><label for="urlInput">URL</label><input id="urlInput" type="text" spellcheck="false" placeholder="https://example.com/article"></div>
      <div class="hint">Only http/https URLs are allowed. Local and private addresses are blocked.</div>
    </div>
    <div class="modal-foot"><span class="lm-status" id="urlStatus">Ready.</span><div class="modal-actions"><button id="urlFetchBtn" class="primary">Fetch URL</button></div></div>
  </div>
</div>
<div class="modal" id="htmlImportModal">
  <div class="modal-card">
    <div class="modal-head"><div><div class="modal-title" id="htmlImportTitle">Choose HTML Import</div><div class="small" id="htmlImportSub">Compare extraction methods before loading</div></div><div class="modal-actions"><button id="htmlImportCancelBtn" class="mini">Cancel</button></div></div>
    <div class="preview-grid"><div class="pane"><div class="pane-head"><span>Simple extraction</span><span id="htmlSimpleMeta" class="small"></span></div><div class="preview-box" id="htmlSimpleBox"></div></div><div class="pane"><div class="pane-head"><span>Main text extraction</span><span id="htmlMainMeta" class="small"></span></div><div class="preview-box" id="htmlMainBox"></div></div></div>
    <div class="modal-foot"><span class="lm-status" id="htmlImportStatus">Choose an import method.</span><div class="modal-actions"><button id="htmlUseSimpleBtn">Use Simple</button><button id="htmlUseMainBtn" class="primary">Use Main Text</button></div></div>
  </div>
</div>
<div class="modal" id="findModal">
  <div class="modal-card" style="max-width:760px;height:auto;min-height:0">
    <div class="modal-head"><div><div class="modal-title">Find / Replace</div><div class="small">Search whole manuscript or current selection</div></div><div class="modal-actions"><button id="closeFindBtn" class="mini">Close</button></div></div>
    <div class="find-body">
      <div class="find-grid">
        <div class="find-field"><label for="findText">Find</label><input id="findText" type="text" spellcheck="false"></div>
        <div class="find-field"><label for="replaceText">Replace with</label><input id="replaceText" type="text" spellcheck="false"></div>
      </div>
      <div class="find-options">
        <label><input type="checkbox" id="findRegex"> Regex</label>
        <label><input type="checkbox" id="findCase"> Case sensitive</label>
        <label>Scope <select id="findScope"><option value="all">Whole text</option><option value="selection">Selection only</option></select></label>
      </div>
    </div>
    <div class="modal-foot"><span class="lm-status" id="findStatus">Ready.</span><div class="modal-actions"><button id="findPrevBtn">Previous</button><button id="findNextBtn" class="primary">Next</button><button id="replaceOneBtn">Replace</button><button id="replaceAllBtn" class="danger">Replace All</button></div></div>
  </div>
</div>
<script>
const $ = sel => document.querySelector(sel);
const editor = $('#editor');
const gutter = $('#gutter');
const filenameEl = $('#filename');
const toastEl = $('#toast');
const VELLUM_CHAPTER_CANDIDATE_RE = /^\s*(?:chapter|hoofdstuk)\s+(?:\d+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|een|twee|drie|vier|vijf|zes|zeven|acht|negen|tien|elf|twaalf)\b|^\s*(?:prologue|epilogue|acknowledg(?:e)?ments|about the author|copyright|dedication|preface|introduction|afterword|part\s+\w+|deel\s+\w+)\s*$/i;
let currentFilename = 'manuscript.txt';
let autosaveTimer = null;
let autosaveWarningShown = false;
let recoveryData = null;
let recoveryPending = false;
let lmAbort = null;
let previewState = {before:'', after:'', start:0, end:0, title:'Preview', finished:false, running:false, partial:false, expectedMarkers:null, applyMode:'replace'};
let lastDiag = null;
let savedSelection = null;
let scrollSyncing = false;
let previewScrollSyncing = false;
let markdownPreviewTimer = null;
let markdownPreviewRequest = 0;
let markdownServerAvailable = true;
let mergeFiles = [];
let htmlImportState = null;
let compareScrollSyncing = false;
let findMatches = [];
let findIndex = -1;
let findScopeRange = null;
const collapsedOutlineHeadings = new Set();
const jumpPointers = {};
const undoStack = [];
const DESTRUCTIVE_FIXES = new Set(['htmlText','join','blankRuns','pageNumbers','repeatLines','indentation','quotesCurly','quotesStraight','dashEm','ellipsis']);
const SAFE_FIX_SEQUENCE = ['mojibake','invisible','specialSpaces','controlChars','unicodeNFC','ligatures','tabsToSpace','trailingSpaces','collapseSpaces','spaceBeforePunct','spaceAfterPunct','trimDoc'];
const FIX_LABELS = {
  mojibake:'Repair mojibake', invisible:'Remove invisible characters', specialSpaces:'Normalize special spaces', controlChars:'Strip control characters',
  unicodeNFC:'Normalize Unicode composition', ligatures:'Replace typographic ligatures', join:'Join wrapped lines conservatively', blankRuns:'Normalize blank lines',
  trailingSpaces:'Trim trailing spaces', indentation:'Remove indentation', pageNumbers:'Remove standalone page-number lines', repeatLines:'Remove repeated header/footer candidates',
  tabsToSpace:'Tabs to spaces', collapseSpaces:'Collapse multiple spaces', spaceBeforePunct:'Remove space before punctuation', spaceAfterPunct:'Add missing space after punctuation',
  quotesCurly:'Convert quotes to curly', quotesStraight:'Convert quotes to straight', dashEm:'Convert double hyphen to em dash', ellipsis:'Convert three dots to ellipsis', trimDoc:'Trim document edges',
  htmlText:'Extract text from HTML'
};

const DEFAULT_PROMPTS = [
 {name:'Text Rewriter — balanced', text:`## Role\nYou are Text Rewriter, a text-cleaning and rewriting utility. Rewrite text supplied by the user into a cleaner, more readable version.\n\n## Rules\nPreserve meaning, events, outcomes, scene order, structure, tone, point of view, tense, character dynamics, and level of detail. Rewrite every part in order; do not summarize, skip, add scenes, add explanations, or moralize. Preserve chapter titles, headings, lists, scene breaks, and dialogue paragraphing. Use clear, natural, modern prose.\n\nRewrite strength: medium.\nOutput only the rewritten text.`},
 {name:'AI Cleanup — conservative', text:`You are a conservative manuscript cleanup utility. Clean the text without rewriting style or changing meaning. Fix typos, spacing, broken line breaks, duplicated words, inconsistent quotes, mojibake artifacts, chatbot clutter, and obvious OCR/PDF errors. Preserve all events, dialogue meaning, headings, scene breaks, sequence, POV, tense, tone, and level of detail. Do not summarize or add content. Output only the cleaned text.`},
 {name:'Light cleanup only', text:`Clean grammar, punctuation, spacing, repeated words, broken paragraphing, and obvious typos. Stay as close as possible to the original wording. Preserve all meaning, dialogue, headings, structure, POV, tense, and scene order. Output only the cleaned text.`},
 {name:'Humanize Rewrite — natural', text:`## Role
You are a natural manuscript rewriter. Rewrite the supplied text so it sounds more human, fluent, and alive while preserving the author's intent.

## Core rules
Rewrite the text. Do not merely proofread it.
Preserve meaning, events, facts, names, numbers, timeline, scene order, POV, tense, tone, character dynamics, and level of detail.
Do not summarize, skip, add new story content, moralize, explain, or comment.
Preserve Markdown structure exactly: headings (#, ##, ###), blockquotes (>), scene breaks (***), lists, tables, and fenced block lines (:::) must remain in place.
Preserve paragraph order and dialogue order. Dialogue may be smoothed, but the speaker's intent, voice, subtext, and information must stay the same.

## Style
Detect the input language. English and Dutch are fully supported. For mixed text, handle each language in its own natural register.
Remove stiff AI-like prose patterns where they weaken the text: overly formal openers, repetitive transitions, passive constructions, redundant phrases, and overlong sentences.
Use active voice where it reads better.
Split long sentences when it improves rhythm, but keep deliberate long sentences if they fit the voice.
Vary sentence openings and rhythm.
Keep contractions or natural spoken forms when they fit the tone.
Remove em dashes by default because they often make AI-written prose feel obvious. Replace them with periods, commas, parentheses, colons, or a clean sentence split. Keep an em dash only when it is clearly intentional in existing dialogue or part of a strong authorial voice.
Avoid making the prose generic, overly casual, or over-polished.

## Dutch guidance
Maak Nederlands natuurlijker en minder stijf. Vervang formele formuleringen zoals "teneinde", "indien", "met betrekking tot", "door middel van", "er dient te worden" en "het is belangrijk op te merken dat" wanneer dat natuurlijker leest.
Gebruik spreektaal alleen subtiel als de tekst daarom vraagt. Niet overdrijven.

## Output
Output only the rewritten text.`},
 {name:'Strong prose polish', text:`Polish the prose strongly while preserving every event, action, outcome, dialogue meaning, POV, tense, scene order, character dynamic, and level of detail. Do not add or remove story content. Make the prose smooth, clear, and modern without purple prose. Output only the rewritten text.`},
 {name:'Nederlands — clean herschrijven', text:`Je bent een Nederlandse tekstredacteur. Herschrijf de aangeleverde tekst naar soepel, helder modern Nederlands. Behoud betekenis, gebeurtenissen, volgorde, perspectief, tijd, toon, dialoogbetekenis, hoofdstuktitels, scene breaks en detailniveau. Niet samenvatten, niets toevoegen, niets weglaten. Output alleen de herschreven tekst.`}
];
const AI_PROVIDERS={
  lmstudio:{label:'LM Studio',base:'http://localhost:1234/v1',note:'LM Studio provider. Load a model and start the Local Server in LM Studio, then Detect Model here.'},
  omlx:{label:'oMLX',base:'http://localhost:8000/v1',note:'oMLX provider for MLX-format models on Apple Silicon. Start oMLX first, then Detect Model here.'},
  ollama:{label:'Ollama',base:'http://localhost:11434/v1',note:'Ollama provider. Pull/run models in Ollama, then Detect Model here.'},
  custom:{label:'Custom OpenAI-compatible',base:'',note:'Custom OpenAI-compatible endpoint. Enter the base URL ending in /v1 if needed.'}
};

function toast(msg){ toastEl.textContent=msg; toastEl.classList.add('show'); clearTimeout(toastEl._t); toastEl._t=setTimeout(()=>toastEl.classList.remove('show'),2600); }
function api(path, data){ return fetch(path,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(data||{})}); }
function getScope(){ return document.querySelector('input[name=scope]:checked')?.value || 'all'; }
function selectedTextLength(){ return Math.max(0, editor.selectionEnd - editor.selectionStart); }
function currentAiProvider(){ return AI_PROVIDERS[$('#lmProvider')?.value] || AI_PROVIDERS.custom; }
function resetModelSelect(placeholder='Detect first available model'){ const sel=$('#lmModel'); if(!sel) return; sel.innerHTML=`<option value="">${escapeHtml(placeholder)}</option>`; }
function updateProviderUi(){ const provider=currentAiProvider(); if($('#lmProvider').value!=='custom' && $('#lmBase').value.trim()!==provider.base){ $('#lmBase').value=provider.base; resetModelSelect(); } $('#lmNote').textContent=provider.note; $('#lmStatus').textContent=`${provider.label} idle.`; }
function getMetadata(){ return {title:$('#metaTitle').value.trim(), subtitle:$('#metaSubtitle').value.trim(), author:$('#metaAuthor').value.trim(), language:$('#metaLanguage').value.trim(), series:$('#metaSeries').value.trim(), bookNumber:$('#metaBookNumber').value.trim(), publisher:$('#metaPublisher').value.trim(), filename:$('#metaFilename').value.trim()}; }
function setMetadata(m){ m=m||{}; $('#metaTitle').value=m.title||''; $('#metaSubtitle').value=m.subtitle||''; $('#metaAuthor').value=m.author||''; $('#metaLanguage').value=m.language||''; $('#metaSeries').value=m.series||''; $('#metaBookNumber').value=m.bookNumber||''; $('#metaPublisher').value=m.publisher||''; $('#metaFilename').value=m.filename||''; }
const MD_PREFIX_MARKERS={chapter:'# ',part:'# ',subhead1:'## ',subhead2:'### ',subhead3:'#### ',subhead4:'##### ',subhead5:'###### ',quote:'> '};
const MD_BLOCK_MARKERS=new Set(['verse','attribution','centered','flush-left','written-note','text-conversation','caption','inline-image','subtitle','element-author','hidden-heading']);
const MD_STYLE_INFO={chapter:['H1','chapter'],part:['H1','chapter'],subhead1:['H2','sub'],subhead2:['H3','sub'],subhead3:['H4','sub'],subhead4:['H5','sub'],subhead5:['H6','sub'],quote:['QUOTE','supp'],verse:['VERSE','supp'],attribution:['ATTR','supp'],centered:['CENTER','supp'],'flush-left':['FLUSH','supp'],'written-note':['NOTE','supp'],'text-conversation':['TEXT','supp'],caption:['CAPTION','supp'],'inline-image':['IMAGE','supp'],subtitle:['SUBTITLE','supp'],'element-author':['AUTHOR','supp'],'hidden-heading':['HIDDEN','supp'],scene:['SCENE','supp']};
function canonicalStructureMarker(marker){ return ({note:'written-note',text:'text-conversation',image:'inline-image'}[marker]||marker||'').toLowerCase(); }
function isPartHeadingText(text){ return /^\s*(?:part|deel)\b/i.test(text||''); }
function markdownLineInfo(line){ let m=line.match(/^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$/); if(m){ const content=m[2].trim(); const level=m[1].length; const map={1:isPartHeadingText(content)?'part':'chapter',2:'subhead1',3:'subhead2',4:'subhead3',5:'subhead4',6:'subhead5'}; return {marker:map[level],content}; } m=line.match(/^\s{0,3}>\s?(.*)$/); if(m) return {marker:'quote',content:m[1].trim()}; if(/^\s*(?:\*\s*){3,}$/.test(line)||/^\s*[-–—]{3,}\s*$/.test(line)||/^\s*~{3,}\s*$/.test(line)) return {marker:'scene',content:''}; return null; }
function stripStructureLine(line){ const info=markdownLineInfo(line); return (info ? info.content : line).trimEnd(); }
function stripInlineMarkup(text){ return String(text||'').replace(/\[\[(?:u|sc|sans|mono):(.+?)\]\]/g,'$1').replace(/!\[([^\]]*)\]\([^)]+\)/g,'$1').replace(/\[([^\]]+)\]\([^)]+\)/g,'$1').replace(/`([^`\n]+?)`/g,'$1').replace(/\*\*\*(.+?)\*\*\*/g,'$1').replace(/\*\*(.+?)\*\*/g,'$1').replace(/(^|[^\*])\*(?!\s)(.+?)(?<!\s)\*(?!\*)/g,'$1$2').replace(/__(.+?)__/g,'$1').replace(/(^|[^_])_(?!\s)(.+?)(?<!\s)_(?!_)/g,'$1$2').replace(/~~(.+?)~~/g,'$1'); }
function stripMarkers(text){ return text.split('\n').filter(l=>!l.match(/^\s*:::\s*[a-z0-9_-]*\s*$/i)).map(l=>{ const info=markdownLineInfo(l); const line=info?.marker==='scene'?'* * *':stripStructureLine(l); return stripInlineMarkup(line); }).join('\n'); }
function safePreviewUrl(url){ const raw=String(url||'').trim(); if(!raw || /[\s<>"']/.test(raw)) return ''; if(/^(https?:|mailto:|#|\/|\.\/|\.\.\/)/i.test(raw)) return escapeHtml(raw); return ''; }
function inlinePreview(text){ let h=escapeHtml(text); const code=[]; h=h.replace(/`([^`\n]+?)`/g,(_,c)=>{ const key=`\u0000${code.length}\u0000`; code.push(`<code>${c}</code>`); return key; }); h=h.replace(/!\[([^\]]*)\]\(([^)\s]+)(?:\s+"[^"]*")?\)/g,(_,alt,url)=>{ const safe=safePreviewUrl(url); return safe?`<img src="${safe}" alt="${alt}">`:alt; }); h=h.replace(/\[([^\]]+)\]\(([^)\s]+)(?:\s+"[^"]*")?\)/g,(_,label,url)=>{ const safe=safePreviewUrl(url); return safe?`<a href="${safe}" target="_blank" rel="noopener noreferrer">${label}</a>`:label; }); h=h.replace(/\[\[u:(.+?)\]\]/g,'<u>$1</u>'); h=h.replace(/\[\[sc:(.+?)\]\]/g,'<span class="sc">$1</span>'); h=h.replace(/\[\[sans:(.+?)\]\]/g,'<span class="sans">$1</span>'); h=h.replace(/\[\[mono:(.+?)\]\]/g,'<code>$1</code>'); h=h.replace(/~~(.+?)~~/g,'<del>$1</del>'); h=h.replace(/\*\*\*(.+?)\*\*\*/g,'<strong><em>$1</em></strong>'); h=h.replace(/\*\*(.+?)\*\*/g,'<strong>$1</strong>'); h=h.replace(/(^|[^\*])\*(?!\s)(.+?)(?<!\s)\*(?!\*)/g,'$1<em>$2</em>'); h=h.replace(/__(.+?)__/g,'<strong>$1</strong>'); h=h.replace(/(^|[^_])_(?!\s)(.+?)(?<!\s)_(?!_)/g,'$1<em>$2</em>'); code.forEach((value,idx)=>{ h=h.replace(`\u0000${idx}\u0000`,value); }); return h; }
function listItemInfo(line){ const m=line.match(/^(\s*)((?:[-+*])|(?:\d+[.)]))\s+(.+)$/); if(!m) return null; const indent=m[1].replace(/\t/g,'    ').length; const ordered=/^\d/.test(m[2]); let text=m[3]; let task=''; const tm=text.match(/^\[([ xX])\]\s+(.*)$/); if(tm){ task=tm[1].toLowerCase()==='x'?'checked':'unchecked'; text=tm[2]; } return {indent,type:ordered?'ol':'ul',text,task}; }
function renderListItems(items,start,indent,type){ let html=`<${type}>`, i=start; while(i<items.length){ const item=items[i]; if(item.indent<indent) break; if(item.indent>indent){ const nested=renderListItems(items,i,item.indent,item.type); html+=nested.html; i=nested.index; continue; } if(item.type!==type) break; const cls=item.task?' class="task"':''; const box=item.task?`<input type="checkbox" disabled ${item.task==='checked'?'checked':''}>`:''; html+=`<li${cls}>${box}${inlinePreview(item.text)}`; i++; while(i<items.length && items[i].indent>indent){ const nested=renderListItems(items,i,items[i].indent,items[i].type); html+=nested.html; i=nested.index; } html+='</li>'; } html+=`</${type}>`; return {html,index:i}; }
function renderListBlock(items){ let html='', i=0; while(i<items.length){ const rendered=renderListItems(items,i,items[i].indent,items[i].type); html+=rendered.html; i=rendered.index; } return html; }
function renderTable(lines){ if(lines.length<2 || !/^\s*\|?[\s:-]+\|[\s|:-]*$/.test(lines[1])) return null; const split=line=>line.trim().replace(/^\||\|$/g,'').split('|').map(c=>c.trim()); const heads=split(lines[0]); const rows=lines.slice(2).map(split); return `<table><thead><tr>${heads.map(h=>`<th>${inlinePreview(h)}</th>`).join('')}</tr></thead><tbody>${rows.map(r=>`<tr>${heads.map((_,i)=>`<td>${inlinePreview(r[i]||'')}</td>`).join('')}</tr>`).join('')}</tbody></table>`; }
function renderMarkdownPreviewLocal(){ const pane=$('#markdownPreview'); if(!pane) return; const text=editor.value; if(!text.trim()){ pane.innerHTML='<div class="empty">Preview appears here when text is loaded.</div>'; return; } const lines=text.split('\n'); let html='', para=[], block=null; const flushPara=()=>{ if(!para.length) return; html+=`<p>${inlinePreview(para.join(' '))}</p>`; para=[]; }; const closeBlock=()=>{ if(block){ flushPara(); html+='</div>'; block=null; } }; for(let i=0;i<lines.length;i++){ const raw=lines[i], t=raw.trim(); const fence=raw.match(/^\s*(```+|~~~+)\s*([a-z0-9_-]+)?\s*$/i); if(fence){ flushPara(); const mark=fence[1], buf=[]; i++; while(i<lines.length && !lines[i].trim().startsWith(mark)){ buf.push(lines[i]); i++; } html+=`<pre><code>${escapeHtml(buf.join('\n'))}</code></pre>`; continue; } const customBlock=raw.match(/^\s*:::\s*([a-z0-9_-]+)?\s*$/i); if(customBlock){ flushPara(); if(block){ closeBlock(); } else { block=canonicalStructureMarker(customBlock[1]||'block'); html+=`<div class="block"><div class="block-label">${escapeHtml(block)}</div>`; } continue; } if(!t){ flushPara(); continue; } const md=markdownLineInfo(raw); if(md?.marker==='scene'){ flushPara(); html+='<div class="scene">***</div>'; continue; } const heading=raw.match(/^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$/); if(heading){ flushPara(); const level=Math.min(6,heading[1].length); html+=`<h${level}>${inlinePreview(heading[2].trim())}</h${level}>`; continue; } if(/^\s{0,3}(?:-{3,}|\*{3,}|_{3,})\s*$/.test(raw)){ flushPara(); html+='<hr>'; continue; } if(/^\s{0,3}>\s?/.test(raw)){ flushPara(); const quoted=[]; while(i<lines.length && /^\s{0,3}>\s?/.test(lines[i])){ quoted.push(lines[i].replace(/^\s{0,3}>\s?/,'')); i++; } i--; html+=`<blockquote>${renderMarkdownFragment(quoted.join('\n'))}</blockquote>`; continue; } const li=listItemInfo(raw); if(li){ flushPara(); const items=[]; while(i<lines.length){ const next=listItemInfo(lines[i]); if(!next) break; items.push(next); i++; } i--; html+=renderListBlock(items); continue; } if(/\|/.test(raw) && i+1<lines.length){ const tableLines=[]; let j=i; while(j<lines.length && /\|/.test(lines[j]) && lines[j].trim()){ tableLines.push(lines[j]); j++; } const table=renderTable(tableLines); if(table){ flushPara(); html+=table; i=j-1; continue; } } para.push(t); } flushPara(); closeBlock(); pane.innerHTML=html; }
function setMarkdownPreviewHtml(html){ const pane=$('#markdownPreview'); if(pane) pane.innerHTML=html || '<div class="empty">Preview appears here when text is loaded.</div>'; }
function previewIsVisible(){ return document.querySelector('.editor-card')?.classList.contains('preview-on'); }
function renderMarkdownPreview(){ const pane=$('#markdownPreview'); if(!pane) return; if(!previewIsVisible()){ clearTimeout(markdownPreviewTimer); return; } const text=editor.value; const requestId=++markdownPreviewRequest; if(!text.trim()){ clearTimeout(markdownPreviewTimer); setMarkdownPreviewHtml(''); return; } if(!markdownServerAvailable){ renderMarkdownPreviewLocal(); return; } clearTimeout(markdownPreviewTimer); markdownPreviewTimer=setTimeout(async()=>{ try{ const res=await api('/api/render_markdown',{text}); if(!res.ok) throw new Error('Markdown renderer unavailable'); const data=await res.json(); if(requestId!==markdownPreviewRequest) return; setMarkdownPreviewHtml(data.html||''); }catch(err){ markdownServerAvailable=false; setTimeout(()=>{ markdownServerAvailable=true; },10000); renderMarkdownPreviewLocal(); } },160); }
function renderMarkdownFragment(text){ const lines=text.split('\n'); let out='', para=[]; const flush=()=>{ if(para.length){ out+=`<p>${inlinePreview(para.join(' '))}</p>`; para=[]; } }; for(const line of lines){ const t=line.trim(); if(!t){ flush(); continue; } const li=listItemInfo(line); if(li){ flush(); out+=renderListBlock([li]); continue; } para.push(t); } flush(); return out; }
function cleanFilename(name){ return (name||'manuscript').replace(/[^A-Za-z0-9._-]+/g,'-').replace(/^[-._]+|[-._]+$/g,'') || 'manuscript'; }
function downloadBlob(blob, name){ const a=document.createElement('a'); a.href=URL.createObjectURL(blob); a.download=name; document.body.appendChild(a); a.click(); setTimeout(()=>{URL.revokeObjectURL(a.href); a.remove();},0); }
function arrayBufferToBase64(buffer){ let binary=''; const bytes=new Uint8Array(buffer); const chunk=0x8000; for(let i=0;i<bytes.length;i+=chunk){ binary += String.fromCharCode.apply(null, bytes.subarray(i,i+chunk)); } return btoa(binary); }
function updateUndoButton(){ $('#undoBtn').disabled=!undoStack.length; }
function pushUndoSnapshot(){ undoStack.push({text:editor.value,start:editor.selectionStart,end:editor.selectionEnd,scrollTop:editor.scrollTop}); if(undoStack.length>20) undoStack.shift(); updateUndoButton(); }
function undoLastOperation(){ const snap=undoStack.pop(); if(!snap) return; editor.value=snap.text; editor.focus(); editor.setSelectionRange(snap.start,snap.end); updateStats(); restoreEditorScroll(snap.scrollTop||0); updateUndoButton(); toast('Undone'); }
$('#undoBtn').addEventListener('click',undoLastOperation);

function activateSideTab(name){ document.querySelectorAll('[data-side-tab]').forEach(b=>b.classList.toggle('active',b.dataset.sideTab===name)); document.querySelectorAll('[data-side-panel]').forEach(p=>p.classList.toggle('active',p.dataset.sidePanel===name)); localStorage.setItem('vp.py.sideTab',name); }
document.querySelectorAll('[data-side-tab]').forEach(btn=>btn.addEventListener('click',()=>activateSideTab(btn.dataset.sideTab)));

function updateStats(){ const text=editor.value; $('#statChars').textContent=text.length.toLocaleString(); $('#statWords').textContent=(text.match(/\S+/g)||[]).length.toLocaleString(); $('#statLines').textContent=text.split('\n').length.toLocaleString(); $('#statChapters').textContent=(text.match(/^\s{0,3}#\s+\S/gm)||[]).length.toLocaleString(); $('#statSubheads').textContent=(text.match(/^\s{0,3}#{2,6}\s+\S/gm)||[]).length.toLocaleString(); const diag=diagnose(text); renderGutter(diag); renderDiagnose(diag); renderOutline(); renderVellumStructure(); renderMarkdownPreview(); scheduleAutosave(); }
function renderGutter(diag){ const lines=diag ? diag.lines : editor.value.split('\n'); if(lines.length>12000){ gutter.innerHTML='<div class="gline">…</div>'; return; } const frag=document.createDocumentFragment(); for(let i=0;i<lines.length;i++){ const f=diag ? diag.flags[i] : new Set(); const d=document.createElement('div'); let cls='gline'; const notes=[]; if(f.has('mojibake')||f.has('invisible')||f.has('control')||f.has('specialSpace')||f.has('ligature')){ cls+=' issue-violet'; notes.push('encoding/hidden character issue'); } else if(f.has('trailingWs')||f.has('tabs')||f.has('extraBlank')){ cls+=' issue-rust'; notes.push('spacing issue'); } else if(f.has('wrap')||f.has('pageNum')||f.has('repeatLine')){ cls+=' issue-amber'; notes.push('structure candidate'); } else if(f.has('doubleSpace')||f.has('spaceBeforePunct')){ cls+=' issue-blue'; notes.push('typography spacing'); } else if(markdownLineInfo(lines[i])||/^\s*:::\s*[a-z0-9_-]*\s*$/i.test(lines[i])) cls+=' issue'; d.className=cls; if(notes.length) d.title=notes.join(', '); d.textContent=i+1; frag.appendChild(d); } gutter.innerHTML=''; gutter.appendChild(frag); gutter.scrollTop=editor.scrollTop; }
function syncScroll(source,target){ if(scrollSyncing||!target) return; scrollSyncing=true; const maxSource=Math.max(1,source.scrollHeight-source.clientHeight); const maxTarget=Math.max(0,target.scrollHeight-target.clientHeight); target.scrollTop=(source.scrollTop/maxSource)*maxTarget; requestAnimationFrame(()=>{ scrollSyncing=false; }); }
editor.addEventListener('scroll',()=>{ gutter.scrollTop=editor.scrollTop; syncFindHighlights(); const pane=$('#markdownPreview'); if(pane && document.querySelector('.editor-card')?.classList.contains('preview-on')) syncScroll(editor,pane); });
$('#markdownPreview').addEventListener('scroll',()=>{ if(!document.querySelector('.editor-card')?.classList.contains('preview-on')) return; syncScroll($('#markdownPreview'),editor); gutter.scrollTop=editor.scrollTop; });
editor.addEventListener('keydown',handleTabIndent);
editor.addEventListener('input',()=>{ updateStats(); if($('#findModal')?.classList.contains('open')) refreshFindMatches(editor.selectionStart); });
editor.addEventListener('select', updateSelInfo); editor.addEventListener('click', updateSelInfo); editor.addEventListener('keyup', updateSelInfo); editor.addEventListener('mouseup', updateSelInfo);
document.addEventListener('keydown',e=>{ if((e.metaKey||e.ctrlKey) && e.key.toLowerCase()==='f'){ e.preventDefault(); openFindModal(); } else if(e.key==='Escape' && $('#findModal')?.classList.contains('open')) closeFindModal(); });
$('#openFindBtn').addEventListener('click',openFindModal);
$('#closeFindBtn').addEventListener('click',closeFindModal);
$('#findNextBtn').addEventListener('click',()=>findMove(1));
$('#findPrevBtn').addEventListener('click',()=>findMove(-1));
$('#replaceOneBtn').addEventListener('click',replaceCurrentFind);
$('#replaceAllBtn').addEventListener('click',replaceAllFind);
['findText','replaceText'].forEach(id=>$('#'+id).addEventListener('keydown',e=>{ if(e.key==='Enter'){ e.preventDefault(); findMove(e.shiftKey?-1:1); } }));
['findText','findRegex','findCase'].forEach(id=>$('#'+id).addEventListener('input',()=>refreshFindMatches(editor.selectionEnd)));
$('#findScope').addEventListener('change',()=>{ if($('#findScope').value==='selection'){ findScopeRange=editor.selectionStart!==editor.selectionEnd?{start:editor.selectionStart,end:editor.selectionEnd}:findScopeRange; } else findScopeRange=null; refreshFindMatches(editor.selectionEnd); });
function updateSelInfo(){ const s=editor.selectionStart,e=editor.selectionEnd; $('#selInfo').textContent=s===e?'no selection':`${(e-s).toLocaleString()} chars selected`; updateSelectionMenu(); }
function updateSelectionMenu(){ const menu=$('#selectionMenu'); if(!menu) return; if($('#findModal')?.classList.contains('open')){ menu.hidden=true; return; } const hasSelection=editor.selectionStart!==editor.selectionEnd && document.activeElement===editor; if(!hasSelection){ if(!menu.matches(':focus-within')){ menu.hidden=true; $('#selectionAction').value=''; savedSelection=null; } return; } savedSelection={start:editor.selectionStart,end:editor.selectionEnd}; const r=editor.getBoundingClientRect(); menu.hidden=false; const width=Math.min(520, Math.max(320, r.width*.55)); menu.style.width=width+'px'; menu.style.left=Math.max(10, Math.min(window.innerWidth-width-10, r.right-width-14))+'px'; menu.style.top=Math.max(10, r.top+10)+'px'; }
document.addEventListener('selectionchange',()=>{ if(document.activeElement===editor) updateSelectionMenu(); });

$('#wrapToggle').addEventListener('change',()=>{ const on=$('#wrapToggle').checked; editor.classList.toggle('wrap',on); $('#editorArea').classList.toggle('wrap',on); editor.setAttribute('wrap',on?'soft':'off'); syncFindHighlights(); saveSettings(); });
$('#togglePreview').addEventListener('click',()=>{ const card=document.querySelector('.editor-card'); const on=!card.classList.contains('preview-on'); card.classList.toggle('preview-on',on); $('#togglePreview').classList.toggle('active',on); if(on) markdownServerAvailable=true; renderMarkdownPreview(); saveSettings(); });
document.querySelectorAll('input[name="scope"]').forEach(input=>input.addEventListener('change',saveSettings));
$('#toggleVellum').addEventListener('click',()=>{ const tb=$('#vellumToolbar'); tb.classList.toggle('collapsed'); $('#toggleVellum').textContent=tb.classList.contains('collapsed')?'Expand':'Collapse'; saveSettings(); });

function lineOffsetsFor(text){ const lines=text.split('\n'); const offs=[]; let acc=0; for(let i=0;i<lines.length;i++){ offs.push(acc); acc += lines[i].length+1; } return {lines,offs}; }
function wordCount(text){ return (String(text||'').match(/\S+/g)||[]).length; }
function headingOutline(text){ const {lines,offs}=lineOffsetsFor(text); const rows=[]; for(let i=0;i<lines.length;i++){ const h=lines[i].match(/^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$/); if(!h) continue; rows.push({line:i+1,level:h[1].length,title:h[2].trim(),start:offs[i],lineEnd:offs[i]+lines[i].length}); } for(let i=0;i<rows.length;i++){ const row=rows[i]; let end=text.length; for(let j=i+1;j<rows.length;j++){ if(rows[j].level<=row.level){ end=Math.max(row.lineEnd, rows[j].start-1); break; } } row.end=end; row.words=wordCount(text.slice(row.lineEnd,row.end)); } return rows; }
function outlineKey(row){ return `${row.level}:${row.start}:${row.title}`; }
function clamp(n,min,max){ return Math.max(min,Math.min(max,n)); }
function measureEditorOffsetTop(pos){ const style=getComputedStyle(editor); const mirror=document.createElement('div'); const before=editor.value.slice(0,Math.max(0,Math.min(pos,editor.value.length))); mirror.style.position='absolute'; mirror.style.visibility='hidden'; mirror.style.pointerEvents='none'; mirror.style.left='-10000px'; mirror.style.top='0'; mirror.style.width=editor.clientWidth+'px'; mirror.style.boxSizing='border-box'; mirror.style.fontFamily=style.fontFamily; mirror.style.fontSize=style.fontSize; mirror.style.fontWeight=style.fontWeight; mirror.style.letterSpacing=style.letterSpacing; mirror.style.lineHeight=style.lineHeight; mirror.style.padding=style.padding; mirror.style.border='0'; mirror.style.tabSize=style.tabSize || '4'; mirror.style.whiteSpace=editor.classList.contains('wrap')?'pre-wrap':'pre'; mirror.style.overflowWrap=editor.classList.contains('wrap')?'break-word':'normal'; mirror.textContent=before; const mark=document.createElement('span'); mark.textContent='\u200b'; mirror.appendChild(mark); document.body.appendChild(mirror); const top=mark.offsetTop; mirror.remove(); return top; }
function scrollEditorToOffset(pos){ const top=measureEditorOffsetTop(pos); const max=Math.max(0,editor.scrollHeight-editor.clientHeight); editor.scrollTop=clamp(top-editor.clientHeight*.45,0,max); gutter.scrollTop=editor.scrollTop; syncFindHighlights(); }
function jumpToOffset(pos){ editor.focus(); editor.setSelectionRange(pos,pos); scrollEditorToOffset(pos); updateSelInfo(); }
function selectChapterRange(row){ editor.focus(); editor.setSelectionRange(row.start,row.end); editor.scrollTop=Math.max(0,(row.line-1)*20-editor.clientHeight/2); gutter.scrollTop=editor.scrollTop; updateSelInfo(); }
function renderOutline(){ const list=$('#outlineList'); if(!list) return; const rows=headingOutline(editor.value); list.innerHTML=''; if(!rows.length){ list.innerHTML='<div class="empty-hint">No # headings yet. Mark chapters with # to build an outline.</div>'; return; } let hiddenLevel=null; rows.forEach((row,index)=>{ if(hiddenLevel!==null){ if(row.level>hiddenLevel) return; hiddenLevel=null; } const key=outlineKey(row); const hasChildren=index<rows.length-1 && rows[index+1].level>row.level; const collapsed=collapsedOutlineHeadings.has(key); const div=document.createElement('div'); div.className=`outline-row l${row.level}`+(row.level>1?' sub':''); div.style.paddingLeft=`${16 + Math.min(Math.max(row.level-1,0),4)*14}px`; const tag=row.level===1?'CH':`H${row.level}`; const title=row.title||'(untitled)'; const toggle=hasChildren?`<button data-act="toggle" class="outline-toggle" title="${collapsed?'Expand':'Collapse'}">${collapsed?'+':'-'}</button>`:''; div.innerHTML=`<span class="outline-tag">${tag}</span><div class="outline-main"><div class="outline-title" title="${escapeHtml(title)}">${escapeHtml(title)}</div><div class="outline-meta">line ${row.line} · ${row.words.toLocaleString()} words</div></div><div class="outline-actions">${toggle}<button data-act="jump">Go</button>${row.level===1?'<button data-act="rewrite">AI</button><button data-act="clean">Clean</button>':''}</div>`; const toggleBtn=div.querySelector('[data-act="toggle"]'); if(toggleBtn) toggleBtn.addEventListener('click',()=>{ if(collapsedOutlineHeadings.has(key)) collapsedOutlineHeadings.delete(key); else collapsedOutlineHeadings.add(key); renderOutline(); }); div.querySelector('[data-act="jump"]').addEventListener('click',()=>jumpToOffset(row.start)); const rewrite=div.querySelector('[data-act="rewrite"]'); if(rewrite) rewrite.addEventListener('click',()=>{ selectChapterRange(row); startLM('rewrite','selection',false); }); const clean=div.querySelector('[data-act="clean"]'); if(clean) clean.addEventListener('click',()=>{ selectChapterRange(row); startLM('clean','selection',false); }); list.appendChild(div); if(collapsed) hiddenLevel=row.level; }); }
function currentParagraphRange(){ const val=editor.value; const {lines,offs}=lineOffsetsFor(val); const pos=editor.selectionStart; let li=0; for(let i=0;i<offs.length;i++){ if(offs[i]<=pos) li=i; else break; } let startLine=li,endLine=li; if(lines[li] && lines[li].trim()!==''){ while(startLine>0 && lines[startLine-1].trim()!=='') startLine--; while(endLine<lines.length-1 && lines[endLine+1].trim()!=='') endLine++; } const start=offs[startLine]||0; const end=(offs[endLine]||0)+(lines[endLine]||'').length; return {start,end}; }
function selectedLineRange(){ const s=editor.selectionStart,e=editor.selectionEnd; const val=editor.value; const start=val.lastIndexOf('\n',s-1)+1; let end=val.indexOf('\n', e); if(end<0) end=val.length; return {start,end}; }
function restoreEditorScroll(top){ editor.scrollTop=top; gutter.scrollTop=top; requestAnimationFrame(()=>{ editor.scrollTop=top; gutter.scrollTop=top; }); }
function replaceRange(start,end,text){ const top=editor.scrollTop; pushUndoSnapshot(); editor.value=editor.value.slice(0,start)+text+editor.value.slice(end); editor.focus(); editor.setSelectionRange(start,start+text.length); updateStats(); restoreEditorScroll(top); updateSelInfo(); }
function escapeRegExp(text){ return String(text||'').replace(/[.*+?^${}()|[\]\\]/g,'\\$&'); }
function syncFindHighlights(){ const layer=$('#findHighlights'); if(!layer || !layer.classList.contains('show')) return; const area=$('#editorArea').getBoundingClientRect(), r=editor.getBoundingClientRect(); layer.style.left=(r.left-area.left)+'px'; layer.style.top=(r.top-area.top)+'px'; layer.style.width=editor.clientWidth+'px'; layer.style.height=editor.clientHeight+'px'; layer.classList.toggle('wrap',editor.classList.contains('wrap')); layer.scrollTop=editor.scrollTop; layer.scrollLeft=editor.scrollLeft; }
function hideFindHighlights(){ const layer=$('#findHighlights'); if(layer){ layer.classList.remove('show'); layer.innerHTML=''; } }
function renderFindHighlights(){ const layer=$('#findHighlights'); if(!layer) return; if(!$('#findModal')?.classList.contains('open') || !findMatches.length){ hideFindHighlights(); return; } const text=editor.value; let html='', pos=0; findMatches.forEach((m,idx)=>{ if(m.start<pos) return; html += escapeHtml(text.slice(pos,m.start)); html += `<span class="find-hit ${idx===findIndex?'current':''}">${escapeHtml(text.slice(m.start,m.end))}</span>`; pos=m.end; }); html += escapeHtml(text.slice(pos)); layer.innerHTML=html || '&nbsp;'; layer.classList.add('show'); syncFindHighlights(); }
function findOptions(){ return {pattern:$('#findText').value, replacement:$('#replaceText').value, regex:$('#findRegex').checked, caseSensitive:$('#findCase').checked, scope:$('#findScope').value}; }
function findRange(){ const opts=findOptions(); if(opts.scope==='selection'){ if(!findScopeRange || findScopeRange.start===findScopeRange.end){ return null; } return {start:findScopeRange.start,end:Math.min(findScopeRange.end,editor.value.length)}; } return {start:0,end:editor.value.length}; }
function buildFindRegex(global=true){ const opts=findOptions(); if(!opts.pattern){ $('#findStatus').textContent='Enter text to find.'; return null; } const source=opts.regex ? opts.pattern : escapeRegExp(opts.pattern); const flags=(global?'g':'')+(opts.caseSensitive?'':'i'); try{ return new RegExp(source,flags); }catch(err){ $('#findStatus').textContent=`Regex error: ${err.message}`; return null; } }
function refreshFindMatches(anchor=null){ const range=findRange(); const re=buildFindRegex(true); findMatches=[]; findIndex=-1; if(!range || !re){ hideFindHighlights(); if(!range) $('#findStatus').textContent='Select text before using Selection only.'; return false; } const text=editor.value.slice(range.start,range.end); let m; while((m=re.exec(text))){ if(m[0].length===0){ re.lastIndex++; continue; } findMatches.push({start:range.start+m.index,end:range.start+m.index+m[0].length}); } if(!findMatches.length){ hideFindHighlights(); $('#findStatus').textContent='No matches.'; return false; } if(anchor!==null){ findIndex=findMatches.findIndex(m=>m.start>=anchor); if(findIndex<0) findIndex=0; } else findIndex=0; $('#findStatus').textContent=`${findMatches.length.toLocaleString()} match${findMatches.length===1?'':'es'}.`; renderFindHighlights(); return true; }
function revealFindMatch(){ const m=findMatches[findIndex]; if(!m) return; editor.focus(); editor.setSelectionRange(m.start,m.end); scrollEditorToOffset(m.start); renderFindHighlights(); updateSelInfo(); $('#findStatus').textContent=`Match ${findIndex+1} of ${findMatches.length}`; }
function findMove(dir){ const anchor=editor.selectionEnd; if(!findMatches.length || !findMatches.some(m=>m.start===editor.selectionStart&&m.end===editor.selectionEnd)){ if(!refreshFindMatches(dir>0?anchor:editor.selectionStart)) return; } else { findIndex=(findIndex+dir+findMatches.length)%findMatches.length; } revealFindMatch(); }
function openFindModal(){ const s=editor.selectionStart,e=editor.selectionEnd, selected=editor.value.slice(s,e); findScopeRange=s!==e?{start:s,end:e}:null; $('#findScope').value='all'; if(selected && selected.length<=120 && !selected.includes('\n')) $('#findText').value=selected; $('#findModal').classList.add('open'); setTimeout(()=>{$('#findText').focus(); $('#findText').select(); refreshFindMatches(editor.selectionEnd);},0); }
function closeFindModal(){ $('#findModal').classList.remove('open'); hideFindHighlights(); }
function replaceCurrentFind(){ if(!findMatches.length && !refreshFindMatches(editor.selectionStart)) return; const m=findMatches[findIndex]; if(!m) return; const opts=findOptions(), top=editor.scrollTop, val=editor.value, matchText=val.slice(m.start,m.end); const re=opts.regex ? buildFindRegex(false) : null; if(opts.regex && !re) return; const replacement=opts.regex ? matchText.replace(re,opts.replacement) : opts.replacement; pushUndoSnapshot(); editor.value=val.slice(0,m.start)+replacement+val.slice(m.end); updateStats(); restoreEditorScroll(top); findScopeRange=findScopeRange?{start:findScopeRange.start,end:findScopeRange.end+(replacement.length-matchText.length)}:null; refreshFindMatches(m.start+replacement.length); revealFindMatch(); toast('Replaced'); }
function replaceAllFind(){ const range=findRange(); const opts=findOptions(); const re=buildFindRegex(true); if(!range || !re) return; refreshFindMatches(); const count=findMatches.length; if(!count){ $('#findStatus').textContent='No matches.'; return; } const top=editor.scrollTop, val=editor.value, chunk=val.slice(range.start,range.end); const replaced=opts.regex ? chunk.replace(re,opts.replacement) : chunk.replace(re,()=>opts.replacement); pushUndoSnapshot(); editor.value=val.slice(0,range.start)+replaced+val.slice(range.end); updateStats(); restoreEditorScroll(top); editor.focus(); editor.setSelectionRange(range.start,range.start+replaced.length); findScopeRange=opts.scope==='selection'?{start:range.start,end:range.start+replaced.length}:null; findMatches=[]; findIndex=-1; hideFindHighlights(); $('#findStatus').textContent=`Replaced ${count.toLocaleString()} match${count===1?'':'es'}.`; updateSelInfo(); toast(`Replaced ${count.toLocaleString()}`); }
function lineRangeForSelection(){ const val=editor.value, s=editor.selectionStart, e=editor.selectionEnd; const effectiveEnd=e>s && val[e-1]==='\n' ? e-1 : e; const start=val.lastIndexOf('\n',s-1)+1; let end=val.indexOf('\n',effectiveEnd); if(end<0) end=val.length; return {start,end}; }
function removeOneIndent(line){ if(line.startsWith('\t')) return {line:line.slice(1),removed:1}; const spaces=(line.match(/^ {1,4}/)||[''])[0].length; return {line:line.slice(spaces),removed:spaces}; }
function handleTabIndent(event){ if(event.key!=='Tab' || event.metaKey || event.ctrlKey || event.altKey) return; event.preventDefault(); const val=editor.value, s=editor.selectionStart, e=editor.selectionEnd, top=editor.scrollTop, unit='    '; pushUndoSnapshot(); if(s===e && !event.shiftKey){ editor.value=val.slice(0,s)+unit+val.slice(e); editor.focus(); editor.setSelectionRange(s+unit.length,s+unit.length); updateStats(); restoreEditorScroll(top); updateSelInfo(); return; } const range=lineRangeForSelection(); const block=val.slice(range.start,range.end); const lines=block.split('\n'); let deltaBeforeStart=0, totalDelta=0, lineOffset=range.start; const changed=lines.map(line=>{ const before=line; let after=line, delta=0; if(event.shiftKey){ const res=removeOneIndent(line); after=res.line; delta=-res.removed; } else { after=unit+line; delta=unit.length; } if(lineOffset<s) deltaBeforeStart += delta; totalDelta += delta; lineOffset += before.length+1; return after; }).join('\n'); editor.value=val.slice(0,range.start)+changed+val.slice(range.end); const newStart=Math.max(range.start, s+deltaBeforeStart); const newEnd=Math.max(newStart, e+totalDelta); editor.focus(); editor.setSelectionRange(newStart,newEnd); updateStats(); restoreEditorScroll(top); updateSelInfo(); }
function markLines(marker){ const r=selectedLineRange(); const block=editor.value.slice(r.start,r.end); if(MD_PREFIX_MARKERS[marker]){ const prefix=MD_PREFIX_MARKERS[marker]; const lines=block.split('\n').map(line=>{ if(!line.trim()) return marker==='quote' ? '>' : line; let content=stripStructureLine(line).replace(/^\s+/,''); if(marker==='part' && !isPartHeadingText(content)) content=`Part ${content}`; return prefix + content; }); replaceRange(r.start,r.end,lines.join('\n')); toast(`Marked as ${marker}`); return; } if(MD_BLOCK_MARKERS.has(marker)){ const body=block.split('\n').filter(line=>!/^\s*:::\s*[a-z0-9_-]*\s*$/i.test(line)).map(stripStructureLine).join('\n').trim(); const wrapped=`::: ${marker}\n${body || ''}\n:::`; replaceRange(r.start,r.end,wrapped); toast(`Marked as ${marker}`); } }
document.querySelectorAll('[data-marker]').forEach(btn=>btn.addEventListener('click',()=>markLines(btn.dataset.marker)));
document.querySelectorAll('[data-insert]').forEach(btn=>btn.addEventListener('click',()=>{ const insert='***'; const pos=editor.selectionStart; replaceRange(pos,pos,(pos&&editor.value[pos-1]!='\n'?'\n':'')+insert+'\n'); }));
document.querySelectorAll('[data-inline]').forEach(btn=>btn.addEventListener('click',()=>applyInline(btn.dataset.inline)));
function clearVellumMark(){ const r=selectedLineRange(); const block=editor.value.slice(r.start,r.end); const cleared=block.split('\n').filter(line=>!/^\s*:::\s*[a-z0-9_-]*\s*$/i.test(line)).map(stripStructureLine).join('\n'); if(cleared===block){ toast('No structure marker on selected/current line.'); return; } replaceRange(r.start,r.end,cleared); toast('Structure marker removed'); }
$('#clearVellumBtn').addEventListener('click',clearVellumMark);
function inlineWrapper(kind){ if(kind==='italic') return {prefix:'*',suffix:'*'}; if(kind==='bold') return {prefix:'**',suffix:'**'}; if(kind==='mono') return {prefix:'`',suffix:'`'}; return {prefix:`[[${kind}:`,suffix:']]'}; }
function countEdgeChars(text, char, fromStart){ let n=0, i=fromStart?0:text.length-1; while(i>=0 && i<text.length && text[i]===char){ n++; i += fromStart?1:-1; } return n; }
function countEditorChars(pos, char, dir){ let n=0, i=dir<0?pos-1:pos; while(i>=0 && i<editor.value.length && editor.value[i]===char){ n++; i += dir; } return n; }
function unwrapInlineText(text, kind){ const w=inlineWrapper(kind); let out=text, changed=false; if(kind==='italic' && countEdgeChars(out,'*',true)===2 && countEdgeChars(out,'*',false)===2) return null; while(out.startsWith(w.prefix) && out.endsWith(w.suffix) && out.length>=w.prefix.length+w.suffix.length){ out=out.slice(w.prefix.length,out.length-w.suffix.length); changed=true; if(kind==='italic' || kind==='mono') break; } return changed ? out : null; }
function unwrapAdjacentInline(start,end,kind){ const w=inlineWrapper(kind), val=editor.value; const beforeStart=start-w.prefix.length, afterEnd=end+w.suffix.length; if(beforeStart<0 || afterEnd>val.length) return false; if(val.slice(beforeStart,start)!==w.prefix || val.slice(end,afterEnd)!==w.suffix) return false; if(kind==='italic' && countEditorChars(start,'*',-1)===2 && countEditorChars(end,'*',1)===2) return false; replaceRange(beforeStart,afterEnd,val.slice(start,end)); return true; }
function applyInline(kind){ const s=editor.selectionStart,e=editor.selectionEnd; if(s===e){ toast('Select text first.'); return; } const sel=editor.value.slice(s,e); const unwrapped=unwrapInlineText(sel,kind); if(unwrapped!==null){ replaceRange(s,e,unwrapped); return; } if(unwrapAdjacentInline(s,e,kind)) return; let out=sel; if(kind==='mono') out='`'+sel.replace(/`/g,'')+'`'; else { const w=inlineWrapper(kind); out=w.prefix+sel+w.suffix; } replaceRange(s,e,out); }
function restoreSavedSelection(){ if(savedSelection){ editor.focus(); editor.setSelectionRange(savedSelection.start,savedSelection.end); } }
function restoreSelectionForQuickAction(){ if(savedSelection){ restoreSavedSelection(); } else { editor.focus(); } }
function applyQuickStyle(marker){ restoreSelectionForQuickAction(); if(marker==='normal') clearVellumMark(); else if(marker) markLines(marker); $('#quickStyle').value=''; updateSelectionMenu(); }
function applyQuickIcon(icon){ if(!icon) return; restoreSelectionForQuickAction(); const s=editor.selectionStart,e=editor.selectionEnd, sel=editor.value.slice(s,e); replaceRange(s,e,`${icon} ${sel}`); $('#quickIcon').value=''; updateSelectionMenu(); }
$('#quickStyle').addEventListener('change',()=>applyQuickStyle($('#quickStyle').value));
$('#quickIcon').addEventListener('change',()=>applyQuickIcon($('#quickIcon').value));
document.querySelectorAll('[data-quick-inline]').forEach(btn=>{ btn.addEventListener('mousedown',e=>e.preventDefault()); btn.addEventListener('click',()=>{ restoreSelectionForQuickAction(); applyInline(btn.dataset.quickInline); updateSelectionMenu(); }); });
$('#quickAiRewrite').addEventListener('mousedown',e=>e.preventDefault());
$('#quickAiRewrite').addEventListener('click',()=>{ restoreSelectionForQuickAction(); startLM('rewrite','selection',false); });
$('#quickAiClean').addEventListener('mousedown',e=>e.preventDefault());
$('#quickAiClean').addEventListener('click',()=>{ restoreSelectionForQuickAction(); startLM('clean','selection',false); });
$('#quickClearMark').addEventListener('mousedown',e=>e.preventDefault());
$('#quickClearMark').addEventListener('click',()=>{ restoreSelectionForQuickAction(); clearVellumMark(); updateSelectionMenu(); });
function runSelectionAction(){ const action=$('#selectionAction').value; if(!action) return; restoreSavedSelection(); if(action==='ai-rewrite') startLM('rewrite','selection',false); else if(action==='ai-clean') startLM('clean','selection',false); else if(action.startsWith('mark-')) markLines(action.slice(5)); else if(action==='clear-vellum') clearVellumMark(); else if(action.startsWith('inline-')) applyInline(action.slice(7)); else if(action==='safe-clean'){ const start=editor.selectionStart,end=editor.selectionEnd; if(start===end){ toast('Select text first.'); return; } const before=editor.value.slice(start,end); const after=safeCleanup(before); applyTextChange(before,after,start,end,'Safe Cleanup applied to selection'); } $('#selectionAction').value=''; savedSelection=null; $('#selectionMenu').hidden=true; }
$('#selectionRunBtn').addEventListener('click',runSelectionAction);
$('#selectionAction').addEventListener('change',runSelectionAction);

function isHtmlFile(file){ return /\.(html?|xhtml)$/i.test(file.name) || /^text\/html\b/i.test(file.type||''); }
function fileDisplaySize(size){ return typeof size==='number' ? `${(size/1024).toFixed(1)} kB` : ''; }
function applyImportedText(data, displayName, sizeLabel=''){ if(editor.value && !confirm(`Replace the current editor text with "${displayName||data.filename||'the imported file'}"?`)) return false; editor.value=data.text||''; currentFilename=data.filename||displayName||'manuscript.txt'; filenameEl.textContent=sizeLabel?`${displayName} (${sizeLabel})`:displayName; $('#metaFilename').value=cleanFilename((data.filename||displayName||'manuscript').replace(/\.[^.]+$/,'')); undoStack.length=0; updateUndoButton(); updateStats(); toast((data.notes||[]).join(' · ')||'File loaded'); return true; }
async function importFileObject(file){ const b64=arrayBufferToBase64(await file.arrayBuffer()); const res=await api('/api/import',{filename:file.webkitRelativePath||file.name,data_b64:b64,html_mode:isHtmlFile(file)?'main':'simple'}); const data=await res.json(); if(!res.ok) throw new Error(data.error||`Import failed: ${file.name}`); return {file,name:file.name,path:file.webkitRelativePath||file.name,size:file.size,text:data.text||'',notes:data.notes||[],filename:data.filename||file.name}; }
async function importHtmlFileOptions(file){ const b64=arrayBufferToBase64(await file.arrayBuffer()); const res=await api('/api/import_html_options',{filename:file.webkitRelativePath||file.name,data_b64:b64}); const data=await res.json(); if(!res.ok) throw new Error(data.error||`HTML import failed: ${file.name}`); openHtmlImportChoice(data,{displayName:file.name,size:file.size}); }
function openHtmlImportChoice(data, source){ htmlImportState={data,source:source||{}}; const simple=data.simple||{}, main=data.main||{}; $('#htmlImportTitle').textContent=source?.url?'Choose URL Import':'Choose HTML Import'; $('#htmlImportSub').textContent=source?.url ? `${source.url}` : `${source?.displayName||data.filename||'HTML'} - compare extraction methods before loading`; $('#htmlSimpleBox').textContent=simple.text||'(No text extracted.)'; $('#htmlMainBox').textContent=main.text || (main.error?`Trafilatura unavailable:\n${main.error}`:'(No main text extracted.)'); $('#htmlSimpleMeta').textContent=`${(simple.chars||0).toLocaleString()} chars`; $('#htmlMainMeta').textContent=main.error&&!main.text?'unavailable':`${(main.chars||0).toLocaleString()} chars`; $('#htmlUseSimpleBtn').disabled=!simple.text; $('#htmlUseMainBtn').disabled=!main.text; $('#htmlImportStatus').textContent=(data.notes||[]).join(' · ')||'Choose an import method.'; $('#htmlImportModal').classList.add('open'); }
function useHtmlImport(kind){ if(!htmlImportState) return; const {data,source}=htmlImportState; const chosen=(kind==='main'?data.main:data.simple)||{}; if(!chosen.text){ toast('No text available for that method.'); return; } if(!applyImportedText({text:chosen.text,filename:data.filename,notes:[kind==='main'?'Imported with Trafilatura main text':'Imported with simple extraction']}, source.displayName||data.filename||'HTML import', source.size?fileDisplaySize(source.size):'')) return; $('#htmlImportModal').classList.remove('open'); htmlImportState=null; }
$('#htmlUseSimpleBtn').addEventListener('click',()=>useHtmlImport('simple'));
$('#htmlUseMainBtn').addEventListener('click',()=>useHtmlImport('main'));
$('#htmlImportCancelBtn').addEventListener('click',()=>{ $('#htmlImportModal').classList.remove('open'); htmlImportState=null; });
function compareStats(text){ return `${String(text||'').length.toLocaleString()} chars · ${(String(text||'').match(/\S+/g)||[]).length.toLocaleString()} words`; }
function updateCompareMeta(){ const original=$('#compareOriginal').value, revised=$('#compareRevised').value; $('#compareOriginalMeta').textContent=original?compareStats(original):'empty'; $('#compareRevisedMeta').textContent=revised?compareStats(revised):'empty'; $('#compareSub').textContent=`Original ${compareStats(original)} · Revised ${compareStats(revised)}`; }
function setCompareText(side,text,label=''){ const el=side==='original'?$('#compareOriginal'):$('#compareRevised'); el.value=text||''; updateCompareMeta(); if(label) $('#compareStatus').textContent=`Loaded ${label} into ${side}.`; }
function openCompare(original='',revised='',status='Ready.'){ if(original!=='' || revised!==''){ setCompareText('original',original); setCompareText('revised',revised); } updateCompareMeta(); $('#compareStatus').textContent=status; $('#compareModal').classList.add('open'); }
function syncCompareScroll(source,target){ if(compareScrollSyncing || !$('#compareLinked').checked || !target) return; compareScrollSyncing=true; const maxSource=Math.max(1,source.scrollHeight-source.clientHeight); const maxTarget=Math.max(0,target.scrollHeight-target.clientHeight); target.scrollTop=(source.scrollTop/maxSource)*maxTarget; requestAnimationFrame(()=>{ compareScrollSyncing=false; }); }
async function loadCompareFile(file,side){ if(!file) return; try{ $('#compareStatus').textContent=`Loading ${file.name}...`; const data=await importFileObject(file); setCompareText(side,data.text||'',file.name); }catch(err){ $('#compareStatus').textContent=err.message; toast(err.message); alert(err.message); } }
$('#compareBtn').addEventListener('click',()=>openCompare('', '', 'Ready.'));
$('#closeCompareBtn').addEventListener('click',()=>$('#compareModal').classList.remove('open'));
$('#compareOriginal').addEventListener('input',updateCompareMeta); $('#compareRevised').addEventListener('input',updateCompareMeta);
$('#compareOriginal').addEventListener('scroll',()=>syncCompareScroll($('#compareOriginal'),$('#compareRevised')));
$('#compareRevised').addEventListener('scroll',()=>syncCompareScroll($('#compareRevised'),$('#compareOriginal')));
$('#compareOriginalInput').addEventListener('change',async e=>{ await loadCompareFile(e.target.files[0],'original'); e.target.value=''; });
$('#compareRevisedInput').addEventListener('change',async e=>{ await loadCompareFile(e.target.files[0],'revised'); e.target.value=''; });
$('#compareUseEditorOriginal').addEventListener('click',()=>setCompareText('original',editor.value,'editor text'));
$('#compareUseEditorRevised').addEventListener('click',()=>setCompareText('revised',editor.value,'editor text'));
$('#compareSwapBtn').addEventListener('click',()=>{ const a=$('#compareOriginal').value,b=$('#compareRevised').value; $('#compareOriginal').value=b; $('#compareRevised').value=a; updateCompareMeta(); $('#compareStatus').textContent='Swapped original and revised.'; });
$('#compareExportRevisedBtn').addEventListener('click',()=>{ const text=$('#compareRevised').value; if(!text.trim()){ toast('Revised side is empty.'); return; } const base=cleanFilename($('#metaFilename').value||currentFilename.replace(/\.[^.]+$/,'')||'revised'); downloadBlob(new Blob([text],{type:'text/markdown;charset=utf-8'}), `${base}-revised.md`); });
$('#compareCopyRevisedBtn').addEventListener('click',()=>{ const text=$('#compareRevised').value; if(!text.trim()){ toast('Revised side is empty.'); return; } if(editor.value && editor.value!==text && !confirm('Replace the current editor text with the revised side?')) return; replaceRange(0,editor.value.length,text); $('#compareStatus').textContent='Revised copied to editor.'; toast('Revised copied to editor'); });
$('#fileInput').addEventListener('change',async e=>{ const file=e.target.files[0]; if(!file) return; try{ if(isHtmlFile(file)){ await importHtmlFileOptions(file); } else { const data=await importFileObject(file); applyImportedText(data,file.name,fileDisplaySize(file.size)); } }catch(err){ toast(err.message); alert(err.message); } e.target.value=''; });
$('#importUrlBtn').addEventListener('click',()=>{ $('#urlModal').classList.add('open'); $('#urlStatus').textContent='Ready.'; setTimeout(()=>$('#urlInput').focus(),0); });
$('#urlCancelBtn').addEventListener('click',()=>$('#urlModal').classList.remove('open'));
$('#urlInput').addEventListener('keydown',e=>{ if(e.key==='Enter'){ e.preventDefault(); $('#urlFetchBtn').click(); } });
$('#urlFetchBtn').addEventListener('click',async()=>{ const url=$('#urlInput').value.trim(); if(!url){ toast('Enter a URL first.'); return; } try{ $('#urlFetchBtn').disabled=true; $('#urlStatus').textContent='Fetching URL...'; const res=await api('/api/import_url',{url}); const data=await res.json(); if(!res.ok) throw new Error(data.error||'URL import failed'); $('#urlModal').classList.remove('open'); $('#urlInput').value=''; openHtmlImportChoice(data,{displayName:data.filename||'URL import',url:data.url||url}); $('#urlStatus').textContent='Ready.'; }catch(err){ $('#urlStatus').textContent=err.message; toast(err.message); } finally { $('#urlFetchBtn').disabled=false; } });
function isMergeableFile(file){ return /\.(txt|md|markdown|rtf|docx|doc|html|htm)$/i.test(file.name); }
function renderMergeList(){ const list=$('#mergeList'); if(!list) return; list.innerHTML=''; if(!mergeFiles.length){ list.innerHTML='<div class="empty-hint">No supported files found in that folder.</div>'; return; } mergeFiles.forEach((item,idx)=>{ const row=document.createElement('div'); row.className='merge-row'; row.innerHTML=`<div><div class="merge-name">${escapeHtml(item.path)}</div><div class="merge-meta">${idx+1}. ${(item.text.length).toLocaleString()} chars · ${(item.notes||[]).join(' · ')}</div></div><div class="merge-actions"><button data-act="up" ${idx===0?'disabled':''}>Up</button><button data-act="down" ${idx===mergeFiles.length-1?'disabled':''}>Down</button></div>`; row.querySelector('[data-act="up"]').addEventListener('click',()=>{ if(idx>0){ [mergeFiles[idx-1],mergeFiles[idx]]=[mergeFiles[idx],mergeFiles[idx-1]]; renderMergeList(); } }); row.querySelector('[data-act="down"]').addEventListener('click',()=>{ if(idx<mergeFiles.length-1){ [mergeFiles[idx+1],mergeFiles[idx]]=[mergeFiles[idx],mergeFiles[idx+1]]; renderMergeList(); } }); list.appendChild(row); }); $('#mergeSub').textContent=`${mergeFiles.length} file(s) ready. Use Up/Down to choose order.`; }
$('#folderInput').addEventListener('change',async e=>{ const files=Array.from(e.target.files||[]).filter(isMergeableFile).sort((a,b)=>(a.webkitRelativePath||a.name).localeCompare(b.webkitRelativePath||b.name,undefined,{numeric:true,sensitivity:'base'})); if(!files.length){ toast('No supported files found.'); e.target.value=''; return; } mergeFiles=[]; $('#mergeStatus').textContent='Importing files...'; $('#mergeModal').classList.add('open'); try{ for(let i=0;i<files.length;i++){ $('#mergeStatus').textContent=`Importing ${i+1}/${files.length}: ${files[i].name}`; mergeFiles.push(await importFileObject(files[i])); renderMergeList(); } $('#mergeStatus').textContent='Arrange order, then merge.'; }catch(err){ $('#mergeStatus').textContent=err.message; toast(err.message); } e.target.value=''; });
$('#mergeCancelBtn').addEventListener('click',()=>{ $('#mergeModal').classList.remove('open'); mergeFiles=[]; });
$('#mergeBtn').addEventListener('click',()=>{ if(!mergeFiles.length){ toast('No files to merge.'); return; } if(editor.value && !confirm('Replace the current editor text with the merged folder?')) return; pushUndoSnapshot(); const merged=mergeFiles.map(item=>item.text.trim()).filter(Boolean).join('\n\n'); editor.value=merged; currentFilename='merged-manuscript.md'; filenameEl.textContent=`merged folder (${mergeFiles.length} files)`; $('#metaFilename').value=$('#metaFilename').value||'merged-manuscript'; undoStack.length=0; updateUndoButton(); updateStats(); $('#mergeModal').classList.remove('open'); toast('Folder merged. You can now export DOCX.'); });
$('#clearBtn').addEventListener('click',()=>{ if(editor.value && !confirm('Clear the editor?')) return; pushUndoSnapshot(); editor.value=''; currentFilename='manuscript.txt'; filenameEl.textContent='empty document'; $('#metaFilename').value=''; updateStats(); });
$('#exportTxtBtn').addEventListener('click',()=>{ const base=cleanFilename($('#metaFilename').value||currentFilename.replace(/\.[^.]+$/,'')); downloadBlob(new Blob([stripMarkers(editor.value)],{type:'text/plain;charset=utf-8'}), `${base}-clean.txt`); });
$('#exportMarkedTxtBtn').addEventListener('click',()=>{ const base=cleanFilename($('#metaFilename').value||currentFilename.replace(/\.[^.]+$/,'')); downloadBlob(new Blob([editor.value],{type:'text/markdown;charset=utf-8'}), `${base}-markdown.md`); });
$('#exportDocxBtn').addEventListener('click',async()=>{ try{ const base=cleanFilename($('#metaFilename').value||currentFilename.replace(/\.[^.]+$/,'')); const res=await api('/api/export_docx',{text:editor.value,metadata:getMetadata(),filename:base}); if(!res.ok){ const data=await res.json().catch(()=>({error:'Export failed'})); throw new Error(data.error||'Export failed'); } const blob=await res.blob(); downloadBlob(blob,`${base}-formatted.docx`); toast('DOCX exported'); }catch(err){ alert(err.message); toast(err.message); } });

$('#preflightBtn').addEventListener('click',runPreflight);
async function runPreflight(){ const res=await api('/api/preflight',{text:editor.value,metadata:getMetadata()}); const data=await res.json(); renderPreflight(data); }
function renderPreflight(data){ const s=data.summary||{}; $('#preflightSummary').innerHTML=`<span class="${s.error?'err':'ok'}">${s.error||0} errors</span> · <span class="${s.warning?'warn':'ok'}">${s.warning||0} warnings</span> · ${data.wordCount||0} words · ${data.lineCount||0} lines`; const list=$('#issueList'); list.innerHTML=''; if(!(data.issues||[]).length){ list.innerHTML='<div class="hint ok">No preflight issues found.</div>'; return; } (data.issues||[]).forEach(issue=>{ const div=document.createElement('div'); div.className='issue'; div.innerHTML=`<div class="top"><span class="dot ${issue.level}"></span><span class="msg">${escapeHtml(issue.message)} ${issue.count>1?`<b>(${issue.count})</b>`:''}</span>${issue.line?`<button class="mini">line ${issue.line}</button>`:''}</div>${issue.sample?`<div class="sample">${escapeHtml(issue.sample)}</div>`:''}`; if(issue.line){ div.querySelector('button').addEventListener('click',()=>jumpToLine(issue.line)); } list.appendChild(div); }); }
function jumpToLine(n){ const lines=editor.value.split('\n'); let pos=0; for(let i=0;i<n-1 && i<lines.length;i++) pos += lines[i].length+1; editor.focus(); editor.setSelectionRange(pos, Math.min(editor.value.length,pos+(lines[n-1]||'').length)); editor.scrollTop=Math.max(0,(n-1)*20-editor.clientHeight/2); gutter.scrollTop=editor.scrollTop; }
function escapeHtml(s){ return String(s||'').replace(/[&<>"']/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c])); }

const CP1252_SPECIALS={0x20AC:0x80,0x201A:0x82,0x0192:0x83,0x201E:0x84,0x2026:0x85,0x2020:0x86,0x2021:0x87,0x02C6:0x88,0x2030:0x89,0x0160:0x8A,0x2039:0x8B,0x0152:0x8C,0x017D:0x8E,0x2018:0x91,0x2019:0x92,0x201C:0x93,0x201D:0x94,0x2022:0x95,0x2013:0x96,0x2014:0x97,0x02DC:0x98,0x2122:0x99,0x0161:0x9A,0x203A:0x9B,0x0153:0x9C,0x017E:0x9E,0x0178:0x9F};
const CONT='\\u0080-\\u00BF\\u20AC\\u201A\\u0192\\u201E\\u2026\\u2020\\u2021\\u02C6\\u2030\\u0160\\u2039\\u0152\\u017D\\u2018\\u2019\\u201C\\u201D\\u2022\\u2013\\u2014\\u02DC\\u2122\\u0161\\u203A\\u0153\\u017E\\u0178';
const SIG_RE=new RegExp('[\\u00C2-\\u00F4]['+CONT+']'); const SIG_RE_G=new RegExp('[\\u00C2-\\u00F4]['+CONT+']','g');
function charToCp1252Byte(code){ if(code<=0xFF) return code; if(CP1252_SPECIALS[code]!==undefined) return CP1252_SPECIALS[code]; return -1; }
function isEligible(code){ return (code>=0x80&&code<=0xFF)||CP1252_SPECIALS[code]!==undefined; }
function decodeMojibakeRun(run){ const bytes=[]; for(const ch of run){ const b=charToCp1252Byte(ch.codePointAt(0)); if(b<0) return null; bytes.push(b); } const decoded=new TextDecoder('utf-8',{fatal:false}).decode(new Uint8Array(bytes)); return decoded.includes('\uFFFD')?null:decoded; }
function fixMojibake(text){ const chars=[...text]; let out='',i=0; while(i<chars.length){ if(isEligible(chars[i].codePointAt(0))){ let j=i; while(j<chars.length&&isEligible(chars[j].codePointAt(0))) j++; const run=chars.slice(i,j).join(''); out += SIG_RE.test(run) ? (decodeMojibakeRun(run) ?? run) : run; i=j; } else { out+=chars[i]; i++; } } return out; }
function isSceneBreak(line){ const t=line.trim(); return /^(?:\*\s*){3,}$/.test(t)||/^(?:[-–—]\s*){3,}$/.test(t)||/^#{1,3}$/.test(t)||/^~{3,}$/.test(t)||/^•{3,}$/.test(t); }
function isListOrBullet(line){ return /^\s*(?:[-*•]\s+|\d{1,3}[.)]\s+|[a-zA-Z][.)]\s+)/.test(line); }
function isLikelySectionTitle(line){ const t=line.trim(); if(!t||t.length>90) return false; if(isSceneBreak(t)||isListOrBullet(t)) return true; if(/^(?:chapter|hoofdstuk|part|deel|book|prologue|epilogue|appendix|contents|inhoud|acknowledg|copyright)\b/i.test(t)) return true; if(/^(?:[ivxlcdm]{1,8}|[IVXLCDM]{1,8}|\d{1,4})$/.test(t)) return true; if(t===t.toUpperCase()&&/[A-ZÀ-Ý]/.test(t)&&t.length<=70) return true; return false; }
function endsLikeSentence(line){ return /[.!?:;…]["'”’)]*$/.test(line.trim()); }
function shouldJoinLines(prev,current){ const a=prev.trim(),b=current.trim(); if(!a||!b) return false; if(isSceneBreak(a)||isSceneBreak(b)||isLikelySectionTitle(a)||isLikelySectionTitle(b)||isListOrBullet(current)||/^\s/.test(current)) return false; if(a.length<45&&!/[A-Za-zÀ-ÖØ-öø-ÿ]-$/.test(a)) return false; return !endsLikeSentence(a); }
function isLikelyPageNumberLine(line,idx,lines){ const t=line.trim(); if(!t) return false; if(/^[-–—]?\s*\d{1,4}\s*[-–—]?$/.test(t)) return true; if(/^(?:p\.?\s*|page\s+|pagina\s+)\d{1,4}(?:\s*(?:of|van|\/|-)\s*\d{1,4})?$/i.test(t)) return true; if(/^\d{1,4}\s*(?:\/|of|van)\s*\d{1,4}$/i.test(t)) return true; if(/^[ivxlcdm]{2,8}$/i.test(t)){ const prevBlank=idx===0||lines[idx-1].trim()===''; const nextBlank=idx===lines.length-1||lines[idx+1].trim()===''; return prevBlank&&nextBlank; } return false; }
function repeatedLineCandidates(lines){ const map=new Map(); lines.forEach((line,idx)=>{ const t=line.trim(); if(t.length<4||t.length>90||isSceneBreak(t)||isLikelyPageNumberLine(t,idx,lines)) return; const key=t.replace(/\s+/g,' '); if(!map.has(key)) map.set(key,[]); map.get(key).push(idx); }); const out=new Set(); for(const positions of map.values()){ if(positions.length<4) continue; if(positions[positions.length-1]-positions[0]<40) continue; positions.forEach(i=>out.add(i)); } return out; }
function fixInvisible(text){ return text.replace(/[\u200B\u200C\u200D\u2060\uFEFF\u00AD]/g,''); }
function fixSpecialSpaces(text){ return text.replace(/[\u00A0\u1680\u2000-\u200A\u202F\u205F\u3000]/g,' '); }
function fixControlChars(text){ return text.replace(/[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F]/g,''); }
function fixUnicodeNFC(text){ return text.normalize('NFC'); }
function fixLigatures(text){ return text.replace(/\uFB00/g,'ff').replace(/\uFB01/g,'fi').replace(/\uFB02/g,'fl').replace(/\uFB03/g,'ffi').replace(/\uFB04/g,'ffl').replace(/\uFB05/g,'st').replace(/\uFB06/g,'st'); }
function fixJoin(text,opts){ const lines=text.split('\n'),out=[]; let buffer=null; for(const raw of lines){ if(raw.trim()===''){ if(buffer!==null){ out.push(buffer); buffer=null; } out.push(''); continue; } if(buffer===null){ buffer=raw.trim(); continue; } if(shouldJoinLines(buffer,raw)){ buffer=(opts.dehyphenate&&/[A-Za-zÀ-ÖØ-öø-ÿ]-$/.test(buffer)) ? buffer.replace(/-$/,'')+raw.trim() : buffer+' '+raw.trim(); } else { out.push(buffer); buffer=raw.trim(); } } if(buffer!==null) out.push(buffer); return out.join('\n'); }
function fixBlankRuns(text){ return text.replace(/\n[ \t]*\n(?:[ \t]*\n)+/g,'\n\n'); }
function fixTrailingSpaces(text){ return text.split('\n').map(l=>l.replace(/[ \t]+$/,'')).join('\n'); }
function fixIndentation(text){ return text.split('\n').map(l=>l.replace(/^[ \t]+/,'')).join('\n'); }
function fixPageNumbers(text){ const lines=text.split('\n'); return lines.filter((l,idx)=>!isLikelyPageNumberLine(l,idx,lines)).join('\n'); }
function fixRepeatedLines(text){ const lines=text.split('\n'); const remove=repeatedLineCandidates(lines); return lines.filter((_,idx)=>!remove.has(idx)).join('\n'); }
function fixTabsToSpace(text){ return text.replace(/\t/g,' '); }
function fixCollapseSpaces(text){ return text.split('\n').map(l=>l.replace(/(?<=\S) {2,}/g,' ')).join('\n'); }
function fixSpaceBeforePunct(text){ return text.replace(/ +([,.;:!?])/g,'$1'); }
function fixSpaceAfterPunct(text){ return text.replace(/([a-zà-ÿ]{2})([.!?])([A-ZÀ-Ý])/g,'$1$2 $3').replace(/([a-zà-ÿ]),([a-zà-ÿA-ZÀ-Ý])/g,'$1, $2'); }
function fixQuotesCurly(text){ return text.replace(/(^|[\s([{—–-])'(?=(?:em|cause|til|tis|twas|n|round|bout|fore|gain|neath|cross)\b)/gi,'$1’').replace(/(^|[\s([{—–-])'(?=\d{2}s\b)/g,'$1’').replace(/(^|[\s([{—–-])"/g,'$1“').replace(/"/g,'”').replace(/(^|[\s([{—–-])'/g,'$1‘').replace(/'/g,'’'); }
function fixQuotesStraight(text){ return text.replace(/[“”]/g,'"').replace(/[‘’]/g,"'"); }
function fixDashEm(text){ return text.replace(/(^|[^-])---(?!-)/g,'$1—').replace(/(^|[^-])--(?!-)/g,'$1—'); }
function fixEllipsis(text){ return text.replace(/\.\s+\.\s+\./g,'…').replace(/\.{3}/g,'…'); }
function fixHtmlText(text){ if(!/<[a-z!/][\s\S]*>/i.test(text) && !/&(?:[a-z][a-z0-9]+|#\d+|#x[0-9a-f]+);/i.test(text)) return text; const doc=new DOMParser().parseFromString(text,'text/html'); doc.querySelectorAll('script,style,noscript,template,svg,canvas,nav,header,footer,aside,form,menu,.sidebar,.navigation,.menu,.footer,.header,#sidebar,#navigation,#menu,#footer,#header').forEach(el=>el.remove()); const selectors=['main','article','[role="main"]','#content','#main','#article','.content','.post','.entry']; let root=null; selectors.forEach(sel=>{ doc.querySelectorAll(sel).forEach(el=>{ if(!root || (el.textContent||'').length>(root.textContent||'').length) root=el; }); }); root=root||doc.body; const blocks=new Set(['address','article','blockquote','body','dd','details','dialog','div','dl','dt','fieldset','figcaption','figure','h1','h2','h3','h4','h5','h6','hr','li','main','ol','p','pre','section','table','tbody','td','tfoot','th','thead','tr','ul']); const skip=new Set(['script','style','noscript','template','svg','canvas','nav','header','footer','aside','form','menu']); const walk=node=>{ if(node.nodeType===Node.TEXT_NODE) return node.nodeValue||''; if(node.nodeType!==Node.ELEMENT_NODE && node.nodeType!==Node.DOCUMENT_FRAGMENT_NODE && node.nodeType!==Node.DOCUMENT_NODE) return ''; const tag=(node.tagName||'').toLowerCase(); if(skip.has(tag)) return ''; if(tag==='br') return '\n'; let inner=''; node.childNodes.forEach(child=>{ inner += walk(child); }); if(tag==='td'||tag==='th') return `\n${inner}\t`; if(blocks.has(tag)) return `\n${inner}\n`; return inner; }; return walk(root).replace(/\u00a0/g,' ').replace(/[ \t\f\v]+/g,' ').replace(/ *\n */g,'\n').replace(/\n{3,}/g,'\n\n').trim(); }
function fixTrimDoc(text){ return text.trim(); }
const FIXES={mojibake:fixMojibake,invisible:fixInvisible,specialSpaces:fixSpecialSpaces,controlChars:fixControlChars,unicodeNFC:fixUnicodeNFC,ligatures:fixLigatures,htmlText:fixHtmlText,join:fixJoin,blankRuns:fixBlankRuns,trailingSpaces:fixTrailingSpaces,indentation:fixIndentation,pageNumbers:fixPageNumbers,repeatLines:fixRepeatedLines,tabsToSpace:fixTabsToSpace,collapseSpaces:fixCollapseSpaces,spaceBeforePunct:fixSpaceBeforePunct,spaceAfterPunct:fixSpaceAfterPunct,quotesCurly:fixQuotesCurly,quotesStraight:fixQuotesStraight,dashEm:fixDashEm,ellipsis:fixEllipsis,trimDoc:fixTrimDoc};
function safeCleanup(text){ let out=text.replace(/\r\n/g,'\n').replace(/\r/g,'\n'); for(const id of SAFE_FIX_SEQUENCE){ out=FIXES[id](out,{dehyphenate:$('#dehyphenate')?.checked??true}); } return out; }
function buildFixResult(id){ const fn=FIXES[id]; if(!fn) return null; const prev=editor.value; const opts={dehyphenate:$('#dehyphenate')?.checked??true}; if(getScope()==='selection'){ const start=editor.selectionStart,end=editor.selectionEnd; if(start===end){ toast('Select some text first, or switch to Whole text.'); return null; } const before=prev.slice(start,end); const after=fn(before,opts); return {id,before,after,start,end,title:FIX_LABELS[id]||id}; } const fixed=fn(prev,opts); return {id,before:prev,after:fixed,start:0,end:prev.length,title:FIX_LABELS[id]||id}; }
function applyTextChange(before,after,start,end,label){ if(after===before){ toast('Nothing to change.'); return; } replaceRange(start,end,after); toast(label||'Applied'); }
function showFixPreview(result){ openPreview(`Preview: ${result.title}`,result.before,result.after,result.start,result.end,false); $('#applyPreviewBtn').disabled=false; previewState.finished=true; }
document.querySelectorAll('[data-fix]').forEach(btn=>{ if(DESTRUCTIVE_FIXES.has(btn.dataset.fix)) btn.classList.add('destructive'); btn.addEventListener('click',()=>{ const result=buildFixResult(btn.dataset.fix); if(!result) return; if(result.after===result.before){ toast('Nothing to change for this operation.'); return; } if(DESTRUCTIVE_FIXES.has(btn.dataset.fix)) showFixPreview(result); else applyTextChange(result.before,result.after,result.start,result.end,FIX_LABELS[btn.dataset.fix]); }); });
$('#safeCleanupBtn').addEventListener('click',()=>{ if(getScope()==='selection'){ const start=editor.selectionStart,end=editor.selectionEnd; if(start===end){ toast('Select some text first, or switch to Whole text.'); return; } const before=editor.value.slice(start,end); const after=safeCleanup(before); applyTextChange(before,after,start,end,'Safe Cleanup applied to selection'); return; } const before=editor.value; const after=safeCleanup(before); applyTextChange(before,after,0,before.length,'Safe Cleanup applied'); });
$('#showMarkersBtn').addEventListener('click',()=>openPreview('Clean text preview', editor.value, stripMarkers(editor.value), 0, editor.value.length, false));

function markLine(offs,count,pos,flags,flagName){ let lo=0,hi=count-1,li=0; while(lo<=hi){ const mid=(lo+hi)>>1; if(offs[mid]<=pos){ li=mid; lo=mid+1; } else hi=mid-1; } if(flags[li]) flags[li].add(flagName); }
function diagnose(text){ const lines=text.split('\n'); const flags=lines.map(()=>new Set()); const matches={wrap:[],trailingWs:[],tabs:[],doubleSpace:[],spaceBeforePunct:[],extraBlank:[],doubleHyphen:[],tripleDot:[],mojibake:[],specialSpace:[],invisible:[],control:[],pageNum:[],repeatLine:[],ligature:[]}; const offs=[]; let acc=0; for(let i=0;i<lines.length;i++){ offs.push(acc); acc += lines[i].length+1; } const repeatCandidates=repeatedLineCandidates(lines); for(let i=0;i<lines.length;i++){ const line=lines[i], lineStart=offs[i], lineEnd=lineStart+line.length; if(line.trim()==='') continue; if(i>0&&shouldJoinLines(lines[i-1],line)){ flags[i].add('wrap'); matches.wrap.push({start:lineStart,end:lineEnd}); } const tw=line.match(/[ \t]+$/); if(tw){ flags[i].add('trailingWs'); matches.trailingWs.push({start:lineStart+tw.index,end:lineEnd}); } let idx=line.indexOf('\t'); while(idx!==-1){ flags[i].add('tabs'); matches.tabs.push({start:lineStart+idx,end:lineStart+idx+1}); idx=line.indexOf('\t',idx+1); } let m; const dsRe=/(?<=\S) {2,}/g; while((m=dsRe.exec(line))){ flags[i].add('doubleSpace'); matches.doubleSpace.push({start:lineStart+m.index,end:lineStart+m.index+m[0].length}); } const sbRe=/ +[,.;:!?]/g; while((m=sbRe.exec(line))){ flags[i].add('spaceBeforePunct'); matches.spaceBeforePunct.push({start:lineStart+m.index,end:lineStart+m.index+m[0].length}); } if(isLikelyPageNumberLine(line,i,lines)){ flags[i].add('pageNum'); matches.pageNum.push({start:lineStart,end:lineEnd}); } if(repeatCandidates.has(i)){ flags[i].add('repeatLine'); matches.repeatLine.push({start:lineStart,end:lineEnd}); } } let i=0; while(i<lines.length){ if(lines[i].trim()===''){ let j=i; while(j<lines.length&&lines[j].trim()==='') j++; if(j-i>=2){ for(let k=i+1;k<j;k++){ flags[k].add('extraBlank'); matches.extraBlank.push({start:offs[k],end:offs[k]}); } } i=j; } else i++; } let m; const dhRe=/(^|[^-])--(?!-)/g; while((m=dhRe.exec(text))) matches.doubleHyphen.push({start:m.index+(m[1]?m[1].length:0),end:m.index+m[0].length}); const tdRe=/\.{3}/g; while((m=tdRe.exec(text))) matches.tripleDot.push({start:m.index,end:m.index+m[0].length}); const ligRe=/[\uFB00-\uFB06]/g; while((m=ligRe.exec(text))){ matches.ligature.push({start:m.index,end:m.index+1}); markLine(offs,lines.length,m.index,flags,'ligature'); } SIG_RE_G.lastIndex=0; while((m=SIG_RE_G.exec(text))){ matches.mojibake.push({start:m.index,end:m.index+m[0].length}); markLine(offs,lines.length,m.index,flags,'mojibake'); } const ssRe=/[\u00A0\u1680\u2000-\u200A\u202F\u205F\u3000]/g; while((m=ssRe.exec(text))){ matches.specialSpace.push({start:m.index,end:m.index+1}); markLine(offs,lines.length,m.index,flags,'specialSpace'); } const invRe=/[\u200B\u200C\u200D\u2060\uFEFF\u00AD]/g; while((m=invRe.exec(text))){ matches.invisible.push({start:m.index,end:m.index+1}); markLine(offs,lines.length,m.index,flags,'invisible'); } const ctlRe=/[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F]/g; while((m=ctlRe.exec(text))){ matches.control.push({start:m.index,end:m.index+1}); markLine(offs,lines.length,m.index,flags,'control'); } return {lines,flags,matches,straightQuotes:(text.match(/["']/g)||[]).length,curlyQuotes:(text.match(/[“”‘’]/g)||[]).length}; }
const DIAG_DEFS=[['mojibake','Mojibake / encoding damage','violet'],['invisible','Zero-width / invisible characters','violet'],['specialSpace','Non-breaking / special spaces','violet'],['control','Control / page-break characters','violet'],['ligature','Typographic ligatures','violet'],['wrap','Likely hard-wrapped lines','amber'],['pageNum','Standalone page-number lines','amber'],['repeatLine','Repeated header/footer candidates','amber'],['extraBlank','Extra blank lines','rust'],['trailingWs','Trailing whitespace','rust'],['tabs','Tabs in the text','rust'],['doubleSpace','Double spaces','blue'],['spaceBeforePunct','Space before punctuation','blue'],['doubleHyphen','Exact -- instead of em dash','blue'],['tripleDot','Three dots instead of ellipsis','blue']];
function renderDiagnose(diag){ const list=$('#diagList'); if(!list) return; list.innerHTML=''; DIAG_DEFS.forEach(([key,label,dot])=>{ const arr=diag.matches[key]||[]; const row=document.createElement('div'); row.className='diag-row'+(arr.length?'':' zero'); row.innerHTML=`<span class="dot ${dot}"></span><span class="lbl">${label}</span><span class="count">${arr.length}</span><button class="jump">next</button>`; row.querySelector('.jump').addEventListener('click',()=>jumpToDiag(key)); list.appendChild(row); }); const mixed=diag.straightQuotes>0&&diag.curlyQuotes>0; const q=document.createElement('div'); q.className='diag-row'+(mixed?'':' zero'); q.innerHTML=`<span class="dot ${mixed?'blue':'grey'}"></span><span class="lbl">Quotes: ${diag.straightQuotes} straight, ${diag.curlyQuotes} curly${mixed?' - mixed':''}</span>`; list.appendChild(q); lastDiag=diag; }
function jumpToDiag(key){ const arr=lastDiag?.matches?.[key]||[]; if(!arr.length) return; jumpPointers[key]=((jumpPointers[key]??-1)+1)%arr.length; const m=arr[jumpPointers[key]]; const start=Math.max(0,Math.min(m.start,editor.value.length)); const end=Math.max(start,Math.min(m.end,editor.value.length)); const lineNo=editor.value.slice(0,start).split('\n').length; editor.focus(); editor.setSelectionRange(start,end>start?end:Math.min(editor.value.length,start+1)); editor.scrollTop=Math.max(0,(lineNo-1)*20-editor.clientHeight/2); gutter.scrollTop=editor.scrollTop; }
function renderVellumStructure(){ const summary=$('#vellumSummary'), list=$('#vellumList'); if(!summary||!list) return; const lines=editor.value.split('\n'); const rows=[]; let chapters=0,subheads=0,special=0,candidates=0,blockMarker=''; lines.forEach((line,idx)=>{ const container=line.match(/^\s*:::\s*([a-z0-9_-]+)?\s*$/i); if(container){ blockMarker=canonicalStructureMarker(container[1]||''); return; } const md=markdownLineInfo(line); const clean=md ? md.content : line.trim(); let marker=md?.marker || ''; if(!marker && blockMarker && clean) marker=blockMarker; if(marker){ const info=MD_STYLE_INFO[marker]||[marker.toUpperCase(),'supp']; const text=clean||'(empty marker)'; rows.push({line:idx+1,marker,text,tag:info[0],group:info[1]}); if(info[1]==='chapter') chapters++; else if(info[1]==='sub') subheads++; else special++; return; } if(clean && VELLUM_CHAPTER_CANDIDATE_RE.test(clean)){ candidates++; rows.push({line:idx+1,marker:'candidate',text:clean,tag:'CHECK',group:'warn'}); } }); summary.innerHTML=`<div class="vsum"><b>${chapters}</b> chapters/parts</div><div class="vsum"><b>${subheads}</b> subheads</div><div class="vsum"><b>${special}</b> special</div><div class="vsum"><b>${candidates}</b> candidates</div>`; list.innerHTML=''; if(!rows.length){ list.innerHTML='<div class="empty-hint">No Markdown structure yet. Possible chapter lines will show here as CHECK.</div>'; return; } rows.slice(0,80).forEach(r=>{ const div=document.createElement('div'); div.className='vellum-row'; const tagClass=r.group==='sub'?'sub':r.group==='supp'?'supp':r.group==='warn'?'warn':''; div.innerHTML=`<span class="tag ${tagClass}">${escapeHtml(r.tag)}</span><span class="vtext">${escapeHtml(r.text)}</span><button class="jump">line ${r.line}</button>`; div.querySelector('button').addEventListener('click',()=>jumpToLine(r.line)); list.appendChild(div); }); if(rows.length>80){ const more=document.createElement('div'); more.className='empty-hint'; more.textContent=`Showing first 80 of ${rows.length} structure items/candidates.`; list.appendChild(more); } }

function mergeDefaultPrompts(prompts){ const list=Array.isArray(prompts)?prompts.slice():[]; const names=new Set(list.map(p=>p?.name).filter(Boolean)); DEFAULT_PROMPTS.forEach(p=>{ if(!names.has(p.name)){ list.push({...p}); names.add(p.name); } }); return list; }
function loadPrompts(){ let prompts; try{ prompts=JSON.parse(localStorage.getItem('vp.py.prompts')||'null'); }catch(e){} if(!Array.isArray(prompts)||!prompts.length) prompts=DEFAULT_PROMPTS; else prompts=mergeDefaultPrompts(prompts); const sel=$('#promptSelect'); sel.innerHTML=''; prompts.forEach((p,i)=>{ const opt=document.createElement('option'); opt.value=i; opt.textContent=p.name; sel.appendChild(opt); }); sel._prompts=prompts; try{ localStorage.setItem('vp.py.prompts', JSON.stringify(prompts)); }catch(e){} sel.value=localStorage.getItem('vp.py.promptIndex')||0; if(!prompts[sel.value]) sel.value=0; $('#promptText').value=prompts[sel.value].text; }
function savePrompts(){ localStorage.setItem('vp.py.prompts', JSON.stringify($('#promptSelect')._prompts||DEFAULT_PROMPTS)); localStorage.setItem('vp.py.promptIndex', $('#promptSelect').value||0); }
$('#promptSelect').addEventListener('change',()=>{ const p=$('#promptSelect')._prompts[$('#promptSelect').value]; $('#promptText').value=p?p.text:''; savePrompts(); });
$('#savePromptBtn').addEventListener('click',()=>{ const prompts=$('#promptSelect')._prompts; prompts[$('#promptSelect').value].text=$('#promptText').value; savePrompts(); toast('Prompt saved'); });
$('#savePromptAsBtn').addEventListener('click',()=>{ const name=prompt('Prompt name?'); if(!name) return; const prompts=$('#promptSelect')._prompts; prompts.push({name,text:$('#promptText').value}); savePrompts(); loadPrompts(); $('#promptSelect').value=prompts.length-1; savePrompts(); toast('Prompt added'); });
$('#deletePromptBtn').addEventListener('click',()=>{ const prompts=$('#promptSelect')._prompts; if(prompts.length<=1){ toast('Keep at least one prompt.'); return; } prompts.splice(Number($('#promptSelect').value),1); savePrompts(); loadPrompts(); });
$('#resetPromptsBtn').addEventListener('click',()=>{ if(!confirm('Reset prompts to defaults?')) return; localStorage.setItem('vp.py.prompts', JSON.stringify(DEFAULT_PROMPTS)); localStorage.setItem('vp.py.promptIndex','0'); loadPrompts(); });

$('#detectModelBtn').addEventListener('click',detectModels);
async function detectModels(){ const provider=currentAiProvider(); try{ $('#lmStatus').textContent=`Detecting ${provider.label} models...`; const res=await api('/api/lm_models',{base_url:$('#lmBase').value,provider:$('#lmProvider').value,api_key:$('#lmApiKey').value}); const data=await res.json(); if(!res.ok) throw new Error(data.error||'Detect failed'); const sel=$('#lmModel'); const previous=sel.value; resetModelSelect(); (data.models||[]).forEach(id=>{ const opt=document.createElement('option'); opt.value=id; opt.textContent=id; sel.appendChild(opt); }); if((data.models||[]).length){ sel.value=(previous && (data.models||[]).includes(previous)) ? previous : data.models[0]; $('#lmBase').value=data.base_url||$('#lmBase').value; $('#lmStatus').textContent=`${provider.label}: detected ${(data.models||[]).length} model(s).`; saveSettings(); } else $('#lmStatus').textContent=`${provider.label}: connection OK, but no models returned.`; }catch(err){ $('#lmStatus').textContent=err.message; toast(`${provider.label} detect failed`); } }
$('#lmTestBtn').addEventListener('click',async()=>{ const provider=currentAiProvider(); try{ $('#lmStatus').textContent=`Testing ${provider.label} connection...`; const res=await api('/api/lm_models',{base_url:$('#lmBase').value,provider:$('#lmProvider').value,api_key:$('#lmApiKey').value}); const data=await res.json(); if(!res.ok) throw new Error(data.error||'Test failed'); $('#lmStatus').textContent=(data.models||[]).length?`${provider.label} connection OK. ${(data.models||[]).length} model(s) available.`:`${provider.label} connection OK, but no models returned.`; toast(`${provider.label} connection OK`); }catch(err){ $('#lmStatus').textContent=err.message; toast(`${provider.label} test failed`); } });
function getScopeRange(scope){ const val=editor.value; if(scope==='selection'){ const s=editor.selectionStart,e=editor.selectionEnd; if(s===e) throw new Error('Select text first.'); return {start:s,end:e,text:val.slice(s,e)}; } if(scope==='paragraph'){ const r=currentParagraphRange(); return {start:r.start,end:r.end,text:val.slice(r.start,r.end)}; } return {start:0,end:val.length,text:val}; }
function cleanPrompt(){ const prompts=$('#promptSelect')._prompts||DEFAULT_PROMPTS; const prompt=prompts.find(p=>p.name==='AI Cleanup — conservative')||DEFAULT_PROMPTS.find(p=>p.name==='AI Cleanup — conservative'); return prompt.text; }
function promptForTask(task){ return task==='clean' ? cleanPrompt() : $('#promptText').value; }
function prepareTextForLm(text){ return {text, preserveStructure:structureSignature(text)!==''}; }
function structureSignature(text){ return text.split('\n').map(line=>{ const container=line.match(/^\s*:::\s*([a-z0-9_-]+)?\s*$/i); if(container) return `block:${canonicalStructureMarker(container[1]||'end')}`; const md=markdownLineInfo(line); return md ? md.marker : ''; }).filter(Boolean).join('|'); }
function markerSignature(text){ return structureSignature(text); }
function markerCount(text){ return structureSignature(text).split('|').filter(Boolean).length; }
function reportPreviewMarkerDrift(){ if(previewState.expectedMarkers===null) return; const beforeCount=markerCount(previewState.before), afterCount=markerCount(previewState.after); const afterSig=markerSignature(previewState.after); if(previewState.expectedMarkers!==afterSig){ $('#previewStatus').textContent=beforeCount===afterCount ? `Structure changed position/order, but count stayed ${afterCount}. Review before applying.` : `Structure count changed: before ${beforeCount}, after ${afterCount}. Review before applying.`; $('#lmStatus').textContent=$('#previewStatus').textContent; } }
const NAME_IGNORE=new Set('A An And As At But By Chapter De Deel Dr For From He Her His I If In It Its Me Mr Mrs Ms My No Not Of On Or Our She Sir So The Their They This To We With You Your Een Het De Een En Van Voor Maar Als Dat Die Dit Deze Naar Zijn Haar Hij Zij Wij Jij U Ze There Then When Where'.split(/\s+/));
function plainForIntegrity(text){ return stripInlineMarkup(String(text||'').split('\n').map(stripStructureLine).join('\n')).replace(/[`*_~>#|[\]()]/g,' '); }
function paragraphCount(text){ return String(text||'').split(/\n\s*\n/).map(p=>p.trim()).filter(Boolean).length; }
function dialogueCount(text){ return String(text||'').split('\n').filter(line=>line.trim().length>1 && (/^\s*(?:["“‘'’-]|—)/.test(line) || /[“”]/.test(line))).length; }
function tokenSet(text, re){ const set=new Set(); let m; re.lastIndex=0; while((m=re.exec(text))) set.add(m[0]); return set; }
function properNouns(text){ const clean=plainForIntegrity(text); const set=new Set(); const re=/\b[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿ'’-]{2,}(?:\s+[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿ'’-]{2,}){0,2}\b/g; let m; while((m=re.exec(clean))){ const name=m[0].trim(); const first=name.split(/\s+/)[0]; if(!NAME_IGNORE.has(first) && !NAME_IGNORE.has(name)) set.add(name); } return set; }
function setDiff(a,b){ return [...a].filter(x=>!b.has(x)).sort((x,y)=>x.localeCompare(y)); }
function analyzeIntegrity(before,after){ const bPlain=plainForIntegrity(before), aPlain=plainForIntegrity(after); const bWords=wordCount(bPlain), aWords=wordCount(aPlain); const ratio=bWords? aWords/bWords : 1; const bPara=paragraphCount(before), aPara=paragraphCount(after); const bDialog=dialogueCount(before), aDialog=dialogueCount(after); const bNames=properNouns(before), aNames=properNouns(after); const missingNames=setDiff(bNames,aNames), addedNames=setDiff(aNames,bNames); const numberRe=/\b(?:\d+(?:[.,:/-]\d+)*|Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?|januari|februari|maart|april|mei|juni|juli|augustus|september|oktober|november|december)\b/gi; const bNums=tokenSet(bPlain,numberRe), aNums=tokenSet(aPlain,numberRe); const missingNums=setDiff(bNums,aNums), addedNums=setDiff(aNums,bNums); const issues=[]; const facts=[]; facts.push(`${bWords.toLocaleString()} → ${aWords.toLocaleString()} words`); facts.push(`${bPara} → ${aPara} paragraphs`); if(ratio<.7||ratio>1.3) issues.push(`length changed ${Math.round((ratio-1)*100)}%`); if(bPara && Math.abs(aPara-bPara)>Math.max(2,Math.ceil(bPara*.25))) issues.push(`paragraph count changed ${bPara} → ${aPara}`); if(bDialog && Math.abs(aDialog-bDialog)>Math.max(2,Math.ceil(bDialog*.25))) issues.push(`dialogue lines changed ${bDialog} → ${aDialog}`); if(missingNames.length) issues.push(`names missing: ${missingNames.slice(0,6).join(', ')}${missingNames.length>6?'...':''}`); if(addedNames.length>=3) issues.push(`new names: ${addedNames.slice(0,6).join(', ')}${addedNames.length>6?'...':''}`); if(missingNums.length) issues.push(`numbers/dates missing: ${missingNums.slice(0,8).join(', ')}${missingNums.length>8?'...':''}`); if(addedNums.length>=3) issues.push(`new numbers/dates: ${addedNums.slice(0,8).join(', ')}${addedNums.length>8?'...':''}`); return {issues,facts}; }
function renderIntegrityGuardrails(before,after){ const panel=$('#integrityPanel'); if(!panel) return; const result=analyzeIntegrity(before,after); panel.className='integrity-panel show '+(result.issues.length?'warn':'ok'); const title=result.issues.length?`Content guardrails: ${result.issues.length} item(s) to review`:'Content guardrails: no major drift detected'; const pills=[...result.issues.map(x=>`<span class="pill issue">${escapeHtml(x)}</span>`),...result.facts.map(x=>`<span class="pill">${escapeHtml(x)}</span>`)]; panel.innerHTML=`<div>${escapeHtml(title)}</div><div class="integrity-items">${pills.join('')}</div>`; }
function clearIntegrityGuardrails(){ const panel=$('#integrityPanel'); if(panel){ panel.className='integrity-panel'; panel.innerHTML=''; } }
function sanitizeLmFinalText(text){ let out=String(text||'').replace(/<think>[\s\S]*?<\/think>/gi,'').trim(); const fence=out.match(/^```(?:[a-z0-9_-]+)?\s*\n([\s\S]*?)\n```$/i); if(fence) out=fence[1].trim(); out=out.replace(/^(?:here is|here's|hier is|dit is)\b[^\n:]*:\s*\n+/i,''); return out; }
function buildLmUserContent(text, prep, task){ const isClean=task==='clean'; const strength=$('#lmStrength')?.value||'medium'; const strengthLine=!isClean && strength!=='default' ? `Rewrite strength for this task: ${strength}.\n` : ''; const markerLine=prep.preserveStructure ? 'Important: Preserve Markdown structure exactly, including heading markers (#, ##), blockquote markers (>), scene breaks (***), and fenced block lines (:::) .\n' : ''; const encodingLine='Use normal UTF-8 Unicode characters. Do not output mojibake sequences such as Ã©, â€™, â€œ, or â€”.\n'; const actionLine=isClean ? 'Clean the following text conservatively. Preserve wording and style unless something is clearly broken. Output only the cleaned text.' : 'Rewrite the following text. Output only the processed text.'; return `${strengthLine}${markerLine}${encodingLine}\n${actionLine}\n\n<<<BEGIN TEXT>>>\n${text}\n<<<END TEXT>>>`; }
function optionalNumber(id){ const raw=$('#'+id)?.value.trim(); if(!raw) return null; const n=Number(raw.replace(',','.')); return Number.isFinite(n)?n:null; }
function optionalInteger(id){ const raw=$('#'+id)?.value.trim(); if(!raw) return null; const n=parseInt(raw,10); return Number.isFinite(n)?n:null; }
function advancedOptions(){ const opts={}; const temperature=optionalNumber('advTemperature'), top_p=optionalNumber('advTopP'), top_k=optionalInteger('advTopK'), repeat_penalty=optionalNumber('advRepeatPenalty'), max_tokens=optionalInteger('advMaxTokens'); if(temperature!==null) opts.temperature=temperature; if(top_p!==null) opts.top_p=top_p; if(top_k!==null) opts.top_k=top_k; if(repeat_penalty!==null) opts.repeat_penalty=repeat_penalty; if(max_tokens!==null) opts.max_tokens=max_tokens; return opts; }
function guidedSystemPrompt(kind){ if(kind==='questions') return `You are a careful writing coach. Review the supplied manuscript text and ask only the clarification questions needed before improving it. Do not rewrite the text. If no clarification is needed, say that and list the main improvement priorities. Preserve privacy and do not invent missing context.`; if(kind==='feedback') return `You are a manuscript editor. Give concrete, useful feedback on clarity, flow, tone, structure, consistency, and language. Do not rewrite the full text. Use concise numbered points with examples where helpful. Do not invent facts.`; if(kind==='translate') return `You are a careful literary translator. Translate the supplied text into the requested target language while preserving meaning, paragraph order, Markdown structure, tone, names, numbers, dialogue intent, and scene order. Output only the translated text.`; return `You are a manuscript editor. Rewrite the supplied text using the user's guidance while preserving meaning, facts, names, numbers, paragraph order, Markdown structure, POV, tense, tone, and scene order. Output only the rewritten text.`; }
function guidedUserContent(kind,text){ const guidance=$('#guidedContext')?.value.trim()||''; const target=$('#translateTarget')?.value.trim()||'English'; if(kind==='questions') return `Context / goal:\n${guidance||'(none provided)'}\n\nAsk clarification questions for this text. Do not rewrite it yet.\n\n<<<BEGIN TEXT>>>\n${text}\n<<<END TEXT>>>`; if(kind==='feedback') return `Context / goal:\n${guidance||'(none provided)'}\n\nGive feedback only. Do not rewrite the full text.\n\n<<<BEGIN TEXT>>>\n${text}\n<<<END TEXT>>>`; if(kind==='translate') return `Target language: ${target}\nGuidance:\n${guidance||'(none provided)'}\n\nTranslate this text. Preserve Markdown structure exactly.\n\n<<<BEGIN TEXT>>>\n${text}\n<<<END TEXT>>>`; return `Guidance / answers:\n${guidance||'(none provided)'}\n\nRewrite this text using the guidance. Preserve Markdown structure exactly.\n\n<<<BEGIN TEXT>>>\n${text}\n<<<END TEXT>>>`; }
function chunkScope(){ return selectedTextLength() ? 'selection' : 'whole'; }
$('#rewriteSelBtn').addEventListener('click',()=>startLM('rewrite','selection',false)); $('#rewriteParaBtn').addEventListener('click',()=>startLM('rewrite','paragraph',false)); $('#rewriteWholeBtn').addEventListener('click',()=>startLM('rewrite','whole',false)); $('#rewriteChunkBtn').addEventListener('click',()=>startLM('rewrite',chunkScope(),true)); $('#cleanSelBtn').addEventListener('click',()=>startLM('clean','selection',false)); $('#cleanParaBtn').addEventListener('click',()=>startLM('clean','paragraph',false)); $('#cleanWholeBtn').addEventListener('click',()=>startLM('clean','whole',false)); $('#cleanChunkBtn').addEventListener('click',()=>startLM('clean',chunkScope(),true));
$('#askImproveBtn').addEventListener('click',()=>startGuidedAI('questions','selection'));
$('#feedbackOnlyBtn').addEventListener('click',()=>startGuidedAI('feedback','selection'));
$('#rewriteGuidedBtn').addEventListener('click',()=>startGuidedAI('guided-rewrite','selection'));
$('#translateSelBtn').addEventListener('click',()=>startGuidedAI('translate','selection'));
function setLmBusy(on){ ['rewriteSelBtn','rewriteParaBtn','rewriteWholeBtn','rewriteChunkBtn','cleanSelBtn','cleanParaBtn','cleanWholeBtn','cleanChunkBtn','askImproveBtn','feedbackOnlyBtn','rewriteGuidedBtn','translateSelBtn'].forEach(id=>{$('#'+id).disabled=on;}); }
async function startLM(task, scope, chunked){ if(previewState.running){ toast('Generation is already running. Stop it first.'); return; } try{ const range=getScopeRange(scope); const promptText=promptForTask(task); if(!range.text.trim()) throw new Error('No text to send.'); if(chunked){ await startLMChunked(task, scope, range, promptText); return; } const prep=prepareTextForLm(range.text); openPreview(`${task==='clean'?'AI Clean':'Rewrite'} ${scope}`, range.text, '', range.start, range.end, true); previewState.expectedMarkers=markerSignature(range.text); await streamOne(promptText, buildLmUserContent(prep.text, prep, task), txt=>{ previewState.after += txt; updatePreview(); }, `${task} ${scope}`); previewState.after=sanitizeLmFinalText(previewState.after); updatePreview(); previewState.finished=true; previewState.running=false; const hasOutput=!!previewState.after.trim(); $('#applyPreviewBtn').disabled=!hasOutput; setStopButtons(false); setLmBusy(false); $('#streamState').textContent='done'; $('#previewStatus').textContent=hasOutput?'Generation complete.':'Generation produced no applicable text.'; $('#lmStatus').textContent=$('#previewStatus').textContent; if(hasOutput){ renderIntegrityGuardrails(previewState.before,previewState.after); reportPreviewMarkerDrift(); } }catch(err){ if(err?.name==='AbortError' || previewState.partial){ markLmStopped(); return; } toast(err.message); $('#lmStatus').textContent=err.message; finishPreviewError(err.message); } }
async function startGuidedAI(kind, scope){ if(previewState.running){ toast('Generation is already running. Stop it first.'); return; } try{ const range=getScopeRange(scope); if(!range.text.trim()) throw new Error('No text to send.'); const applyMode=(kind==='questions'||kind==='feedback')?'none':'replace'; const titles={questions:'Improve Questions',feedback:'Feedback Only','guided-rewrite':'Rewrite with Guidance',translate:'Translate Selection'}; openPreview(titles[kind]||'AI Tool', range.text, '', range.start, range.end, true, applyMode); previewState.expectedMarkers=applyMode==='replace'?markerSignature(range.text):null; await streamOne(guidedSystemPrompt(kind), guidedUserContent(kind, range.text), txt=>{ previewState.after += txt; updatePreview(); }, kind); previewState.after=sanitizeLmFinalText(previewState.after); updatePreview(); previewState.finished=true; previewState.running=false; const hasOutput=!!previewState.after.trim(); $('#applyPreviewBtn').disabled=applyMode==='none'||!hasOutput; setStopButtons(false); setLmBusy(false); $('#streamState').textContent='done'; $('#previewStatus').textContent=applyMode==='none'?'Guidance complete. Review only; not applicable.':(hasOutput?'Generation complete.':'Generation produced no applicable text.'); $('#lmStatus').textContent=$('#previewStatus').textContent; if(applyMode==='replace' && hasOutput){ renderIntegrityGuardrails(previewState.before,previewState.after); reportPreviewMarkerDrift(); } }catch(err){ if(err?.name==='AbortError' || previewState.partial){ markLmStopped(); return; } toast(err.message); $('#lmStatus').textContent=err.message; finishPreviewError(err.message); } }
function makeChunks(text, maxChars){ const lines=text.split('\n'); const chunks=[]; let buf=[]; let len=0; const push=()=>{ if(buf.length){ chunks.push(buf.join('\n')); buf=[]; len=0; } }; for(const line of lines){ const isChapter=/^\s{0,3}#\s+\S/.test(line); if(isChapter && len>0) push(); if(len+line.length+1>maxChars && len>0 && line.trim()==='') push(); if(len+line.length+1>maxChars*1.25 && len>0) push(); buf.push(line); len += line.length+1; } push(); return chunks; }
async function startLMChunked(task, scope, range, promptText){ const maxChars=Math.max(1200, parseInt($('#chunkSize').value,10)||6000); const chunks=makeChunks(range.text,maxChars); openPreview(`${task==='clean'?'AI Clean':'Rewrite'} chunked ${scope}`, range.text, '', range.start, range.end, true); previewState.expectedMarkers=markerSignature(range.text); setPreviewProgress(0,chunks.length); for(let i=0;i<chunks.length;i++){ if(!previewState.running) break; const pct=progressPercent(i,chunks.length); $('#streamState').textContent=`${pct}% · chunk ${i+1}/${chunks.length}`; $('#previewStatus').textContent=`Generating chunk ${i+1} of ${chunks.length} (${pct}% done)...`; if(i>0) previewState.after += '\n\n'; const prep=prepareTextForLm(chunks[i]); await streamOne(promptText, buildLmUserContent(prep.text, prep, task), txt=>{ previewState.after += txt; updatePreview(); }, `${task} chunk ${i+1}/${chunks.length}`); setPreviewProgress(i+1,chunks.length); }
 previewState.after=sanitizeLmFinalText(previewState.after); updatePreview(); previewState.finished=!previewState.partial && previewState.running; previewState.running=false; setStopButtons(false); setLmBusy(false); const hasOutput=!!previewState.after.trim(); const finalPct=previewState.progress?.percent ?? 0; $('#streamState').textContent=previewState.partial?`stopped · ${finalPct}%`:'done · 100%'; $('#applyPreviewBtn').disabled=!!previewState.partial || !hasOutput; $('#previewStatus').textContent=previewState.partial?`Stopped at ${finalPct}%. Partial output is not applicable.`:(hasOutput?'Chunked generation complete.':'Chunked generation produced no applicable text.'); $('#lmStatus').textContent=$('#previewStatus').textContent; if(!previewState.partial && hasOutput){ renderIntegrityGuardrails(previewState.before,previewState.after); reportPreviewMarkerDrift(); } }
async function streamOne(promptText, inputText, onText, label){ const provider=currentAiProvider(); lmAbort=new AbortController(); setStopButtons(true); $('#showPreviewBtn').disabled=false; const res=await fetch('/api/lm_stream',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({base_url:$('#lmBase').value,provider:$('#lmProvider').value,api_key:$('#lmApiKey').value,model:$('#lmModel').value,prompt:promptText,input:inputText,options:advancedOptions()}),signal:lmAbort.signal}); if(!res.ok){ const msg=await res.text(); throw new Error(msg||`${provider.label} request failed`); } const reader=res.body.getReader(); const decoder=new TextDecoder(); while(true){ const {value,done}=await reader.read(); if(done) break; const chunk=decoder.decode(value,{stream:true}); if(chunk) onText(chunk); } }
function previewSubText(){ const base=`${previewState.before.length.toLocaleString()} chars → ${previewState.after.length.toLocaleString()} chars`; return previewState.progress ? `${base} · ${previewState.progress.percent}% done` : base; }
function progressPercent(done,total){ return total ? Math.max(0,Math.min(100,Math.round((done/total)*100))) : 0; }
function setPreviewProgress(done,total){ previewState.progress={done,total,percent:progressPercent(done,total)}; $('#previewSub').textContent=previewSubText(); }
function openPreview(title,before,after,start,end,running,applyMode='replace'){ previewState={title,before,after,start,end,finished:!running,running,partial:false,expectedMarkers:null,applyMode,progress:null}; clearIntegrityGuardrails(); $('#previewTitle').textContent=title; $('#previewSub').textContent=previewSubText(); $('#beforeBox').textContent=before; $('#afterBox').textContent=after; $('#previewModal').classList.add('open'); setPreviewView(localStorage.getItem('vp.py.previewView')||'split'); $('#applyPreviewBtn').disabled=running || !after || applyMode==='none'; $('#comparePreviewBtn').disabled=!before&&!after; setStopButtons(running); setLmBusy(running); $('#showPreviewBtn').disabled=false; $('#streamState').textContent=running?'generating...':'ready'; $('#previewStatus').textContent=running?'Generating...':(applyMode==='none'?'Review ready.':'Preview ready.'); }
function updatePreview(){ $('#afterBox').textContent=previewState.after; $('#previewSub').textContent=previewSubText(); $('#comparePreviewBtn').disabled=!previewState.before&&!previewState.after; if($('#autoScroll').checked){ const b=$('#afterBox'); b.scrollTop=b.scrollHeight; } }
function finishPreviewError(msg){ previewState.running=false; setStopButtons(false); setLmBusy(false); $('#streamState').textContent='error'; $('#previewStatus').textContent=msg; }
function syncPreviewScroll(source,target){ if(previewScrollSyncing || !$('#previewLinked').checked || !target) return; previewScrollSyncing=true; const maxSource=Math.max(1,source.scrollHeight-source.clientHeight); const maxTarget=Math.max(0,target.scrollHeight-target.clientHeight); target.scrollTop=(source.scrollTop/maxSource)*maxTarget; requestAnimationFrame(()=>{ previewScrollSyncing=false; }); }
$('#showPreviewBtn').addEventListener('click',()=>{ $('#previewModal').classList.add('open'); updatePreview(); });
$('#closePreviewBtn').addEventListener('click',()=>$('#previewModal').classList.remove('open'));
$('#comparePreviewBtn').addEventListener('click',()=>openCompare(previewState.before||'',previewState.after||'','Loaded preview into Compare.'));
$('#applyPreviewBtn').addEventListener('click',()=>{ if(previewState.applyMode==='none'){ toast('This preview is review-only.'); return; } if(!previewState.finished || previewState.partial) return; reportPreviewMarkerDrift(); replaceRange(previewState.start,previewState.end,previewState.after); $('#previewModal').classList.remove('open'); previewState.after=''; $('#showPreviewBtn').disabled=true; toast('Applied'); });
function setStopButtons(enabled){ $('#stopBtn').disabled=!enabled; $('#lmStopBtn').disabled=!enabled; }
function markLmStopped(){ previewState.partial=true; previewState.running=false; setStopButtons(false); setLmBusy(false); $('#applyPreviewBtn').disabled=true; $('#streamState').textContent='stopped'; $('#previewStatus').textContent='Stopped. Partial output is viewable but not applicable.'; $('#lmStatus').textContent='Stopped; partial preview available.'; }
function stopLmGeneration(){ markLmStopped(); if(lmAbort) lmAbort.abort(); }
$('#stopBtn').addEventListener('click',stopLmGeneration);
$('#lmStopBtn').addEventListener('click',stopLmGeneration);
$('#bottomBtn').addEventListener('click',()=>{ $('#autoScroll').checked=true; const b=$('#afterBox'); b.scrollTop=b.scrollHeight; });
$('#afterBox').addEventListener('scroll',()=>{ const b=$('#afterBox'); if(b.scrollTop + b.clientHeight < b.scrollHeight - 40) $('#autoScroll').checked=false; });
$('#beforeBox').addEventListener('scroll',()=>syncPreviewScroll($('#beforeBox'),$('#afterBox')));
$('#afterBox').addEventListener('scroll',()=>syncPreviewScroll($('#afterBox'),$('#beforeBox')));
document.querySelectorAll('.viewBtn').forEach(b=>b.addEventListener('click',()=>setPreviewView(b.dataset.view)));
function setPreviewView(view){ if(!['split','after-only','before-only'].includes(view)) view='split'; const modal=$('#previewModal'), card=$('#modalCard'); [modal,card].forEach(el=>{ el.classList.remove('after-only','before-only'); if(view==='after-only') el.classList.add('after-only'); if(view==='before-only') el.classList.add('before-only'); }); document.querySelectorAll('.viewBtn').forEach(btn=>btn.classList.toggle('active',btn.dataset.view===view)); localStorage.setItem('vp.py.previewView',view); }

function scheduleAutosave(){ if(recoveryPending && !editor.value){ return; } clearTimeout(autosaveTimer); autosaveTimer=setTimeout(()=>{ if(recoveryPending && !editor.value) return; try{ const settings=getSettings(); delete settings.lmApiKey; const data={version:1,text:editor.value,metadata:getMetadata(),settings,savedAt:new Date().toISOString()}; localStorage.setItem('vp.py.autosave', JSON.stringify(data)); }catch(err){ if(!autosaveWarningShown){ autosaveWarningShown=true; toast('Autosave failed: browser storage is full. Export your work to disk.'); } } },700); }
function getSettings(){ return {wrap:$('#wrapToggle').checked, previewOn:document.querySelector('.editor-card')?.classList.contains('preview-on')||false, scope:getScope(), vellumCollapsed:$('#vellumToolbar').classList.contains('collapsed'), sideTab:document.querySelector('[data-side-tab].active')?.dataset.sideTab||'check', lmProvider:$('#lmProvider')?.value||'lmstudio', lmBase:$('#lmBase').value, lmApiKey:$('#lmApiKey').value, lmModel:$('#lmModel').value, lmStrength:$('#lmStrength')?.value||'medium', chunkSize:$('#chunkSize').value, guidedContext:$('#guidedContext')?.value||'', translateTarget:$('#translateTarget')?.value||'English', advTemperature:$('#advTemperature')?.value||'', advTopP:$('#advTopP')?.value||'', advTopK:$('#advTopK')?.value||'', advRepeatPenalty:$('#advRepeatPenalty')?.value||'', advMaxTokens:$('#advMaxTokens')?.value||''}; }
function saveSettings(){ try{ localStorage.setItem('vp.py.settings', JSON.stringify(getSettings())); }catch(err){} scheduleAutosave(); }
function loadSettings(){ let s={}; try{s=JSON.parse(localStorage.getItem('vp.py.settings')||'{}');}catch(e){} $('#wrapToggle').checked=!!s.wrap; $('#wrapToggle').dispatchEvent(new Event('change')); const card=document.querySelector('.editor-card'); card.classList.toggle('preview-on',!!s.previewOn); $('#togglePreview').classList.toggle('active',!!s.previewOn); const scopeInput=document.querySelector(`input[name="scope"][value="${s.scope==='selection'?'selection':'all'}"]`); if(scopeInput) scopeInput.checked=true; const tb=$('#vellumToolbar'); const collapsed=s.vellumCollapsed!==false; tb.classList.toggle('collapsed',collapsed); $('#toggleVellum').textContent=collapsed?'Expand':'Collapse'; const savedTab=s.sideTab||localStorage.getItem('vp.py.sideTab')||'check'; activateSideTab(document.querySelector(`[data-side-tab="${savedTab}"]`)?savedTab:'check'); if(s.lmProvider && AI_PROVIDERS[s.lmProvider]) $('#lmProvider').value=s.lmProvider; if(s.lmBase) $('#lmBase').value=s.lmBase; else if($('#lmProvider').value!=='custom') $('#lmBase').value=currentAiProvider().base; if(s.lmApiKey) $('#lmApiKey').value=s.lmApiKey; updateProviderUi(); if(s.lmBase) $('#lmBase').value=s.lmBase; if(s.lmModel){ const sel=$('#lmModel'); if(!Array.from(sel.options).some(o=>o.value===s.lmModel)){ const opt=document.createElement('option'); opt.value=s.lmModel; opt.textContent=s.lmModel; sel.appendChild(opt); } sel.value=s.lmModel; } if(s.lmStrength) $('#lmStrength').value=s.lmStrength; if(s.chunkSize) $('#chunkSize').value=s.chunkSize; if(s.guidedContext) $('#guidedContext').value=s.guidedContext; if(s.translateTarget) $('#translateTarget').value=s.translateTarget; ['advTemperature','advTopP','advTopK','advRepeatPenalty','advMaxTokens'].forEach(id=>{ if(s[id]) $('#'+id).value=s[id]; }); }
['metaTitle','metaSubtitle','metaAuthor','metaLanguage','metaSeries','metaBookNumber','metaPublisher','metaFilename','lmBase','lmApiKey','lmModel','chunkSize','guidedContext','translateTarget','advTemperature','advTopP','advTopK','advRepeatPenalty','advMaxTokens'].forEach(id=>$('#'+id).addEventListener('input',()=>{ if(id==='lmBase' && $('#lmProvider').value!=='custom'){ $('#lmProvider').value='custom'; updateProviderUi(); } saveSettings(); updateStats(); }));
$('#lmProvider').addEventListener('change',()=>{ updateProviderUi(); saveSettings(); });
$('#lmStrength').addEventListener('change',saveSettings);
$('#lmModel').addEventListener('change',saveSettings);
function showRecoveryIfNeeded(){ const raw=localStorage.getItem('vp.py.autosave'); if(!raw) return; try{ const data=JSON.parse(raw); if(data.text && !editor.value){ recoveryData=data; recoveryPending=true; $('#recovery').classList.add('show'); } }catch(e){} }
$('#restoreBtn').addEventListener('click',()=>{ const data=recoveryData||JSON.parse(localStorage.getItem('vp.py.autosave')||'{}'); recoveryPending=false; editor.value=data.text||''; setMetadata(data.metadata||{}); if(data.settings){ try{ localStorage.setItem('vp.py.settings',JSON.stringify(data.settings)); loadSettings(); }catch(e){} } $('#recovery').classList.remove('show'); updateStats(); toast(`Autosave restored: ${(editor.value.match(/\S+/g)||[]).length.toLocaleString()} words`); });
$('#discardRestoreBtn').addEventListener('click',()=>{ recoveryPending=false; recoveryData=null; localStorage.removeItem('vp.py.autosave'); $('#recovery').classList.remove('show'); updateStats(); });
window.addEventListener('resize',syncFindHighlights);
loadPrompts(); loadSettings(); showRecoveryIfNeeded(); updateStats();
</script>
</body>
</html>
'''


class ReusableThreadingHTTPServer(ThreadingHTTPServer):
    allow_reuse_address = True


class ManuscriptWorkbenchHandler(BaseHTTPRequestHandler):
    server_version = "ManuscriptWorkbenchHTTP/1.0"

    def log_message(self, fmt: str, *args: Any) -> None:
        sys.stderr.write("[%s] %s\n" % (_dt.datetime.now().strftime("%H:%M:%S"), fmt % args))

    def do_OPTIONS(self) -> None:
        self.send_response(204)
        self.end_headers()

    def do_GET(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path in ("/", "/index.html"):
            text_response(self, HTML, content_type="text/html; charset=utf-8")
        elif parsed.path == "/api/health":
            json_response(self, {"ok": True, "app": APP_NAME, "python_docx": Document is not None, "requests": requests is not None, "markdown": markdown_lib is not None, "lxml": lxml_html is not None, "trafilatura": trafilatura is not None})
        else:
            text_response(self, "Not found", HTTPStatus.NOT_FOUND)

    def do_POST(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        self._streaming_response_started = False
        try:
            if parsed.path == "/api/import":
                payload = get_json_body(self)
                filename = payload.get("filename") or "untitled.txt"
                data_b64 = payload.get("data_b64") or ""
                data = base64.b64decode(data_b64)
                json_response(self, import_file(filename, data, payload.get("html_mode") or "simple"))
            elif parsed.path == "/api/import_html_options":
                payload = get_json_body(self)
                filename = payload.get("filename") or "import.html"
                data_b64 = payload.get("data_b64") or ""
                data = base64.b64decode(data_b64)
                decoded, notes = decode_text_bytes(data)
                result = html_import_options(decoded, filename=filename)
                result["notes"] = notes + result.get("notes", [])
                json_response(self, result)
            elif parsed.path == "/api/import_url":
                payload = get_json_body(self)
                data, final_url, content_type = fetch_url_bytes(payload.get("url") or "")
                decoded, notes = decode_text_bytes(data)
                parsed_url = urllib.parse.urlparse(final_url)
                path_name = Path(urllib.parse.unquote(parsed_url.path or "")).name or (parsed_url.hostname or "import")
                filename = safe_filename(path_name.rsplit(".", 1)[0] or parsed_url.hostname or "url-import") + ".md"
                if "html" in content_type.lower() or re.search(r"<(?:!doctype\s+html|html|head|body|article|main)\b", decoded, re.I):
                    result = html_import_options(decoded, filename=filename, url=final_url)
                else:
                    text = normalize_imported_text(decoded)
                    result = {
                        "filename": filename,
                        "url": final_url,
                        "simple": {"text": text, "chars": len(text)},
                        "main": {"text": "", "chars": 0, "error": "URL did not look like HTML."},
                        "notes": ["Imported URL as text"],
                        "engine": "text",
                    }
                result["content_type"] = content_type
                result["notes"] = notes + result.get("notes", [])
                json_response(self, result)
            elif parsed.path == "/api/render_markdown":
                payload = get_json_body(self)
                json_response(self, {"html": render_markdown_html(payload.get("text") or ""), "engine": "python-markdown"})
            elif parsed.path == "/api/export_docx":
                payload = get_json_body(self)
                text = payload.get("text") or ""
                metadata = payload.get("metadata") or {}
                docx_bytes = create_manuscript_docx(text, metadata)
                filename = safe_filename(payload.get("filename") or metadata.get("filename") or "manuscript") + "-formatted.docx"
                self.send_response(200)
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
                self.send_header("Content-Length", str(len(docx_bytes)))
                self.end_headers()
                self.wfile.write(docx_bytes)
            elif parsed.path == "/api/preflight":
                payload = get_json_body(self)
                json_response(self, preflight(payload.get("text") or "", payload.get("metadata") or {}))
            elif parsed.path == "/api/lm_models":
                self.handle_lm_models()
            elif parsed.path == "/api/lm_stream":
                self.handle_lm_stream()
            else:
                text_response(self, "Not found", HTTPStatus.NOT_FOUND)
        except Exception as exc:
            traceback.print_exc()
            if getattr(self, "_streaming_response_started", False):
                return
            json_response(self, {"error": str(exc)}, status=500)

    def handle_lm_models(self) -> None:
        if requests is None:
            raise RuntimeError("requests is not installed. Run: pip install -r requirements.txt")
        payload = get_json_body(self)
        base_url = normalize_base_url(payload.get("base_url") or "", payload.get("provider") or "lmstudio")
        resp = requests.get(f"{base_url}/models", headers=ai_request_headers(payload.get("api_key") or ""), timeout=8)
        resp.raise_for_status()
        data = resp.json()
        models = [m.get("id") for m in data.get("data", []) if m.get("id")]
        json_response(self, {"models": models, "base_url": base_url})

    def handle_lm_stream(self) -> None:
        if requests is None:
            raise RuntimeError("requests is not installed. Run: pip install -r requirements.txt")
        payload = get_json_body(self)
        base_url = normalize_base_url(payload.get("base_url") or "", payload.get("provider") or "lmstudio")
        model = payload.get("model") or ""
        if not model.strip():
            try:
                resp = requests.get(f"{base_url}/models", headers=ai_request_headers(payload.get("api_key") or ""), timeout=6)
                resp.raise_for_status()
                models = resp.json().get("data") or []
                if models:
                    model = models[0].get("id") or ""
            except Exception:
                pass
        model = model.strip() or "local-model"
        system_prompt = payload.get("prompt") or "Rewrite the text. Output only the result."
        input_text = payload.get("input") or ""
        body = {
            "model": model,
            "stream": True,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": input_text},
            ],
        }
        options = payload.get("options") or {}
        if isinstance(options, dict):
            for key in ("temperature", "top_p", "top_k", "repeat_penalty", "max_tokens"):
                value = options.get(key)
                if value is not None and value != "":
                    body[key] = value
        with requests.post(f"{base_url}/chat/completions", json=body, headers=ai_request_headers(payload.get("api_key") or ""), stream=True, timeout=(10, 300)) as resp:
            if resp.status_code >= 400:
                raise RuntimeError(resp.text[:4000])
            self.send_response(200)
            self.send_header("Content-Type", "text/plain; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.end_headers()
            self._streaming_response_started = True
            in_think = False
            def strip_think_chunk(content: str) -> str:
                nonlocal in_think
                out: List[str] = []
                i = 0
                lower = content.lower()
                while i < len(content):
                    if in_think:
                        end = lower.find("</think>", i)
                        if end < 0:
                            return "".join(out)
                        i = end + len("</think>")
                        in_think = False
                    else:
                        start = lower.find("<think>", i)
                        if start < 0:
                            out.append(content[i:])
                            break
                        out.append(content[i:start])
                        i = start + len("<think>")
                        in_think = True
                return "".join(out)
            for raw in resp.iter_lines(decode_unicode=False):
                if not raw:
                    continue
                line = raw.decode("utf-8", errors="replace").strip()
                if line.startswith("data:"):
                    line = line[5:].strip()
                if line == "[DONE]":
                    break
                try:
                    data = json.loads(line)
                    choice = (data.get("choices") or [{}])[0]
                    delta = choice.get("delta") or {}
                    content = delta.get("content")
                    if content is None:
                        msg = choice.get("message") or {}
                        content = msg.get("content")
                    if content:
                        content = strip_think_chunk(content)
                    if content:
                        self.wfile.write(content.encode("utf-8"))
                        self.wfile.flush()
                except (json.JSONDecodeError, BrokenPipeError):
                    if isinstance(sys.exc_info()[1], BrokenPipeError):
                        break
                    continue


def main() -> int:
    missing = []
    if Document is None:
        missing.append("python-docx")
    if requests is None:
        missing.append("requests")
    if markdown_lib is None:
        missing.append("markdown")
    if missing:
        print("Missing dependencies:", ", ".join(missing))
        print("Run: pip install -r requirements.txt")
        return 2
    url = f"http://{APP_HOST}:{APP_PORT}"
    try:
        server = ReusableThreadingHTTPServer((APP_HOST, APP_PORT), ManuscriptWorkbenchHandler)
    except OSError as exc:
        if exc.errno == errno.EADDRINUSE:
            print(f"{APP_NAME}")
            print(f"Port {APP_PORT} is already in use.")
            print(f"If Manuscript Workbench is already open, use: {url}")
            print(f"To start another copy, run: MANUSCRIPT_WORKBENCH_PORT={APP_PORT + 1} ./start_mac_linux.sh")
            return 1
        raise
    print(f"{APP_NAME}")
    print(f"Open: {url}")
    print("Press Ctrl+C to stop.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping.")
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
